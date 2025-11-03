import os
import sys
import json
import argparse
import traceback
import inspect
import pandas as pd
import re
from docx.table import _Cell
from docx.table import Table
# ========= FILL HELPERS (add to this file) =========
from docx import Document

_P_PLACEHOLDER = re.compile(r"^[\s\u2000-\u200B\u00A0._\-–—]{3,}$")
_P_HAS_ALNUM   = re.compile(r"[A-Za-z0-9]")

def _norm_key(s: str) -> str:
    t = re.sub(r"\s+", " ", str(s or "")).strip().lower()
    t = t.replace("—", "-").replace("–", "-")
    # drop trailing colon-like
    t = re.sub(r"[:：﹕꞉˸፡︓]\s*$", "", t)
    # collapse punctuation that often differs slightly
    t = re.sub(r"[ \t]+", " ", t)
    return t

def _para_text(p) -> str:
    return (p.text or "").strip()

def _is_placeholder_like(s: str) -> bool:
    s = (s or "").strip()
    if not s:
        return True
    if len(s) <= 2:
        return True
    if _P_PLACEHOLDER.match(s) and not _P_HAS_ALNUM.search(s):
        return True
    return False

def _write_value_into_paragraph_after_label(par, value: str):
    """
    If there's no explicit placeholder paragraph after label,
    inline the value after the label text with a separating space.
    """
    txt = par.text
    if not txt.endswith(" "):
        txt = txt + " "
    for r in par.runs:
        r.clear()  # nuke existing runs to avoid weird styling splits
    par.add_run(txt + str(value))

def _toggle_checkbox_glyphs_in_text(raw_text: str, val: str) -> str:
    """
    Toggle only EXISTING checkboxes. If val is truthy (yes/true/1/x),
    turn first 'unchecked' into '☒' or '☑'. If falsy, prefer '☐'.
    Does not create or delete glyphs.
    """
    truthy = str(val or "").strip().lower() in {"true", "yes", "y", "1", "x", "✓", "checked"}
    # Replace in a conservative, left-to-right way
    text = raw_text

    # If there is at least one unchecked box, and value is truthy, flip the FIRST one.
    if truthy and "☐" in text and ("☒" not in text and "☑" not in text):
        text = text.replace("☐", "☒", 1)
    # If falsy and there is a checked one, flip the FIRST checked back to unchecked.
    elif (not truthy) and ("☒" in text or "☑" in text):
        text = text.replace("☒", "☐", 1).replace("☑", "☐", 1)

    # Also handle ASCII-style [ ] / [x]
    if truthy and "[ ]" in text and ("[x]" not in text and "[X]" not in text):
        text = text.replace("[ ]", "[x]", 1)
    elif (not truthy) and ("[x]" in text or "[X]" in text):
        text = text.replace("[x]", "[ ]", 1).replace("[X]", "[ ]", 1)

    return text

def _set_paragraph_text(par, new_text: str):
    # Replace all runs with a single run containing new_text.
    for r in par.runs:
        r.clear()
    par.clear()
    par.add_run(new_text)

# --- Inline checkbox splitting helpers (generic, pattern-based) ---
_BOX_TOKEN_RE = re.compile(r'(\[\s*[xX✓]?\s*\]|[□☐☒])')
_INLINE_SEP_RE = re.compile(r'(?:\t+|\s{2,}|\s+/\s+|\s+\|\s+|\s+OR\s+)', re.IGNORECASE)
# Strict checkbox detection: only mark checkboxes when a real checkbox control is found in the DOCX.
STRICT_CHECKBOX_DETECTION = True

def _strict() -> bool:
    """Single source of truth for 'strict' checks throughout the module."""
    try:
        return bool(globals().get("STRICT_CHECKBOX_DETECTION", False))
    except Exception:
        return False

# ---------- PDF row emission helpers (new) ----------

_RECIP_SEC_RE = re.compile(r"\brecipient\s*(\d+)\b", re.IGNORECASE)

def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", str(s or "").strip())

def _short_label_like(s: str, max_words=10, max_len=140) -> bool:
    t = _norm(s)
    if not t or len(t) > max_len:
        return False
    # avoid giant sentences with punctuation
    if t.count(".") >= 2:
        return False
    return len(t.split()) <= max_words

def _token_len(s: str) -> int:
    return len(re.findall(r"[A-Za-z0-9]", s or ""))

def _infer_recipient_index(section: str, field: str, index: int) -> int:
    """
    Prefer index inferred from 'Recipient N' in Section or Field.
    """
    for src in (section, field):
        m = _RECIP_SEC_RE.search(src or "")
        if m:
            try:
                n = int(m.group(1))
                return max(1, min(4, n))
            except Exception:
                pass
    return index or 1

def _is_pdf_fillable_fdef(fdef: dict) -> tuple[bool, str]:
    """
    Return (should_emit, type_tag) where type_tag ∈ {'checkbox','underline','table','dropdown','text',''}
    Uses fdef['type'], fdef['field_type'], fdef['placement'], fdef['choices'].
    """
    placement = _norm(fdef.get("placement", "")).lower()
    kind      = _norm(fdef.get("type", "") or fdef.get("field_type", "")).lower()
    label     = _norm(fdef.get("label_short") or fdef.get("label") or fdef.get("label_full") or "")

    # Widgets
    if "check" in kind or "acro_check" in placement or "acro-checkbox" in placement:
        return True, "checkbox"
    if "choice" in kind or "combo" in kind or "list" in kind or "acro_choice" in placement:
        return True, "dropdown"
    if ("text" in kind and "check" not in kind) or "acro_text" in placement:
        return (_short_label_like(label), "text")

    # Inline/visual checkbox hints
    if fdef.get("choices"):
        return True, "checkbox"
    if re.search(r"(?:\[\s*\]|\[\s*[xX✓]\s*\]|[□☐☑☒])", label):
        return True, "checkbox"

    # Underlines / blanks / rules
    if any(k in placement for k in ("underline", "line_rule", "line", "blank", "rule")):
        return (_short_label_like(label), "underline")

    # Table/grids
    if any(k in placement for k in ("grid_header", "table_header", "grid", "table")):
        return (_short_label_like(label, max_words=6, max_len=60), "table")

    # Rect boxes near labels → treat as text if label is short-ish
    if any(k in placement for k in ("rect_box", "box", "square")) and _short_label_like(label, 8, 80):
        return True, "text"

    return False, ""


def _final_split_compound_checkbox_fields(df_in: pd.DataFrame) -> pd.DataFrame:
    """
    After all other expansions, split any single 'Field' that is really a
    compound checkbox line (e.g., 'Initial Subscription Additional Subscription')
    into separate rows, one per option, keeping Section/Page and setting Index=1.
    """
    additions = []
    drop_idx = []

    for i, r in df_in.iterrows():
        field_text = str(r.get("Field", "")).strip()
        chunks = _split_compound_checkboxish(field_text)
        if not chunks:
            continue

        drop_idx.append(i)
        for ch in chunks:
            additions.append({
                "Section": r.get("Section", ""),
                "Page": r.get("Page", ""),
                "Field": ch,
                "Index": 1,
                "Value": "",
                "Choices": (r.get("Choices", "") or "checkbox")
            })

    if not additions:
        return df_in

    df_out = df_in.drop(index=drop_idx).reset_index(drop=True)
    df_out = pd.concat([df_out, pd.DataFrame(additions, columns=df_out.columns)], ignore_index=True)
    return df_out


def _split_compound_checkboxish(label: str) -> list:
    """
    Split lines like 'Initial Subscription Additional Subscription' into
    ['Initial Subscription', 'Additional Subscription'].

    Defensive heuristics; now tolerant to multiple/nbsp spaces.
    """
    t = (label or "").strip()
    if not t:
        return []
    if ":" in t:
        return []
    if len(t) > 60:
        return []
    if re.search(r"\d", t):  # avoid account numbers etc.
        return []

    # normalize whitespace (collapses multiple spaces / NBSP to single space)
    t_norm = re.sub(r"\s+", " ", t.replace("\u00A0", " ")).strip()
    words = t_norm.split()
    if len(words) < 3:  # need at least two chunks of 1–4 words each
        return []

    def _is_titleish(w: str) -> bool:
        # Accept TitleCase (Bank, State) or ALLCAPS (USA)
        return (w[:1].isupper() and (w[1:].islower() or w[1:] == "")) or w.isupper()

    if not all(_is_titleish(w) for w in words):
        return []

    # Greedy grouping into 1–4-word chunks
    chunks, cur = [], []
    for w in words:
        cur.append(w)
        if 2 <= len(cur) <= 4:
            chunks.append(" ".join(cur))
            cur = []
    if cur:
        if chunks and len(cur) <= 2:
            chunks[-1] = f"{chunks[-1]} {' '.join(cur)}"
            cur = []
        else:
            return []

    # Need 2–4 chunks, and each chunk 1–4 words
    if len(chunks) < 2 or len(chunks) > 4:
        return []
    if not all(1 <= len(c.split()) <= 4 for c in chunks):
        return []

    # Relaxed guard: match after whitespace normalization OR ensure chunks appear in order
    combined = " ".join(chunks)
    if combined != t_norm:
        # in-order containment fallback
        pos, ok = 0, True
        key = t_norm.lower()
        for c in chunks:
            ck = c.lower()
            j = key.find(ck, pos)
            if j < 0:
                ok = False
                break
            pos = j + len(ck)
        if not ok:
            return []

    return chunks



def _looks_short_option(s: str, max_words: int = 5) -> bool:
    return len((s or "").split()) <= max_words


def _is_title_case_phrase(s: str) -> bool:
    # Accept phrases where most words start uppercase (Title-Case-ish).
    words = (s or "").split()
    if not words:
        return False
    caps = sum(1 for w in words if w[:1].isupper())
    return caps >= max(1, int(0.6 * len(words)))


def _extract_inline_checkbox_options(text: str):
    """
    Extract inline options from a single-line label such as:
      '□ Initial Subscription    □ Additional Subscription'
      '[ ] Initial Subscription   [ ] Additional Subscription'
      'Initial Subscription    Additional Subscription' (2+ spaces)
      'Initial Subscription OR Additional Subscription'
      'Initial Subscription Additional Subscription'  <-- (title-case split)
    Returns list[str] or None.
    """
    if not text:
        return None

    # 0) Normalize a cheap test string
    t_norm = re.sub(r'\s+', ' ', text).strip()

    # 1) Split on explicit checkbox tokens (□, ☐, ☒, [ ], [x], etc.)
    t = _BOX_TOKEN_RE.sub(' ::: ', text)
    t = re.sub(r'\s+', ' ', t).strip()
    parts = [p.strip(' -:') for p in t.split(':::') if p and p.strip()]
    parts = [p for p in parts if _looks_short_option(p)]
    if len(parts) >= 2 and len(set(p.lower() for p in parts)) >= 2:
        return parts

    # 2) Fallback: split on tabs, 2+ spaces, '/', '|', or ' OR '
    parts = [p.strip(' -:') for p in _INLINE_SEP_RE.split(text) if p and p.strip()]
    parts = [p for p in parts if _looks_short_option(p)]
    if 2 <= len(parts) <= 8 and len(set(p.lower() for p in parts)) >= 2:
        return parts

    # 3) Title-case chunk heuristic (handles "Initial Subscription Additional Subscription")
    if ':' not in t_norm:
        words = t_norm.split()
        if 2 <= len(words) <= 8 and _is_title_case_phrase(t_norm):
            # Equal halves first
            if len(words) % 2 == 0:
                mid = len(words) // 2
                left = " ".join(words[:mid]).strip(' -:')
                right = " ".join(words[mid:]).strip(' -:')
                if _looks_short_option(left) and _looks_short_option(right) and left.lower() != right.lower():
                    return [left, right]

            # Otherwise, split before a capitalized word
            for i in range(2, len(words) - 1):
                if words[i][0].isupper():
                    left = " ".join(words[:i]).strip(' -:')
                    right = " ".join(words[i:]).strip(' -:')
                    if (_looks_short_option(left) and _looks_short_option(right)
                            and _is_title_case_phrase(left) and _is_title_case_phrase(right)
                            and left.lower() != right.lower()):
                        return [left, right]

    return None

PLACEHOLDER_RE = re.compile(r"^[_\-\u2014\.\s\u2002\u2003\u2007\u2009\u00A0]{3,}$")
ALNUM_RE = re.compile(r"[A-Za-z0-9]")

def _cell_has_placeholder_like(cell: _Cell) -> bool:
    # True if cell contains grey-box/underscore-only runs or paragraph shading
    try:
        if cell._tc.xpath('.//w:tcPr/w:shd', namespaces=cell._tc.nsmap):
            return True
    except Exception:
        pass
    for p in cell.paragraphs:
        # paragraph shading
        try:
            if p._element.xpath('.//w:pPr/w:shd'):
                return True
        except Exception:
            pass
        # any run that looks like placeholder only
        for r in p.runs:
            t = (r.text or "").strip()
            if t and not ALNUM_RE.search(t) and len(t.replace(" ", "")) >= 3:
                return True
    return False

def _extract_label_from_same_cell(cell: _Cell) -> str:
    """
    Heuristic:
    - If cell has multiple paragraphs and the *next* paragraph/run looks like placeholder,
      use the text of the *previous* paragraph as label.
    - Else, if a paragraph mixes text + placeholder in the same line, take the left text.
    - Else, fall back to the first non-empty alphanumeric line.
    """
    # pass 1: paragraph followed by placeholder paragraph
    pars = cell.paragraphs
    for i, p in enumerate(pars):
        txt = (p.text or "").strip()
        if not txt:
            continue
        # if *next* paragraph is placeholder-like, this is likely the label
        if i + 1 < len(pars):
            nxt = (pars[i+1].text or "").strip()
            if nxt and not ALNUM_RE.search(nxt):
                return re.sub(r"[:：﹕꞉˸፡︓]\s*$", "", txt).strip()

        # inline placeholder in the same paragraph
        # split when we find a long whitespace/underscore/dash cluster
        m = re.search(r"([ \t\u00A0\u2002\u2003\u2007\u2009]{3,}|_{3,}|[\u2014\-]{3,})", txt)
        if m:
            left = txt[:m.start()].strip()
            if left:
                return re.sub(r"[:：﹕꞉˸፡︓]\s*$", "", left).strip()

    # pass 2: first line that has letters/digits
    for p in pars:
        txt = (p.text or "").strip()
        if txt and ALNUM_RE.search(txt):
            return re.sub(r"[:：﹕꞉˸፡︓]\s*$", "", txt).strip()
    return ""

def _row0_has_contact_headers(tbl: Table) -> bool:
    try:
        hdrs = [(_cell_text(c) or "").strip() for c in tbl.rows[0].cells]
    except Exception:
        return False
    # Look for at least two columns whose text looks like "Contact N"
    pat = re.compile(r"^contact\s*\d+\b[:)?]?$", re.IGNORECASE)
    hits = sum(1 for h in hdrs if h and pat.search(h))
    return hits >= 2



# -----------------------------------------------------------------------------
# Try to find build template no matter where you place this helper.
# -----------------------------------------------------------------------------
_build_pdf_template_external = None
_build_pdf_template_import_error = None
_build_pdf_template_import_tb = None

# Optional DOCX builder import
_docx_builder = None
try:
    # Most repos expose this as build_docs_template
    from docx_prefill import build_docs_template as _docx_builder  # type: ignore
except Exception:
    try:
        # Back-compat: some repos call it build_docx_template
        from docx_prefill import build_docx_template as _docx_builder  # type: ignore
    except Exception:
        _docx_builder = None

# Optional PDF builder import (legacy/default)
try:
    from pdf_prefill import build_pdf_template as _build_pdf_template_external  # type: ignore
except Exception as _e:
    _build_pdf_template_external = None
    _build_pdf_template_import_error = _e
    _build_pdf_template_import_tb = traceback.format_exc()


def _describe_builder(fn) -> str:
    try:
        mod = getattr(fn, "__module__", None)
        src = inspect.getsourcefile(fn) or inspect.getfile(fn)
        return f"{mod or '<unknown module>'} :: {src}"
    except Exception:
        return "<uninspectable builder>"



# ---- Minimal local DOCX builder (only used when input is .docx) ----
# ---- Minimal local DOCX builder (only used when input is .docx) ----
# ---- Minimal local DOCX builder (only used when input is .docx) ----
def build_docx_template(input_path: str, out_json: str | None = None) -> dict:
    import re, json, unicodedata
    from collections import defaultdict
    from docx import Document
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.text.paragraph import Paragraph
    from docx.table import Table, _Cell  # ensure _Cell is imported

    doc = Document(input_path)

    def _norm_space(s: str) -> str:
        return re.sub(r"\s+", " ", (s or "").strip())

    def _strip_colon_like(s: str) -> str:
        return re.sub(r"[:：﹕꞉˸]\s*$", "", (s or "").strip())

    def _looks_like_section_title(text: str) -> bool:
        if not text: return False
        t = unicodedata.normalize("NFKC", text).strip()
        if len(t) > 180: return False
        if t.endswith(":"): return True
        FIELD_LABEL_DENY = {
            "Applicant Name","Last name","First name","Middle Initial",
            "Date of Birth","Place of Birth","Citizen of","Country of Residence",
            "Tax Identification Number","GIIN (if applicable)","Jurisdiction of Organisation",
            "Residence (Individual Applicants) or Registered Address (Entity Applicants)",
            "Street 1","Street 2","Street 3","City, ST Zip","Country",
            "Work Phone","Home Phone","Cell Phone","Fax","Email","E-Mail","Name",
            "Principal Place of Business (if differs from Registered Address)","Signature",
        }
        if t in FIELD_LABEL_DENY: return False
        words = t.split()
        letters = [c for c in t if c.isalpha()]
        caps_ratio = (sum(1 for c in letters if c.isupper()) / max(1, len(letters))) if letters else 0.0
        KNOWN_BLOCKS = {
            "Individual Applicants","Entity Applicants","All Applicants",
            "Authorised Signatories","Authorized Signatories",
            "Investor Registration Form","Investor Profile Form",
        }
        if t in KNOWN_BLOCKS: return True
        return caps_ratio >= 0.85 and len(words) <= 15

    def _iter_blocks(d):
        parent_elm = d._element.body
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P): yield Paragraph(child, d)
            elif isinstance(child, CT_Tbl): yield Table(child, d)

    def _cell_text(cell) -> str:
        return "\n".join(p.text for p in cell.paragraphs).strip()

    def _header_for_col(tbl: Table, ci: int) -> str:
        """
        Resolve the 'Contact N' header for a given column by scanning
        the first few rows (some templates don't put the header strictly in row 0).
        """
        pat = re.compile(r"^contact\s*\d+\b[:)?]?$", re.IGNORECASE)
        limit = min(3, len(tbl.rows))  # scan up to first 3 rows
        for r in range(limit):
            try:
                raw = (_cell_text(tbl.rows[r].cells[ci]) or "").strip()
            except Exception:
                raw = ""
            if raw and pat.search(raw):
                return _strip_colon_like(_norm_space(raw))
        return ""

    def _extract_all_labels_from_same_cell(cell: _Cell) -> list[str]:
        """
        Extract one or more label lines from a cell where label(s) and the
        grey box/underline live together. Skips inner 'Contact N' lines.
        """
        contact_pat = re.compile(r"^contact\s*\d+\b[:)?]?$", re.IGNORECASE)

        labels, seen = [], set()
        pars = cell.paragraphs
        for i, p in enumerate(pars):
            txt = (p.text or "").strip()
            if not txt:
                continue

            # ignore inner lines that are themselves 'Contact N'
            if contact_pat.search(txt):
                continue

            # next paragraph looks like a placeholder => current is a label
            nxt = (pars[i + 1].text or "").strip() if i + 1 < len(pars) else ""
            if nxt and not ALNUM_RE.search(nxt):
                lab = re.sub(r"[:：﹕꞉˸፡︓]\s*$", "", txt).strip()
                if lab and lab.lower() not in seen:
                    seen.add(lab.lower()); labels.append(lab)
                continue

            # inline placeholder split in same paragraph
            m = re.search(r"([ \t\u00A0\u2000-\u200B]{3,}|_{3,}|[\u2014\-]{3,})", txt)
            if m:
                left = txt[:m.start()].strip()
                if left and not contact_pat.search(left):
                    lab = re.sub(r"[:：﹕꞉˸፡︓]\s*$", "", left).strip()
                    if lab and lab.lower() not in seen:
                        seen.add(lab.lower()); labels.append(lab)
                continue

        # final fallback: first alnum line that's not 'Contact N'
        if not labels:
            for p in pars:
                txt = (p.text or "").strip()
                if txt and not contact_pat.search(txt) and ALNUM_RE.search(txt):
                    lab = re.sub(r"[:：﹕꞉˸፡︓]\s*$", "", txt).strip()
                    if lab and lab.lower() not in seen:
                        labels.append(lab)
                    break
        return labels



    UNDERLINE_RUN = re.compile(r"_+")
    MOSTLY_SPACEY = re.compile(r"^[\s\u2000-\u200B\u00A0.-]{3,}$")
    INLINE_UNDERLINE = re.compile(r"^(?P<label>.*?[:：﹕꞉˸])\s*[_\u2000-\u200B\u00A0\.\-–—]{3,}\s*$")
    CHECKBOX_LINE   = re.compile(r"^\s*(?:[□☐☑☒]|\[\s?[xX✓]?\s?\])\s*(?P<label>.+?)\s*$")
    INLINE_SPACEY_NO_COLON = re.compile(r"^(?P<label>[A-Z][^:]{1,80}?)\s+[_\u2000-\u200B\u00A0\.\-–—]{3,}\s*$")
    DUO_SPACEY_NO_COLON    = re.compile(
        r"^(?P<label1>[A-Z][^:]{1,80}?)\s+[_\u2000-\u200B\u00A0\.\-–—]{3,}\s+"
        r"(?P<label2>[A-Z][^:]{1,80}?)\s+[_\u2000-\u200B\u00A0\.\-–—]{3,}\s*$"
    )

    # NEW helpers for per-column header + same-cell placeholders
    ALNUM_RE = re.compile(r"[A-Za-z0-9]")

    def _blankish_cell(c: _Cell) -> bool:
        t2 = _cell_text(c).strip()
        return (not t2) or bool(UNDERLINE_RUN.search(t2) or MOSTLY_SPACEY.match(t2) or len(t2) <= 1)

    def _short_header(txt: str) -> bool:
        t = _strip_colon_like(_norm_space(txt))
        return bool(t) and len(t) <= 60 and len(t.split()) <= 6

    def _row0_has_contact_headers(tbl: Table) -> bool:
        try:
            hdrs = [(_cell_text(c) or "").strip() for c in tbl.rows[0].cells]
        except Exception:
            return False
        return any(re.match(r"(?i)^contact\s*\d+\b", h) for h in hdrs if h)

    def _cell_has_placeholder_like(cell: _Cell) -> bool:
        # Shading on cell or paragraph, or runs that are mostly non-alnum (underline/box)
        try:
            if cell._tc.xpath('.//w:tcPr/w:shd', namespaces=cell._tc.nsmap):
                return True
        except Exception:
            pass
        for p in cell.paragraphs:
            try:
                if p._element.xpath('.//w:pPr/w:shd'):
                    return True
            except Exception:
                pass
            vis = (p.text or "").strip()
            if vis and not ALNUM_RE.search(vis) and len(vis.replace(" ", "")) >= 3:
                return True
        return False

    def _extract_label_from_same_cell(cell: _Cell) -> str:
        """
        If cell has multiple paragraphs and the next one looks placeholder-like,
        use current paragraph as label. Otherwise try left-of-inline-blank;
        fallback to first alnum line.
        """
        pars = cell.paragraphs
        for i, p in enumerate(pars):
            txt = (p.text or "").strip()
            if not txt:
                continue
            if i + 1 < len(pars):
                nxt = (pars[i+1].text or "").strip()
                if nxt and not ALNUM_RE.search(nxt):
                    return _strip_colon_like(_norm_space(txt))
            m = re.search(r"([ \t\u00A0\u2000-\u200B]{3,}|_{3,}|[\u2014\-]{3,})", txt)
            if m:
                left = txt[:m.start()].strip()
                if left:
                    return _strip_colon_like(_norm_space(left))
        for p in pars:
            txt = (p.text or "").strip()
            if txt and ALNUM_RE.search(txt):
                return _strip_colon_like(_norm_space(txt))
        return ""



    fields = []
    current_section = ""
    blocks = list(_iter_blocks(doc))

    # NEW: per-(section,label) stable indexing
    field_counts = defaultdict(int)
    def _push(section: str, label: str, placement: str, ftype: str = "text", idx: int | None = None):
        lab_short = label
        key = (section.lower(), lab_short.lower())
        if idx is None:
            field_counts[key] += 1
            index_val = field_counts[key]
        else:
            # force the desired index (e.g., column number for Contact N)
            index_val = int(idx)
            field_counts[key] = max(field_counts[key], index_val)
        fields.append({
            "section": section, "label": label, "label_short": lab_short,
            "page": "", "index": index_val, "placement": placement, "type": ftype,
        })


    i = 0
    while i < len(blocks):
        advanced = False
        b = blocks[i]

        # ============== PARAGRAPHS ==============
        if isinstance(b, Paragraph):
            txt = (b.text or "").strip()
            if txt:
                m2 = DUO_SPACEY_NO_COLON.match(txt)
                if m2:
                    for lab in (_strip_colon_like(_norm_space(m2.group("label1"))),
                                _strip_colon_like(_norm_space(m2.group("label2")))):
                        if lab and len(lab) <= 80:
                            _push(current_section, lab, "docx_underline_inline_spacey", "text")
                    i += 1; advanced = True; continue

                m1 = INLINE_SPACEY_NO_COLON.match(txt)
                if m1:
                    lab = _strip_colon_like(_norm_space(m1.group("label")))
                    if lab and len(lab) <= 80:
                        _push(current_section, lab, "docx_underline_inline_spacey", "text")
                    i += 1; advanced = True; continue

                if _looks_like_section_title(txt):
                    current_section = _strip_colon_like(txt)
                    i += 1; advanced = True; continue

                m = INLINE_UNDERLINE.match(txt)
                if m:
                    lab = _strip_colon_like(_norm_space(m.group("label")))
                    if lab:
                        _push(current_section, lab, "docx_underline_inline", "text")
                    i += 1; advanced = True; continue

                m = CHECKBOX_LINE.match(txt)
                if m:
                    lab = _norm_space(m.group("label"))
                    if lab and 1 <= len(lab.split()) <= 10:
                        _push(current_section, lab, "docx_checkbox_line", "checkbox")
                    i += 1; advanced = True; continue

                if i + 1 < len(blocks) and isinstance(blocks[i + 1], Paragraph):
                    nxt_text = (blocks[i + 1].text or "").strip()
                    if UNDERLINE_RUN.search(nxt_text) or MOSTLY_SPACEY.match(nxt_text):
                        lab = _strip_colon_like(_norm_space(txt))
                        if lab and len(lab) <= 80:
                            _push(current_section, lab, "docx_underline_nextline", "text")
                        i += 2; advanced = True; continue

        # ============== TABLES ==============
        # ============== TABLES ==============
        if isinstance(b, Table):
            nrows = len(b.rows)
            ncols = len(b.rows[0].cells) if nrows else 0
            if ncols < 1:
                i += 1
                advanced = True
                continue

            def _row_texts(row):
                return [_strip_colon_like(_norm_space(_cell_text(c))) for c in row.cells]

            # --- NEW: row-agnostic "Contact N" column sections ---
            # As we scan rows, whenever a cell says "Contact N" we mark that column's
            # active section. Subsequent fillable cells in that column get emitted
            # under that Contact section.
            contact_pat = re.compile(r"^contact\s*(\d+)\b[:)?]?$", re.IGNORECASE)
            col_active_section: list[str | None] = [None] * ncols
            emitted_contact_any = False

            for ri in range(nrows):
                row = b.rows[ri]
                # 1) Arm columns with "Contact N" when we see the header anywhere
                for ci in range(ncols):
                    raw = (_cell_text(row.cells[ci]) or "").strip()
                    if not raw:
                        continue
                    m = contact_pat.search(raw)
                    if m:
                        # normalize header text (e.g., "Contact 3")
                        hdr = f"Contact {m.group(1)}"
                        col_active_section[ci] = hdr
                        continue

                # 2) For columns that have an active Contact section, try to extract labels
                for ci in range(ncols):
                    sec_hdr = col_active_section[ci]
                    if not sec_hdr:
                        continue  # this column hasn't been armed by a "Contact N" cell yet

                    cell = row.cells[ci]

                    # Only consider cells that look like they expect a value OR
                    # actually contain visible placeholder-ish content
                    looks_fillable = _cell_has_placeholder_like(cell)
                    if not looks_fillable:
                        vis = (_cell_text(cell) or "").strip()
                        # if totally blank or too long narrative, skip
                        if not vis or len(vis) > 180:
                            continue

                    # Extract a label from the same cell (skip inner "Contact N")
                    label = _extract_label_from_same_cell(cell)
                    if not label:
                        continue

                    # Derive a stable per-contact index from the header if present; else fallback to column index
                    msec = contact_pat.search(sec_hdr)
                    idx_for_contact = int(msec.group(1)) if msec else (ci + 1)

                    _push(sec_hdr, label, "docx_table_samecell", "text", idx=idx_for_contact)
                    emitted_contact_any = True

            if emitted_contact_any:
                # We already emitted Contact-based fields; skip other detectors for this table
                i += 1
                advanced = True
                continue

            # ---------- Sliding matrix detector ----------
            # ---------- Row-aware matrix detector ----------
            def _cell_is_boxlike(c: _Cell) -> bool:
                # box/underline/shaded/checkbox-ish cells count as a field “box”
                t = _cell_text(c).strip()
                if not t:
                    return True
                if re.search(r"_{3,}", t) or re.match(r"^[\s\u2000-\u200B\u00A0.\-–—]{3,}$", t):
                    return True
                try:
                    if c._tc.xpath('.//w:tcPr/w:shd', namespaces=c._tc.nsmap):
                        return True
                except Exception:
                    pass
                return False

            def _clean(s: str) -> str:
                return re.sub(r"[:：﹕꞉˸]\s*$", "", (s or "").strip())

            emitted_matrix = False

            if nrows >= 2 and ncols >= 3:
                # Read header cells from row 0, excluding the first column (left label column)
                header_row = b.rows[0]
                headers = [_clean(_cell_text(header_row.cells[c])) for c in range(1, ncols)]
                headers = [h for h in headers if h]  # nonempty
                # Require at least 2 short headers to consider matrix behavior
                if len([h for h in headers if _short_header(h)]) >= 2:
                    # Walk each data row independently and decide whether it's a matrix row or a single-field row
                    for ri in range(1, nrows):
                        row = b.rows[ri]
                        left_label = _clean(_cell_text(row.cells[0]))
                        # Skip rows without a clear left label
                        if not left_label or len(left_label) > 140:
                            continue

                        # Count how many right-side cells look like independent boxes
                        box_cols = []
                        for ci in range(1, ncols):
                            if _cell_is_boxlike(row.cells[ci]):
                                box_cols.append(ci)

                        if len(box_cols) >= 2:
                            # Treat as a matrix row: emit one field per header for THIS row/section
                            for ci in box_cols:
                                htxt = _clean(_cell_text(header_row.cells[ci])) if ci < len(header_row.cells) else ""
                                if not htxt or not _short_header(htxt):
                                    continue
                                _push(left_label, htxt, "docx_table_grid_row", "text")
                            emitted_matrix = True
                        else:
                            # Not a matrix row (e.g., a single long merged box) => emit only the row label as a single field
                            _push(left_label, left_label, "docx_table_pair_row", "text")

            if emitted_matrix:
                i += 1
                advanced = True
                continue


            if emitted_matrix:
                i += 1
                advanced = True
                # IMPORTANT: since we handled this table as a grid, do not run the
                #             simple pair-table fallback below for this table.
                continue


            # ---------- Fallback: simple pair tables ----------
            if ncols >= 2:
                for r in b.rows:
                    left = _strip_colon_like(_norm_space(_cell_text(r.cells[0])))
                    if not left or len(left) > 100:
                        continue
                    _push(current_section, left, "docx_table_pair", "text")

            i += 1
            advanced = True
            continue


        # default advance
        if not advanced:
            i += 1

    tpl = {"fields": fields}
    if out_json:
        with open(out_json, "w", encoding="utf-8") as f:
            json.dump(tpl, f, ensure_ascii=False, indent=2)
    return tpl





def _pick_builder_by_input(input_path: str, verbose: bool = False):
    """
    Decide which builder to call based on the input file extension.
      - *.docx -> prefer docx_prefill.build_docx_template if available
      - otherwise -> pdf_prefill.build_pdf_template
    Also allows a local `build_pdf_template` or `build_docx_template` to override via globals().
    """
    local_build_pdf = globals().get("build_pdf_template")
    local_build_docx = globals().get("build_docx_template")

    ext = (os.path.splitext(input_path)[1] or "").lower()

    if ext == ".docx":
        if callable(local_build_docx):
            if verbose:
                print(f"🔧 Using local DOCX builder: {_describe_builder(local_build_docx)}")
            return local_build_docx, "docx"
        if callable(_docx_builder):
            if verbose:
                print(f"🔧 Using docx builder: {_describe_builder(_docx_builder)}")
            return _docx_builder, "docx"
        if callable(local_build_pdf):
            if verbose:
                print(f"🔧 Using local PDF builder (fallback for DOCX): {_describe_builder(local_build_pdf)}")
            return local_build_pdf, "pdf"
        if callable(_build_pdf_template_external):
            if verbose:
                print(f"🔧 Using imported PDF builder (fallback for DOCX): {_describe_builder(_build_pdf_template_external)}")
            return _build_pdf_template_external, "pdf"

    if callable(local_build_pdf):
        if verbose:
            print(f"🔧 Using local PDF builder: {_describe_builder(local_build_pdf)}")
        return local_build_pdf, "pdf"
    if callable(_build_pdf_template_external):
        if verbose:
            print(f"🔧 Using imported PDF builder: {_describe_builder(_build_pdf_template_external)}")
        return _build_pdf_template_external, "pdf"

    detail = ""
    if _build_pdf_template_import_error is not None:
        sys.stderr.write("----- pdf_prefill import traceback -----\n")
        if _build_pdf_template_import_tb:
            sys.stderr.write(_build_pdf_template_import_tb + "\n")
        else:
            sys.stderr.write(repr(_build_pdf_template_import_error) + "\n")
        sys.stderr.write("----- end traceback -----\n")
        detail = (
            f"\n(Import error: {type(_build_pdf_template_import_error).__name__}: "
            f"{_build_pdf_template_import_error})"
        )
    raise RuntimeError(
        "No builder found. Define build_docx_template/build_pdf_template in this file "
        "OR ensure docx_prefill/pdf_prefill are importable." + detail
    )


def _ensure_parent_dir(path: str):
    d = os.path.dirname(os.path.abspath(path))
    if d and not os.path.exists(d):
        os.makedirs(d, exist_ok=True)


def _field_title_from_fdef(fdef: dict) -> str:
    """Prefer a concise key when available, otherwise fallback to label."""
    title = (fdef.get("label_short") or fdef.get("label") or "").strip()
    if ":" in title:
        left = title.split(":", 1)[0].strip()
        if len(left) >= 6:
            title = left
    return title or (fdef.get("label_full") or "").strip()

# --- NEW: helpers to detect inline checkbox groups (no hard-coding of names)
_INLINE_SPLIT_RE = re.compile(r"(?:\t+|\s{2,}|(?:\s+OR\s+))", re.IGNORECASE)
_BOX_PREFIX_RE = re.compile(r"^\s*(?:[□☐☒]\s*|\[\s?[xX✓]?\s?\]\s*)")


def _split_inline_checkbox_options(display_text: str):
    """
    If display_text looks like a single line containing multiple short options
    separated by tabs / 2+ spaces / ' OR ', split into options.
    Returns list[str] or None if not applicable.
    """
    if not display_text or not display_text.strip():
        return None
    t = _BOX_PREFIX_RE.sub("", display_text).strip()
    parts = [p.strip() for p in _INLINE_SPLIT_RE.split(t) if p and p.strip()]
    parts = [p for p in parts if _looks_short_option(p, max_words=4)]
    if 2 <= len(parts) <= 6 and len(set(p.lower() for p in parts)) >= 2:
        return parts
    return None

from collections import defaultdict
_RECIP_RE = re.compile(r"\brecipient\s*\d+\b", re.IGNORECASE)

# running item counter PER section (so each RECIPIENT N gets item 1..K)
_section_item_counter = defaultdict(int)


def export_lookup_template_from_json(
        template_json: str,
        out_path: str = "lookup_template.xlsx",
        docx_path: str = None
) -> str:
    """
    Produce an Excel (or CSV) with columns:
    Section | Page | Field | Index | Value | Choices

    If docx_path is a DOCX, we also scan for 'matrix' tables whose header row
    contains multiple field labels and there are multiple blank rows beneath.
    For each header field found in the template, we auto-append Index=2..N rows
    so you can enter values for each data row in the document.
    """
    if not os.path.exists(template_json):
        raise FileNotFoundError(f"Template JSON not found: {template_json}")

    import json
    import pandas as pd
    import re
    from collections import OrderedDict, defaultdict

    with open(template_json, "r", encoding="utf-8") as f:
        tpl = json.load(f)

    rows = []
    seq = 0  # used to make a stable unique fallback field name when needed

    # --- NEW (scoped to this function): helpers for recipients & inline options
    _RECIP_RE = re.compile(r"\brecipient\s*(\d+)\b", re.IGNORECASE)

    # rely on your global _INLINE_SPLIT_RE if present; else provide a conservative default
    _INLINE_SPLIT_DEFAULT = re.compile(r"(?:\t+|\s{2,}|\s+\bOR\b\s+|[;•/])", re.IGNORECASE)
    try:
        _INLINE_SPLIT_RE  # noqa: F401
    except NameError:    # fallback only if it's not defined elsewhere
        _INLINE_SPLIT_RE = _INLINE_SPLIT_DEFAULT  # type: ignore

    def _recipient_index_from(section: str, default_idx: int) -> int:
        m = _RECIP_RE.search(section or "")
        if not m:
            return default_idx
        try:
            return int(m.group(1))
        except Exception:
            return default_idx

    def _split_inline_options_text(s: str) -> list[str]:
        t = (s or "").strip()
        if not t:
            return []
        parts = [p.strip(" -:") for p in _INLINE_SPLIT_RE.split(t) if p.strip()]
        # keep only short-ish items (≤ 8 words; ≥ 3 token chars)
        def _tok_len(x: str) -> int:
            return len(re.findall(r"[A-Za-z0-9]", x))
        parts = [p for p in parts if 1 <= len(p.split()) <= 8 and _tok_len(p) >= 3]
        # require at least 2 distinct parts to treat it as a group
        seen, out = set(), []
        for p in parts:
            k = p.lower()
            if k not in seen:
                seen.add(k)
                out.append(p)
        if len(out) >= 2 and len(out) <= 10:
            return out
        return []

    for fdef in (tpl.get("fields") or []):
        seq += 1

        raw_label = (fdef.get("label") or "").strip()
        section = (fdef.get("section") or "").strip()

        page_val = fdef.get("page", None)
        try:
            page = int(page_val) if page_val not in (None, "") else ""
            if isinstance(page, int) and page <= 0:
                page = ""
        except Exception:
            page = ""

        try:
            index = int(fdef.get("index", 1) or 1)
        except Exception:
            index = 1

        title_from_def = _field_title_from_fdef(fdef).strip()
        looks_unknown = (raw_label.lower().startswith("unknown_") if raw_label else True)
        if not title_from_def or title_from_def.lower().startswith("unknown_") or looks_unknown:
            sec_display = section if section else "No Section"
            title_from_def = f"Field #{seq} (Sec: {sec_display}, Idx: {index})"

        # Existing behavior: only pre-fill the "Choices" cell for AcroForm choice fields.
        choices_cell = ""
        if (fdef.get("placement") or "").lower() == "acro_choice":
            ch = fdef.get("choices") or []
            if isinstance(ch, list) and ch:
                choices_cell = " | ".join(str(x) for x in ch)

        # ---- NEW: Expand checkbox-like choice groups into separate rows for recipients & general checkboxes
        # 1) If this is a real choice list (AcroForm), expand to one row per option.
        if (fdef.get("placement") or "").lower() == "acro_choice":
            choice_list = fdef.get("choices") or []
            if isinstance(choice_list, list) and choice_list:
                if _RECIP_RE.search(section or ""):
                    # RECIPIENT N → use actual option labels and Index=N
                    recip_idx = _recipient_index_from(section, index)
                    for opt in choice_list:
                        opt_label = str(opt).strip()
                        if not opt_label:
                            continue
                        rows.append({
                            "Section": section,
                            "Page": page,
                            "Field": opt_label,   # ← actual label, not "item k"
                            "Index": recip_idx,   # ← recipient number
                            "Value": "",
                            "Choices": "checkbox",
                        })
                else:
                    # Non-recipient: keep each option as its own checkbox row
                    for opt in choice_list:
                        opt_label = str(opt).strip()
                        if not opt_label:
                            continue
                        rows.append({
                            "Section": section,
                            "Page": page,
                            "Field": opt_label,
                            "Index": 1,
                            "Value": "",
                            "Choices": "checkbox",
                        })
                continue  # skip the default single-row append

        # 2) If no explicit choices, try to split inline compact options in the label text.
        #    When STRICT is on, do NOT invent checkboxes from inline prose.
        if not choices_cell and not _strict():
            split_source = (fdef.get("label_full") or fdef.get("label") or title_from_def or "").strip()
            if not split_source.endswith(":"):
                opts = _split_inline_options_text(split_source)
                if opts:
                    if _RECIP_RE.search(section or ""):
                        # RECIPIENT N → use actual option labels and Index=N
                        recip_idx = _recipient_index_from(section, index)
                        for opt in opts:
                            opt_label = str(opt).strip()
                            if not opt_label:
                                continue
                            rows.append({
                                "Section": section,
                                "Page": page,
                                "Field": opt_label,   # ← actual label, not "item k"
                                "Index": recip_idx,   # ← recipient number
                                "Value": "",
                                "Choices": "checkbox",
                            })
                    else:
                        for opt in opts:
                            opt_label = str(opt).strip()
                            if not opt_label:
                                continue
                            rows.append({
                                "Section": section,
                                "Page": page,
                                "Field": opt_label,
                                "Index": 1,
                                "Value": "",
                                "Choices": "checkbox",
                            })
                    continue  # expanded; skip default

        # Default: original single-row emission (text/underline/table/etc.)
        rows.append({
            "Section": section,
            "Page": page,
            "Field": title_from_def,
            "Index": index,
            "Value": "",
            "Choices": choices_cell,
        })

    # Build initial DF
    cols = ["Section", "Page", "Field", "Index", "Value", "Choices"]
    df = pd.DataFrame(rows, columns=cols).fillna("")

    # --- Cleanup: drop obvious non-fillable narrative lines on PDF pages ---
    def _too_narrative(row) -> bool:
        if str(row.get("Choices","")).strip().lower() == "checkbox":
            return False  # keep checkboxes
        field = str(row.get("Field","")).strip()
        if not field:
            return True
        if len(field) > 120 or len(field.split()) > 20:
            return True
        bad_snippets = (
            "subscription agreement", "page ", "the investor may also contact",
            "copy documents for any particular recipient", "attach additional pages if needed"
        )
        f_low = field.lower()
        if any(s in f_low for s in bad_snippets):
            return True
        return False

    df = df[~df.apply(_too_narrative, axis=1)].reset_index(drop=True)

    # --- NEW: Recipient safety net with fused-option splitter (adds the 4th item) ---
    _FUSED_RE = re.compile(
        r"^(?P<a>Account statements)\s+(?P<b>Annual\s*\(audited\)\s*/\s*quarterly\s*\(unaudited\)\s*financial\s*statements)$",
        re.IGNORECASE
    )
    def _split_fused_option(label: str) -> list[str]:
        t = (label or "").strip()
        m = _FUSED_RE.match(t)
        if m:
            return [m.group("a").strip(), m.group("b").strip()]
        # generic fallback: split before a Capitalized word followed by '('
        m2 = re.search(r"\s+(?=[A-Z][a-z]+\s*\()", t)
        if m2:
            left, right = t[:m2.start()].strip(), t[m2.start():].strip()
            if len(left.split()) >= 2 and len(right.split()) >= 2:
                return [left, right]
        return [t]

    def _is_short_option(label: str) -> bool:
        t = (label or "").strip()
        if not t or t.endswith(":"):
            return False
        # explicitly keep the long “Annual (audited)/quarterly (unaudited) financial statements”
        if re.search(r"\bannual\s*\(audited\)\s*/\s*quarterly\s*\(unaudited\)\s*financial\s*statements\b",
                     t, flags=re.IGNORECASE):
            return True
        if len(t) > 140 or len(t.split()) > 16:
            return False
        lower = t.lower()
        if lower in {"name", "e-mail", "email", "telephone", "facsimile", "fax", "address"}:
            return False
        return True

    new_rows, drop_idx = [], []
    for (sec, pg), grp in df.groupby(["Section", "Page"], dropna=False):
        if not _RECIP_RE.search(sec or ""):
            continue
        recip_idx = _recipient_index_from(sec or "", 1)

        # collect & expand candidates in original order
        expanded = []
        for i, r in grp.iterrows():
            fld = str(r["Field"]).strip()
            if _is_short_option(fld):
                for part in _split_fused_option(fld):
                    expanded.append((i, part))

        if len(expanded) >= 2:
            # drop originals and emit standardized numbered items (keeps order)
            for k, (irow, _label) in enumerate(expanded, start=1):
                drop_idx.append(irow)
                new_rows.append({
                    "Section": sec,
                    "Page": pg,
                    "Field": f"{sec} – item {k}",
                    "Index": recip_idx,
                    "Value": "",
                    "Choices": "checkbox",
                })

    if new_rows:
        df = df.drop(index=list(set(drop_idx))).reset_index(drop=True)
        df = pd.concat([df, pd.DataFrame(new_rows, columns=df.columns)], ignore_index=True)

    # ---------- OPTIONAL: expand matrix tables by scanning the DOCX ----------
    def _clean_label(s: str) -> str:
        s = (s or "").strip()
        return re.sub(r"[:：﹕꞉˸]\s*$", "", s).strip()

    def _first_cell_text(cell) -> str:
        return " ".join(p.text for p in cell.paragraphs).strip()

    def _is_boxlike(cell) -> bool:
        txt = "".join(p.text or "" for p in cell.paragraphs).strip()
        if not txt:
            return True
        vis = re.sub(r"[\s_\-\u2014\.\u2002\u2003\u2007\u2009\u00A0]+", "", txt)
        return not re.search(r"[A-Za-z0-9]", vis)

    def _expand_matrix_rows_with_docx(df_in: pd.DataFrame, docx_file: str) -> pd.DataFrame:
        try:
            from docx import Document
        except Exception:
            return df_in

        if not (docx_file and os.path.exists(docx_file) and docx_file.lower().endswith(".docx")):
            return df_in

        doc = Document(docx_file)
        additions = []

        for tbl in doc.tables:
            if not tbl.rows:
                continue
            header = tbl.rows[0]
            ncols = len(header.cells)
            if ncols < 2 or len(tbl.rows) < 2:
                continue

            headers = [_clean_label(_first_cell_text(c)) for c in header.cells]
            nonempty = [h for h in headers if h]
            if len(nonempty) < 2:
                continue

            row1 = tbl.rows[1]
            def _cell_blankish(c):
                t = _first_cell_text(c)
                if not t:
                    return True
                t2 = t.strip()
                return bool(re.search(r"_+", t2) or re.match(r"^[\s\u2000-\u200B\u00A0.-]{3,}$", t2)) or (len(t2) <= 2)
            box_cnt = sum(1 for c in row1.cells if _cell_blankish(c))
            if box_cnt < max(2, int(0.6 * ncols)):
                continue  # not a matrix grid

            total_rows = max(1, len(tbl.rows) - 1)  # data rows
            if total_rows <= 1:
                continue

            for h in range(ncols):
                field_label = headers[h]
                if not field_label:
                    continue

                base_mask = (
                        (df_in["Field"].str.strip().str.lower() == field_label.strip().lower())
                        & (df_in["Index"].astype(int) == 1)
                )
                if not base_mask.any():
                    continue

                base_row = df_in[base_mask].iloc[0]
                for idx in range(2, total_rows + 1):
                    exists = (
                            (df_in["Field"].str.strip().str.lower() == field_label.strip().lower())
                            & (df_in["Index"].astype(int) == idx)
                    ).any()
                    if exists:
                        continue
                    additions.append({
                        "Section": base_row["Section"],
                        "Page": base_row["Page"],
                        "Field": field_label,
                        "Index": idx,
                        "Value": "",
                        "Choices": base_row.get("Choices", ""),
                    })

        if additions:
            df_out = pd.concat([df_in, pd.DataFrame(additions, columns=df_in.columns)], ignore_index=True)
            return df_out
        return df_in

    # ---- STRICT-aware checkbox handling helpers (only mark when glyphs exist) ----
    def _norm_text_for_match(s: str) -> str:
        t = re.sub(r"\s+", " ", str(s or "")).strip()
        t = t.strip(":-—–·•").lower()
        t = re.sub(r"[^\w\s]", "", t)
        t = re.sub(r"\s+", " ", t).strip()
        return t

    def _collect_real_checkbox_labels(docx_file: str) -> set[str]:
        """
        Return a set of normalized labels that truly have a checkbox glyph next to them
        in the DOCX (strict detection: no inference, no creation).
        """
        try:
            from docx import Document
        except Exception:
            return set()
        if not (docx_file and os.path.exists(docx_file) and docx_file.lower().endswith(".docx")):
            return set()

        BOX_RE = re.compile(r"(?:\[\s*[xX✓]?\s*\]|[□☐☑☒])")
        SEP_RE = re.compile(r"\s{2,}|(?:\s+OR\s+)|\||/|;|,|$")

        def _extract_from_line(text: str) -> list[str]:
            labels = []
            if not text:
                return labels

            # glyph -> label
            for m in BOX_RE.finditer(text):
                tail = re.sub(r"^\s+", "", text[m.end():])
                sep = SEP_RE.search(tail)
                cand = tail[:sep.start()] if sep else tail
                cand = re.sub(r"\s+", " ", cand).strip(" :-—–·•")
                if cand and len(cand) <= 140 and re.search(r"[A-Za-z0-9]", cand):
                    labels.append(cand)

            # label -> glyph
            for m in BOX_RE.finditer(text):
                head = re.sub(r"\s+$", "", text[:m.start()])
                parts = SEP_RE.split(head)
                if parts:
                    cand = re.sub(r"\s+", " ", parts[-1]).strip(" :-—–·•")
                    if cand and len(cand) <= 140 and re.search(r"[A-Za-z0-9]", cand):
                        labels.append(cand)

            # de-dup
            out, seen = [], set()
            for lb in labels:
                k = _norm_text_for_match(lb)
                if k and k not in seen:
                    seen.add(k); out.append(lb)
            return out

        def _iter_lines(doc):
            for p in doc.paragraphs:
                yield p.text or ""
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p.text or ""

        doc = Document(docx_file)
        found = []
        for line in _iter_lines(doc):
            if "[" not in line and not re.search(r"[□☐☑☒]", line or ""):
                continue
            found.extend(_extract_from_line(line or ""))

        return {_norm_text_for_match(x) for x in found if x}

    def _mark_checkboxes_from_docx(df_in: pd.DataFrame, docx_file: str) -> pd.DataFrame:
        """
        Only mark Choices='checkbox' for rows whose Field matches a real checkbox label in the DOCX.
        Does NOT add any rows.
        """
        if not docx_file or not os.path.exists(docx_file):
            return df_in

        # If parsing fails, leave DF untouched
        try:
            real_keys = _collect_real_checkbox_labels(docx_file)
        except Exception:
            return df_in

        df = df_in.copy()
        if "Field" not in df.columns:
            return df

        BOX_RE = re.compile(r"(?:\[\s*[xX✓]?\s*\]|[□☐☑☒])")

        def _is_real_checkbox(field_val: str) -> bool:
            f = str(field_val or "")
            # If the field itself already contains a glyph, that's certainly a checkbox
            if BOX_RE.search(f):
                return True
            # Otherwise match against extracted labels
            return _norm_text_for_match(f) in real_keys

        mask = df["Field"].apply(_is_real_checkbox)
        if "Choices" not in df.columns:
            df["Choices"] = ""
        df.loc[mask, "Choices"] = "checkbox"
        return df

    def _split_inline_checkboxes_with_docx(df_in: pd.DataFrame, docx_file: str) -> pd.DataFrame:
        """
        Split a single lookup row into multiple checkbox rows ONLY if we see
        actual checkbox glyph tokens in the DOCX lines (□, ☐, ☑, ☒, [ ], [x], [X]).
        If STRICT is on, we do not use context heuristics and never create rows
        that don't match an existing field.
        """
        try:
            from docx import Document
        except Exception:
            return df_in

        if not (docx_file and os.path.exists(docx_file) and docx_file.lower().endswith(".docx")):
            return df_in

        import re as _re
        import unicodedata as _ud
        import string as _st

        doc = Document(docx_file)

        def _norm(s: str) -> str:
            s = _ud.normalize("NFKC", str(s or "")).strip()
            s = _re.sub(r"\s+", " ", s)
            return s

        def _norm_key(s: str) -> str:
            t = _norm(s).lower()
            return t.translate(str.maketrans("", "", _st.punctuation))

        def _iter_all_paragraphs(d):
            for p in d.paragraphs:
                yield p
            for t in d.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p

        def _looks_short_label(s: str, max_words: int = 12, max_len: int = 120) -> bool:
            s = (s or "").strip(" -:")
            if not s or len(s) > max_len:
                return False
            words = s.split()
            return 1 <= len(words) <= max_words

        # Glyph-based group splitter
        BOX_TOKEN = r"(?:\[\s*[xX✓]?\s*\]|[□☐☑☒])"
        LABEL = r"([A-Z][\w\.\-/&()' ]{0,40}?)(?<!:)"
        LABEL_BOX_RE = _re.compile(LABEL + r"\s*" + BOX_TOKEN)

        glyph_groups = []
        paras = [p for p in _iter_all_paragraphs(doc)]
        for p in paras:
            txt = _norm(p.text)
            if not txt:
                continue
            labels_here = [m.group(1).strip(" -:") for m in LABEL_BOX_RE.finditer(txt)]
            labels_here = [lb for lb in labels_here if _looks_short_label(lb, max_words=8, max_len=60)]
            if len(labels_here) >= 2:
                glyph_groups.append(labels_here)

        df = df_in.copy()

        def _apply_split_by_labels(df0: pd.DataFrame, labels: list[str]) -> tuple[pd.DataFrame, bool]:
            combined = _norm_key(" ".join(labels))
            hit_idx = []
            for idx, r in df0.iterrows():
                f_raw = _norm(str(r["Field"]))
                f_key = _norm_key(f_raw)
                if f_key == combined:
                    hit_idx.append(idx)
                    continue
                # in-order containment fallback
                pos, ok = 0, True
                for lb in labels:
                    lk = _norm_key(lb)
                    j = f_key.find(lk, pos)
                    if j < 0:
                        ok = False
                        break
                    pos = j + len(lk)
                if ok:
                    hit_idx.append(idx)

            # If no matches, no split
            if not hit_idx:
                return df0, False

            new_rows = []
            for idx in hit_idx:
                base = df0.loc[idx]
                for lb in labels:
                    new_rows.append({
                        "Section": base["Section"],
                        "Page": base["Page"],
                        "Field": lb,
                        "Index": 1,
                        "Value": "",
                        "Choices": "checkbox",
                    })
            out = df0.drop(index=hit_idx).reset_index(drop=True)
            out = pd.concat([out, pd.DataFrame(new_rows, columns=df0.columns)], ignore_index=True)
            return out, True

        did_any_split = False
        for g in glyph_groups:
            df, did = _apply_split_by_labels(df, g)
            did_any_split = did_any_split or did

        # STRICT: stop after glyph-based groups (no context heuristics)
        if _strict():
            return df

        # ---- Non-strict legacy context heuristics (optional) ----
        TICK_TRIGGER_RE = _re.compile(r"\bplease\s+tick\s+one\b", _re.IGNORECASE)
        ENUM_BUL_RE     = _re.compile(r"^\s*(?:\((?:[a-eA-E]|i{1,3}|iv|v)\)|[-–•])\s+(.*)$")
        OR_SPLIT_RE     = _re.compile(r"\s+\bOR\b\s+", _re.IGNORECASE)

        ctx_groups = []
        for i, p in enumerate(paras):
            line = _norm(p.text)
            if not line:
                continue
            if TICK_TRIGGER_RE.search(line):
                opts: list[str] = []
                j = i + 1
                while j < len(paras) and len(opts) < 6:
                    t = _norm(paras[j].text)
                    if not t:
                        j += 1
                        continue
                    m = ENUM_BUL_RE.match(t)
                    if m:
                        label = m.group(1).strip()
                        label = _re.sub(r"\s*\(tick[^)]*\)\s*$", "", label, flags=_re.IGNORECASE).strip()
                        parts = [s.strip() for s in OR_SPLIT_RE.split(label) if s.strip()]
                        for part in parts:
                            if _looks_short_label(part, max_words=24, max_len=140):
                                opts.append(part)
                    else:
                        parts = [s.strip() for s in OR_SPLIT_RE.split(t) if s.strip()]
                        for part in parts:
                            if _looks_short_label(part, max_words=24, max_len=140):
                                opts.append(part)
                    if len(t) > 200 and opts:
                        break
                    j += 1

                uniq, seen = [], set()
                for o in opts:
                    k = _norm_key(o)
                    if k and k not in seen:
                        seen.add(k)
                        uniq.append(o)
                if len(uniq) >= 2:
                    ctx_groups.append(uniq[:6])

        for g in ctx_groups:
            df, did = _apply_split_by_labels(df, g)
            did_any_split = did_any_split or did

            # In non-strict mode only: if no hit, we may append pending rows (legacy behavior)
            if not did:
                pending = []
                for lb in g:
                    exists = (df["Field"].str.strip().str.lower() == lb.strip().lower()).any()
                    if not exists:
                        pending.append({
                            "Section": "",
                            "Page": "",
                            "Field": lb,
                            "Index": 1,
                            "Value": "",
                            "Choices": "checkbox",
                        })
                if pending:
                    df = pd.concat([df, pd.DataFrame(pending, columns=df.columns)], ignore_index=True)

        return df

    def _mark_checkbox_rows(df_in: pd.DataFrame) -> pd.DataFrame:
        # In strict mode, do nothing here.
        if _strict():
            return df_in

        df = df_in.copy()
        def _is_checkboxish(s: str) -> bool:
            t = (str(s or "").strip())
            if not t or ":" in t or len(t) > 60:
                return False
            words = t.split()
            if len(words) > 6:
                return False
            caps = sum(1 for w in words if w[:1].isalpha() and w[:1].upper() == w[:1])
            return caps >= max(1, int(0.6 * len(words)))
        mask = df["Field"].apply(_is_checkboxish)
        df.loc[mask, "Choices"] = df.loc[mask, "Choices"].replace({"": "checkbox"}, regex=False)
        return df

    # Apply DOCX-aware expansions/splits (if a DOCX path was provided)
    # Apply DOCX-aware expansions/splits (if a DOCX path was provided)
    if docx_path:
        # If the builder already emitted multiple Index values per (Section, Field),
        # we assume it's a row-label-aware grid and *skip* synthetic expansion.
        has_multi_index = (
                df.groupby(["Section", "Field"])["Index"]
                .nunique(dropna=False)
                .reset_index(name="n")
                .query("n > 1")
                .shape[0] > 0
        )

        if not has_multi_index:
            # Safe: grid/matrix index expansion (only when needed)
            df = _expand_matrix_rows_with_docx(df, docx_path)

        # Splitting only when allowed (strict mode disables inference)
        df = _split_inline_checkboxes_with_docx(df, docx_path)
        if not _strict():
            df = _final_split_compound_checkbox_fields(df)
            df = _mark_checkbox_rows(df)

        # Always run strict pass last: mark only rows that truly have glyphs in the DOCX
        df = _mark_checkboxes_from_docx(df, docx_path)


    # ---------- sort & write ----------
    def _sort_key_row(r):
        page_sort = (999999 if r["Page"] == "" else int(r["Page"]))
        return page_sort, str(r["Section"]).lower(), str(r["Field"]).lower(), int(r["Index"])

    df = df.sort_values(
        by=["Section", "Field", "Index", "Page"],
        key=lambda s: s if s.name != "Page" else s.apply(lambda x: 999999 if x == "" else int(x)),
    ).reset_index(drop=True)


    _ensure_parent_dir(out_path)
    if out_path.lower().endswith(".csv"):
        df.to_csv(out_path, index=False, encoding="utf-8-sig")
    else:
        try:
            df.to_excel(out_path, index=False)
        except Exception as e:
            fallback = os.path.splitext(out_path)[0] + ".csv"
            df.to_csv(fallback, index=False, encoding="utf-8-sig")
            print(f"⚠️  Could not write Excel ({e}). Wrote CSV instead: {fallback}")
            return fallback

    print(f"📤 Lookup template written to {out_path} with {len(df)} rows.")
    return out_path


def _iter_all_paragraph_sequences(doc: Document):
    """
    Yield sequences of paragraphs that belong together, so we can look-ahead.
    Yields lists:
      - the top-level doc paragraphs (single list)
      - paragraphs for each table cell (one list per cell)
    """
    # top-level
    yield doc.paragraphs
    # tables
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                yield cell.paragraphs


def fill_docx_from_lookup(input_path: str, lookup_path: str, out_docx: str):
    """
    Minimal, conservative filler:
      • For a row with Choices='checkbox': toggle only existing glyphs (☐/☒ or [ ]/[x]) on the *label line itself*.
      • For a text row: if the label paragraph is immediately followed by a placeholder-like paragraph,
        write the value into that next paragraph; otherwise inline after the label.
      • Uses Field + Index to pick the N-th match (per normalized Field).
    """
    import pandas as _pd
    from docx import Document as _Doc

    if not os.path.exists(input_path):
        raise FileNotFoundError(input_path)
    if not os.path.exists(lookup_path):
        raise FileNotFoundError(lookup_path)

    df = _pd.read_excel(lookup_path) if lookup_path.lower().endswith((".xlsx", ".xls")) else _pd.read_csv(lookup_path)
    # Normalize expected columns
    exp = ["Section","Page","Field","Index","Value","Choices"]
    for c in exp:
        if c not in df.columns:
            df[c] = ""
    df = df.fillna("")
    # Only keep rows that actually have a value (avoid blank overwrites)
    df = df[df["Value"].astype(str).str.strip() != ""].copy()

    # Build fill plan grouped by normalized field name
    plan = []
    for _, r in df.iterrows():
        field_key = _norm_key(str(r["Field"]))
        try:
            idx = int(r["Index"]) if str(r["Index"]).strip() != "" else 1
        except Exception:
            idx = 1
        plan.append({
            "key": field_key,
            "raw": str(r["Field"]),
            "index": max(1, idx),
            "val": str(r["Value"]),
            "is_checkbox": str(r.get("Choices","")).strip().lower() == "checkbox",
        })

    if not plan:
        # nothing to fill; just copy through
        doc = _Doc(input_path)
        doc.save(out_docx)
        print(f"⚠️  No non-empty rows in lookup. Wrote passthrough to {out_docx}.")
        return

    doc = _Doc(input_path)

    # per-key match counter so we can honor Index=N
    seen_counts: dict[str, int] = {}

    def _match_field(par_text: str, key: str) -> bool:
        """
        Consider a paragraph a match if it equals the label (normalized) OR starts with that label
        (e.g., 'Email:' or 'Email –').
        """
        t = _norm_key(par_text)
        if not t:
            return False
        if t == key:
            return True
        # allow label followed by punctuation/colon-ish
        return t.startswith(key + " ") or t.startswith(key + ":")

    # Build a quick index: for each key, the list of rows (could have multiple different indices)
    by_key = {}
    for row in plan:
        by_key.setdefault(row["key"], []).append(row)

    # Walk every paragraph sequence with look-ahead capability
    for pars in _iter_all_paragraph_sequences(doc):
        i = 0
        while i < len(pars):
            p = pars[i]
            raw = p.text or ""
            if not raw.strip():
                i += 1
                continue

            # Test each field key that appears in plan (cheap filter)
            for key, rows in by_key.items():
                if not _match_field(raw, key):
                    continue

                # bump counter & pick the row whose Index matches this occurrence
                seen_counts[key] = seen_counts.get(key, 0) + 1
                occ = seen_counts[key]

                target = next((r for r in rows if r["index"] == occ), None)
                if not target:
                    # No specific row for this occurrence; skip
                    continue

                val = target["val"]
                is_cb = target["is_checkbox"]

                if is_cb:
                    # Toggle glyphs on THIS line only (we don't invent checkboxes)
                    new_text = _toggle_checkbox_glyphs_in_text(raw, val)
                    if new_text != raw:
                        _set_paragraph_text(p, new_text)
                else:
                    # Try to write into the placeholder paragraph right after the label
                    wrote = False
                    if i + 1 < len(pars):
                        nxt = pars[i+1]
                        nxt_text = (nxt.text or "").strip()
                        if _is_placeholder_like(nxt_text):
                            # clear and write the value in the next paragraph
                            _set_paragraph_text(nxt, str(val))
                            wrote = True

                    if not wrote:
                        # inline after the label text
                        _write_value_into_paragraph_after_label(p, str(val))
                # done with this paragraph; move to next
                break

            i += 1

    # Save result
    _ensure_parent_dir(out_docx)
    doc.save(out_docx)
    print(f"✅ Filled DOCX written to {out_docx}")



def _call_builder_with_compat(builder, input_path: str, template_json: str):
    """
    Call a builder with broad signature compatibility:
    - build_XXX_template(path, template_json, lookup_rows=None, dry_run=False)
    - build_XXX_template(path, template_json=..., lookup_rows=None, dry_run=False)
    - build_XXX_template(doc_or_path=path, template_json=..., ...)
    """
    _ensure_parent_dir(template_json)
    try:
        return builder(input_path, template_json, lookup_rows=None, dry_run=False)
    except TypeError:
        pass
    try:
        return builder(input_path, template_json=template_json, lookup_rows=None, dry_run=False)
    except TypeError:
        pass
    try:
        return builder(doc_or_path=input_path, template_json=template_json, lookup_rows=None, dry_run=False)
    except TypeError:
        pass
    return builder(input_path, template_json)


def _load_template_field_count(template_json: str) -> int:
    try:
        with open(template_json, "r", encoding="utf-8") as f:
            tpl = json.load(f)
        return len(tpl.get("fields", []) or [])
    except Exception:
        return -1


def export_lookup_template(
        input_path: str,
        template_json: str = "template_fields.json",
        out_path: str = "lookup_template.xlsx",
        rebuild_template: bool = False,
        debug_import: bool = False
) -> str:
    """
    Ensures a template JSON exists (building it if needed), then exports the lookup sheet.
    Adds diagnostics so you can see which builder ran and how many fields were detected.
    """
    builder, kind = _pick_builder_by_input(input_path, verbose=True or debug_import)

    must_build = rebuild_template or not os.path.exists(template_json)
    if must_build:
        print(f"🧩 Building template → {template_json}")
        _call_builder_with_compat(builder, input_path, template_json)
        cnt = _load_template_field_count(template_json)
        print(f"🧩 Template saved to {template_json} with {cnt} fields.")
        if cnt == 0:
            print("⚠️  No fields were detected.")
            if kind == "pdf":
                print("   • If using PDF, confirm you imported the correct pdf_prefill.py.")
                print("     → Run:  python -c \"import pdf_prefill,inspect; print(pdf_prefill.__file__)\"")
                print("   • If it’s a dynamic/XFA PDF or lacks detectable lines/widgets, try flattening (Print to PDF).")
                print("   • Or run the builder in DRY mode to see detection logs.")
            else:
                print("   • For DOCX, verify the document has tables/boxes/underlines the builder can detect.")
    else:
        print(f"📄 Using existing template: {template_json} ({_load_template_field_count(template_json)} fields)")

    return export_lookup_template_from_json(
        template_json,
        out_path,
        docx_path=input_path if input_path.lower().endswith(".docx") else None
    )


# ---------------------------
# CLI
# ---------------------------
def main():
    ap = argparse.ArgumentParser(
        description="Export lookup sheet and/or fill a DOCX from a lookup"
    )
    ap.add_argument("--input", required=True, help="Input file (PDF or DOCX)")
    ap.add_argument("--template", default="template_fields.json",
                    help="Template JSON (rebuilt if missing or --rebuild-template)")
    ap.add_argument("--make-lookup", metavar="OUT.xlsx",
                    help="Write blank lookup sheet (xlsx or csv)")
    ap.add_argument("--rebuild-template", action="store_true",
                    help="Force rebuild of the template JSON even if it already exists")
    ap.add_argument("--debug-import", action="store_true",
                    help="Print extra import info for the builder used.")

    # NEW: filling options (backward-compatible)
    ap.add_argument("--lookup", metavar="LOOKUP.xlsx",
                    help="Existing lookup sheet to use for filling (if omitted, uses --make-lookup output in this run)")
    ap.add_argument("--fill-docx", "--output", dest="out_docx", metavar="OUT.docx",
                    help="Fill the input DOCX and write OUT.docx")

    args = ap.parse_args()

    did_anything = False

    # 1) Build lookup if requested
    if args.make_lookup:
        export_lookup_template(
            input_path=args.input,
            template_json=args.template,
            out_path=args.make_lookup,
            rebuild_template=args.rebuild_template,
            debug_import=args.debug_import,
        )
        did_anything = True

    # 2) Fill DOCX if requested
    if args.out_docx:
        # prefer explicit --lookup; else if we just built one this run, reuse that path
        lookup_path = args.lookup or args.make_lookup
        if not lookup_path:
            raise SystemExit("When using --fill-docx, provide --lookup LOOKUP.xlsx (or also pass --make-lookup to build one).")
        if not os.path.exists(lookup_path):
            raise FileNotFoundError(f"Lookup file not found: {lookup_path}")
        if not args.input.lower().endswith(".docx"):
            raise SystemExit("Filling currently supports DOCX input only.")
        fill_docx_from_lookup(args.input, lookup_path, args.out_docx)
        did_anything = True

    if not did_anything:
        print("Nothing to do.\n"
              "• To export a lookup: --make-lookup lookup_template.xlsx\n"
              "• To fill a DOCX: --lookup lookup_template.xlsx --fill-docx out.docx")



if __name__ == "__main__":
    main()
