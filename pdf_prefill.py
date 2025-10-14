# pdf_prefill.py
# Robust prefill for AcroForm & non-form PDFs using PyMuPDF,
# with Section / Page / Index control, stable widget order,
# field-based occurrence counting, Yes/No option handling,
# and checkbox-square detection for bullet lists.

import argparse
import csv
import fitz  # PyMuPDF
import json
import math
import os
import pandas as pd
import re
import string
import unicodedata
from difflib import SequenceMatcher
from typing import List, Dict, Any, Tuple, Optional, Set
# -- Word→PDF adapter (non-invasive) -----------------------------------------
import tempfile, shutil, subprocess
from contextlib import contextmanager
# --- DOCX support ---
from typing import Iterable
from numbers import Number
@contextmanager
def _as_pdf(input_path: str):
    """
    Yield a path to a PDF version of `input_path`.
    If it's already a PDF, yield it as-is.
    If it's a Word doc (.docx/.doc), attempt conversion via:
      1) docx2pdf  (Word on Windows/macOS)
      2) LibreOffice 'soffice' (cross-platform)
    Cleans up temp files afterwards.
    """
    ext = os.path.splitext(input_path)[1].lower()
    if ext == ".pdf":
        yield input_path
        return

    if ext not in (".docx", ".doc"):
        raise ValueError(f"Unsupported input: {ext}. Expected .pdf, .docx, or .doc")

    tmpdir = tempfile.mkdtemp(prefix="prefill2pdf_")
    try:
        base = os.path.splitext(os.path.basename(input_path))[0]
        out_pdf = os.path.join(tmpdir, base + ".pdf")

        # 1) Try docx2pdf
        try:
            from docx2pdf import convert as _docx2pdf_convert
            _docx2pdf_convert(input_path, tmpdir)
            if os.path.exists(out_pdf):
                yield out_pdf
                return
        except Exception:
            pass

        # 2) Try LibreOffice soffice
        try:
            cmd = ["soffice", "--headless", "--convert-to", "pdf", "--outdir", tmpdir, input_path]
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            if os.path.exists(out_pdf):
                yield out_pdf
                return
        except Exception:
            pass

        raise RuntimeError(
            "Unable to convert Word document to PDF. "
            "Install 'docx2pdf' (with Word) or LibreOffice ('soffice') "
            "or provide a PDF."
        )
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


# ---------------------------
# Config / globals
# ---------------------------
PUNCT = str.maketrans("", "", string.punctuation)

# Optional global defaults (lowest priority)
SUBSCRIPTION_DEFAULTS = {
    # "Name of Investor": "John Doe",
}

# Optional label aliases (normalize both sides)
FIELD_ALIASES = {
    "investor name": "name of investor",
    "name of investor printed or typed": "name of investor",
    # add more if needed...
}

# ---------------------------
# String / matching utils
# ---------------------------
def _normalize(s: str) -> str:
    s = str(s or "")
    s = unicodedata.normalize("NFKC", s)
    s = re.sub(r"\(.*?\)", "", s)
    s = re.sub(r"\s+", " ", s).strip().lower()
    s = s.translate(PUNCT)
    return s

def _sim(a: str, b: str) -> float:
    return SequenceMatcher(None, _normalize(a), _normalize(b)).ratio()

def _best_match_scored(label: str, candidates: List[str]) -> Tuple[Optional[str], float]:
    if not candidates:
        return None, 0.0
    best, best_score = None, 0.0
    for c in candidates:
        score = _sim(label, c)
        if score > best_score:
            best, best_score = c, score
    return best, best_score

def alias_normal(norm_label: str) -> str:
    return FIELD_ALIASES.get(norm_label, norm_label)

# truth helpers for Yes/No
def _truthy(val: str) -> Optional[bool]:
    s = (str(val or "")).strip().lower()
    if s in {"y", "yes", "true", "1", "x", "✓", "check", "checked"}:
        return True
    if s in {"n", "no", "false", "0", "uncheck", "unchecked"}:
        return False
    return None

def _rect_key(x0, y0, x1, y1):
    # coarse rounding to avoid tiny numeric jitter
    return (int(round(x0)), int(round(y0)), int(round(x1)), int(round(y1)))



# ---------------------------
# Lookup loader (Excel/CSV) with Section/Page/Index
# ---------------------------


def to_text_value(v) -> str:
    """
    Robustly convert Excel/CSV values to clean text:
      - ints stay ints ("123456")
      - floats lose trailing .0 / zeros ("123456.5", "123456")
      - NaN/None -> ""
      - strings like "123456.0" -> "123456"
    """
    # Pandas NaN / None
    try:
        import pandas as _pd
        if _pd.isna(v):
            return ""
    except Exception:
        pass
    if v is None:
        return ""

    # Numeric types (covers Python & NumPy via Number)
    if isinstance(v, Number):
        # guard weird non-finite floats
        try:
            if isinstance(v, float) and not math.isfinite(v):
                return ""
        except Exception:
            pass
        # integer-like
        try:
            if float(v).is_integer():
                return str(int(v))
        except Exception:
            pass
        # general float formatting, trim trailing zeros
        s = f"{float(v):.12f}".rstrip("0").rstrip(".")
        return s or "0"

    # Strings – normalize common “.0” artifacts
    s = str(v).strip()
    if not s:
        return ""
    # if it looks like 123.000 -> 123 ;  123.4500 -> 123.45
    m = re.fullmatch(r"([+-]?\d+)\.0+\b", s)
    if m:
        return m.group(1)
    m = re.fullmatch(r"([+-]?\d+\.\d*?[1-9])0+\b", s)
    if m:
        return m.group(1)
    return s



def read_lookup_rows(path: str) -> List[Dict[str, Any]]:
    """
    Returns rows with keys:
      Field, Value, Section (optional), Page (optional, int), Index (optional, int)
    Adds: field_norm, section_norm
    """
    if not os.path.exists(path):
        print(f"⚠️  Lookup file not found: {path}")
        return []

    # Load
    try:
        if path.lower().endswith((".xlsx", ".xls")):
            df = pd.read_excel(path)
        else:
            df = pd.read_csv(path)
    except Exception as e:
        print(f"⚠️  Could not load {path}: {e}")
        return []

    # Validate minimal columns
    if {"Field", "Value"} - set(df.columns):
        raise ValueError("Lookup must have at least columns: Field, Value. Optional: Section, Page, Index")

    # Clean text cols
    for col in ["Field", "Value", "Section"]:
        if col in df.columns:
            df[col] = df[col].astype(str).fillna("").map(lambda x: x.strip())

    # Page → Int
    if "Page" in df.columns:
        df["Page"] = pd.to_numeric(df["Page"], errors="coerce").astype("Int64")

    # Index → Int (after stripping junk like '#3')
    if "Index" in df.columns:
        idx_series = (
            df["Index"]
            .astype(str)
            .str.replace(r"[^\d\-]+", "", regex=True)
            .replace({"": None})
        )
        df["Index"] = pd.to_numeric(idx_series, errors="coerce").astype("Int64")

    # Drop empty values
    df = df[(df["Field"].astype(str).str.strip() != "") &
            (df["Value"].astype(str).str.strip() != "") &
            (df["Value"].astype(str).str.lower() != "nan")]

    rows: List[Dict[str, Any]] = []

    # ...
    for _, r in df.iterrows():
        field = str(r.get("Field", "")).strip()
        # 👇 use the cleaner
        value = to_text_value(r.get("Value", ""))
        section = str(r.get("Section", "")).strip() if "Section" in df.columns else ""
        page = int(r["Page"]) if ("Page" in df.columns and pd.notna(r["Page"])) else None
        index = int(r["Index"]) if ("Index" in df.columns and pd.notna(r["Index"])) else None

        rows.append({
            "Field": field,
            "Value": value,
            "Section": section,
            "Page": page,
            "Index": index,
            "field_norm": alias_normal(_normalize(field)),
            "section_norm": _normalize(section) if section else "",
        })


    # Append defaults (lowest priority)
    for k, v in SUBSCRIPTION_DEFAULTS.items():
        rows.append({
            "Field": k,
            "Value": v,
            "Section": "",
            "Page": None,
            "Index": None,
            "field_norm": alias_normal(_normalize(k)),
            "section_norm": "",
        })
    return rows

def expected_section_set(lookup_rows: List[Dict[str, Any]]) -> Set[str]:
    return {r["section_norm"] for r in lookup_rows if r.get("section_norm")}

# ---------------------------
# DOCX (native) prefill
# ---------------------------
def _iter_all_paragraphs_and_cells(doc) -> Iterable:
    """Yield every paragraph-like container: doc paragraphs + each table cell's paragraphs."""
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def _field_value_for_name(lookup_rows, name: str):
    """Resolve a value for a simple field name, ignoring page/section/index (best-effort)."""
    v = resolve_value(
        rows=lookup_rows,
        field_label=name,
        page=None,
        section_norm="",
        occurrence_index=1,
        min_field_fuzzy=0.82,
        return_row=False,
        strict_index=False,
        require_page_match=False,
        require_section_match=False,
    )
    return v

def _replace_placeholders_in_text(text: str, lookup_rows) -> str:
    """
    Replace {{Field}}, [[Field]], ${Field} placeholders using lookup values.
    Keeps original text otherwise.
    """
    import re
    patterns = [
        re.compile(r"\{\{\s*(.*?)\s*\}\}"),
        re.compile(r"\[\[\s*(.*?)\s*\]\]"),
        re.compile(r"\$\{\s*(.*?)\s*\}"),
    ]
    def _sub_one(m):
        key = (m.group(1) or "").strip()
        val = _field_value_for_name(lookup_rows, key)
        return str(val) if val is not None else m.group(0)
    for pat in patterns:
        text = pat.sub(_sub_one, text)
    return text

def _set_paragraph_text_keep_simple_format(p, new_text: str):
    """
    Safer than 'p.text = ...' when we want to keep the paragraph object alive.
    We rebuild runs minimally. Complex inline styles may be lost, but this is robust.
    """
    while p.runs:
        p.runs[0].clear()
        p.runs[0].text = ""
        p.runs[0].element.getparent().remove(p.runs[0].element)
    r = p.add_run(new_text)

def _fill_table_right_cells(doc, lookup_rows, dry_run=False):
    """
    If a cell looks like a label (short, ends with ':' or not too long),
    write the value into the immediate right cell (if exists and empty-ish).
    """
    import re
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = row.cells
            for i in range(len(cells) - 1):
                left = cells[i].text.strip()
                right = cells[i+1]
                if not left:
                    continue
                # Prefer lines that look like labels
                lab = left
                lab = re.sub(r"[:：﹕꞉˸፡︓]\s*$", "", lab).strip()
                if not lab or len(lab) > 80:
                    continue
                val = _field_value_for_name(lookup_rows, lab)
                if val is None:
                    continue
                if dry_run:
                    print(f"[DRY][DOCX] table: '{lab}' → '{val}'")
                else:
                    # only replace if the right cell is empty-ish
                    if right.text.strip() in {"", "________", "_____"} or True:
                        # overwrite regardless; comment the 'or True' to make it conditional
                        right.text = str(val)

def _fill_colon_underscore_lines(doc, lookup_rows, dry_run=False):
    """
    For paragraphs like 'Name: ________', replace the underscores with the value.
    """
    import re
    us_pat = re.compile(r"^(.*?[:：﹕꞉˸፡︓])\s*_+\s*$")
    for p in _iter_all_paragraphs_and_cells(doc):
        txt = p.text.strip()
        if not txt:
            continue
        m = us_pat.match(txt)
        if not m:
            continue
        label_full = m.group(1)
        lab = re.sub(r"[:：﹕꞉˸፡︓]\s*$", "", label_full).strip()
        if not lab:
            continue
        val = _field_value_for_name(lookup_rows, lab)
        if val is None:
            continue
        new_text = f"{label_full} {val}"
        if dry_run:
            print(f"[DRY][DOCX] underline: '{lab}' → '{val}'")
        else:
            _set_paragraph_text_keep_simple_format(p, new_text)

def _fill_checkboxes(doc, lookup_rows, dry_run=False):
    """
    Very simple checkbox logic:
    - Turns '☐ Label' / '[ ] Label' into '☒ Label' / '[x] Label' if value is truthy.
      If value is falsey -> keep empty box.
    - Matching tries the text after the box as a field name.
    NOTE: This does NOT manipulate Word content-control checkboxes (rich XML).
    """
    import re
    box_patterns = [
        re.compile(r"^\s*[□☐]\s+(.*)$"),           # unicode empty box
        re.compile(r"^\s*\[\s?\]\s+(.*)$"),        # [ ] Label
        re.compile(r"^\s*\[\s?[xX✓]\s?\]\s+(.*)$") # already ticked; we re-evaluate
    ]
    for p in _iter_all_paragraphs_and_cells(doc):
        t = p.text.strip()
        if not t:
            continue
        for pat in box_patterns:
            m = pat.match(t)
            if not m:
                continue
            lab = m.group(1).strip()
            if not lab:
                break
            val = _field_value_for_name(lookup_rows, lab)
            yn = _truthy(val)
            if yn is None:
                # if label contains Yes/No token try to be clever
                # e.g., '☐ Yes'/'☐ No' under a section; skip unless explicit
                continue
            new = (f"☒ {lab}" if yn else f"☐ {lab}")
            if dry_run:
                print(f"[DRY][DOCX] checkbox: '{lab}' → {'CHECK' if yn else 'UNCHECK'}")
            else:
                _set_paragraph_text_keep_simple_format(p, new)
            break  # only first pattern

def _replace_placeholders_everywhere(doc, lookup_rows, dry_run=False):
    """
    Replace placeholders in all paragraphs and table cells.
    """
    for p in _iter_all_paragraphs_and_cells(doc):
        old = p.text
        new = _replace_placeholders_in_text(old, lookup_rows)
        if new != old:
            if dry_run:
                print(f"[DRY][DOCX] placeholder: '{old}' → '{new}'")
            else:
                _set_paragraph_text_keep_simple_format(p, new)

# ---------- DOCX helpers: block order + cell text ----------
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

def _iter_block_items(parent):
    """
    Yield paragraphs and tables in document order.
    """
    parent_elm = parent._element
    for child in parent_elm.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def _cell_text(cell: _Cell) -> str:
    # Consolidate runs, strip whitespace/newlines
    return "\n".join(p.text for p in cell.paragraphs).strip()

def _set_cell_text(cell: _Cell, text: str):
    # Replace all content with a single paragraph containing 'text'
    for p in list(cell.paragraphs):
        p.clear()
    cell.text = str(text)

def _looks_like_section_title(text: str) -> bool:
    if not text:
        return False
    t = unicodedata.normalize("NFKC", text).strip()
    if len(t) > 180:
        return False
    # headings or shouty short lines or lines ending with colon
    letters = [c for c in t if c.isalpha()]
    caps_ratio = (sum(1 for c in letters if c.isupper()) / max(1, len(letters))) if letters else 0.0
    return t.endswith(":") or (caps_ratio >= 0.45 and len(t.split()) <= 15)

# ---------- Exact, write-once DOCX table filler ----------
def _fill_table_right_cells(doc, lookup_rows, dry_run: bool = False, label_col: int = 0, value_col: int = 1):
    """
    For each table:
      - detect a section title from recent paragraphs above
      - for each row, use the left cell as label, write value into the right cell
      - EXACT field match (normalized), no fuzzy reuse
      - respect explicit Index in Excel; otherwise count occurrences per (Section, Field)
      - write each (Section, Field, Index) at most once
    """
    written_once = set()  # (section_norm, field_norm, idx)
    occ_counters: Dict[Tuple[str, str], int] = {}
    used_indices: Dict[Tuple[str, str], Set[int]] = {}

    # Walk document in order so we can find the paragraph(s) immediately above each table
    last_section_text = ""
    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            txt = (block.text or "").strip()
            if txt:
                last_section_text = txt
            continue

        # It's a table
        table: Table = block
        # section detection: use the last non-empty paragraph if it "looks like" a title
        section_text = last_section_text if _looks_like_section_title(last_section_text) else ""
        section_norm = _normalize(_strip_colon_like(section_text)) if section_text else ""

        # Iterate rows; treat first column as label, second as value
        for row in table.rows:
            # Guard against short rows
            if len(row.cells) <= max(label_col, value_col):
                continue

            label = _cell_text(row.cells[label_col])
            if not label:
                continue

            # Normalize the field key
            field_norm = alias_normal(_normalize(label))

            # Compute occurrence index: explicit first, else rolling per (section, field)
            bucket = (section_norm, field_norm)
            explicit = _explicit_indices_for(lookup_rows, field_norm, page=None, section_norm=section_norm)
            if explicit:
                used = used_indices.setdefault(bucket, set())
                next_idx = next((i for i in explicit if i not in used), None)
                if next_idx is None:
                    occ_counters[bucket] = occ_counters.get(bucket, 0) + 1
                    next_idx = occ_counters[bucket]
                used.add(next_idx)
                idx = next_idx
            else:
                occ_counters[bucket] = occ_counters.get(bucket, 0) + 1
                idx = occ_counters[bucket]

            logical_key = (section_norm, field_norm, idx)
            if logical_key in written_once:
                continue

            # EXACT match only (min_field_fuzzy ~ exact after normalization)
            val = resolve_value(
                lookup_rows,
                field_label=label,
                page=None,
                section_norm=section_norm,
                occurrence_index=idx,
                min_field_fuzzy=0.999,
                strict_index=True,
                require_page_match=False,
                require_section_match=bool(section_norm),
            )
            if val is None:
                # Try relaxing section requirement if we didn't detect a section
                if not section_norm:
                    val = resolve_value(
                        lookup_rows,
                        field_label=label,
                        page=None,
                        section_norm="",
                        occurrence_index=idx,
                        min_field_fuzzy=0.999,
                        strict_index=True,
                        require_page_match=False,
                        require_section_match=False,
                    )

            if val is None:
                continue

            # Write once
            if dry_run:
                print(f"[DRY][DOCX] '{label}' (sec='{section_text or ''}', idx={idx}) -> {val}")
                written_once.add(logical_key)
                continue

            _set_cell_text(row.cells[value_col], str(val))
            written_once.add(logical_key)


def sanitize_lookup_for_docx(lookup_rows: List[Dict[str, Any]]):
    """
    Generic sanitizer for DOCX filling:
      • Ignore Page (Word reflows)
      • De-duplicate by (section_norm, field_norm, Index), keep first non-empty Value
      • Merge Choices when present
    Returns: (sanitized_rows, report_dict)
    """
    total_rows = len(lookup_rows)

    # Page-agnostic copy, normalize keys
    page_free_rows: List[Dict[str, Any]] = []
    for r in lookup_rows:
        d = dict(r)
        d["Page"] = None  # DOCX never uses Page
        # normalize Index
        try:
            d["Index"] = int(d["Index"]) if d.get("Index") is not None else None
        except Exception:
            d["Index"] = None
        # ensure norms
        if not d.get("field_norm"):
            d["field_norm"] = alias_normal(_normalize(d.get("Field", "")))
        if "section_norm" not in d:
            d["section_norm"] = _normalize(d.get("Section", "")) if d.get("Section") else ""
        page_free_rows.append(d)

    # De-dupe
    def _idx_or_one(x):
        try:
            return int(x) if x is not None else 1
        except Exception:
            return 1

    dedup: Dict[Tuple[str, str, int], Dict[str, Any]] = {}
    conflicts: List[Tuple[str, str, int, str, str]] = []

    for d in page_free_rows:
        sec = d.get("section_norm", "") or ""
        fld = d.get("field_norm", "") or alias_normal(_normalize(d.get("Field", "")))
        idx = _idx_or_one(d.get("Index"))
        key = (sec, fld, idx)
        v = str(d.get("Value", "")).strip()

        if key not in dedup:
            dedup[key] = d
        else:
            kept = dedup[key]
            kept_v = str(kept.get("Value", "")).strip()
            # prefer non-empty values
            if kept_v == "" and v != "":
                dedup[key] = d
            elif v != "" and kept_v != "" and v != kept_v:
                conflicts.append((sec, fld, idx, kept_v, v))
            # merge choices if needed
            if "Choices" in d and d["Choices"] and not kept.get("Choices"):
                kept["Choices"] = d["Choices"]

    sanitized_rows = list(dedup.values())
    report = {
        "total_rows": total_rows,
        "kept_rows": len(sanitized_rows),
        "deduped": total_rows - len(sanitized_rows),
        "conflicts": conflicts,
    }
    return sanitized_rows, report


def prefill_docx(input_docx: str,
                 output_docx: str,
                 lookup_rows: List[Dict[str, Any]],
                 dry_run: bool = False):
    """
    Native DOCX fill (no conversion). Strategies:
      1) Replace placeholders: {{Field}}, [[Field]], ${Field}
      2) Table label (left cell) -> write value in the right cell
      3) 'Label: ________' underline-style lines
      4) Basic checkboxes '☐ Label' / '[ ] Label' -> checked if value truthy
    """
    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError(
            "python-docx is required for native DOCX filling. Install with: pip install python-docx"
        ) from e

    # -----------------------
    # Canonicalize lookup rows
    # -----------------------
    def _norm(s: str) -> str:
        import unicodedata, re, string
        s = unicodedata.normalize("NFKC", str(s or ""))
        s = re.sub(r"\s+", " ", s).strip().lower()
        return s.translate(str.maketrans("", "", string.punctuation))

    def _val_to_text(v):
        if isinstance(v, float):
            if v.is_integer():
                return str(int(v))
            s = f"{v:.6f}".rstrip("0").rstrip(".")
            return s if s else "0"
        return str(v).strip()

    rows_in = len(lookup_rows or [])
    seen, conflicts, rows_use = {}, {}, []

    for r in (lookup_rows or []):
        sec_raw  = r.get("Section") or r.get("section") or r.get("section_norm") or ""
        fld_raw  = r.get("Field")   or r.get("label")   or r.get("label_norm")   or ""
        idx_raw  = r.get("Index")   or r.get("index")   or 1
        try:
            idx = int(idx_raw)
        except Exception:
            idx = 1

        sec_norm = _norm(sec_raw)
        fld_norm = _norm(fld_raw)
        if not fld_norm:
            continue

        val = _val_to_text(r.get("Value", ""))

        key = (sec_norm, fld_norm, idx)
        if key in seen:
            prior_val = seen[key]
            if val and prior_val and val != prior_val:
                conflicts.setdefault(key, set()).update([prior_val, val])
            continue

        seen[key] = val
        rows_use.append({
            "Section": sec_raw,
            "section_norm": sec_norm,
            "Field":   fld_raw,
            "field_norm": fld_norm,
            "Index": idx,
            "Value": val,
        })

    lookup_rows = rows_use

    print(f"📋 DOCX prefill (generic): rows in -> {rows_in}, kept -> {len(rows_use)}, deduped -> {rows_in - len(rows_use)}")
    if conflicts:
        print(f"   • {len(conflicts)} conflicting non-empty value key(s) detected; first match wins in DOCX mode.")
        for (sec_norm, fld_norm, idx), vals in list(conflicts.items())[:20]:
            sec_disp = sec_norm or "(no section)"
            print(f"     - [{sec_disp}] {fld_norm} (Index {idx}) : {sorted(vals)}")

    # -----------------------
    # Do the actual filling
    # -----------------------
    doc = Document(input_docx)

    # 1) placeholders
    _replace_placeholders_everywhere(doc, lookup_rows, dry_run=dry_run)
    # 2) tables (left label -> right value)
    _fill_table_right_cells(doc, lookup_rows, dry_run=dry_run)
    # 3) colon + underscores
    _fill_colon_underscore_lines(doc, lookup_rows, dry_run=dry_run)
    # 4) checkboxes
    _fill_checkboxes(doc, lookup_rows, dry_run=dry_run)

    # --- Section-aware table refill (handles repeated labels across sections) ---
    from collections import defaultdict

    def _norm_key(s: str) -> str:
        import unicodedata, re, string
        s = unicodedata.normalize("NFKC", str(s or "")).strip().lower()
        s = re.sub(r"\s+", " ", s)
        return s.translate(str.maketrans("", "", string.punctuation))

    def _to_text_value(v):
        try:
            import math
            if isinstance(v, float) and math.isfinite(v):
                if v.is_integer():
                    return str(int(v))
                s = f"{v:.6f}".rstrip("0").rstrip(".")
                return s if s else "0"
        except Exception:
            pass
        return str(v).strip()

    values_exact = {}
    values_by_field = defaultdict(dict)
    values_sectionless = defaultdict(dict)

    for r in lookup_rows:
        sec = _norm_key(r.get("Section", ""))
        fld = _norm_key(r.get("Field", ""))
        idx = int(r.get("Index") or 1)
        val = _to_text_value(r.get("Value", ""))
        if not fld or val == "":
            continue
        values_exact[(sec, fld, idx)] = val
        if not sec:
            values_sectionless[fld][idx] = val
        if idx not in values_by_field[fld]:
            values_by_field[fld][idx] = val

    def _first_cell_text(cell):
        return " ".join(p.text for p in cell.paragraphs).strip()

    def _guess_pair_headers(tbl):
        hdr = {}
        try:
            row0 = tbl.rows[0]
            ncols0 = len(row0.cells)
            for ci in range(0, ncols0, 2):
                try:
                    hdr[ci] = _first_cell_text(row0.cells[ci]).strip()
                except Exception:
                    hdr[ci] = ""
        except Exception:
            pass
        return hdr

    def _strip_trailing_paren(s: str) -> str:
        import re
        return re.sub(r"\s*\([^()]*\)\s*$", "", s or "").strip()

    # --- scan upwards for a non-empty label in same column (fixes merged/blank labels) ---
    def _scan_up_label(tbl, ri: int, ci: int) -> str:
        try:
            hdr = _first_cell_text(tbl.rows[0].cells[ci]).strip()
            if hdr:
                return hdr
        except Exception:
            pass
        for rj in range(ri - 1, -1, -1):
            try:
                t = _first_cell_text(tbl.rows[rj].cells[ci]).strip()
            except Exception:
                t = ""
            if t:
                return t
        return ""

    # -- write the value inside the form "box" in a table cell.
    #    Priority: underscores → shaded paragraph → cell shading → first box-like para → last para
    #    Keeps labels & grey boxes; removes visible "FORMTEXT" remnants only.
    def _write_into_form_area(cell, text: str):
        import re
        val = text or ""

        def _para_is_shaded(para) -> bool:
            try:
                return bool(para._element.xpath('.//w:pPr/w:shd', namespaces=para._element.nsmap))
            except Exception:
                return False

        def _cell_is_shaded(c) -> bool:
            try:
                return bool(c._tc.xpath('.//w:tcPr/w:shd', namespaces=c._tc.nsmap))
            except Exception:
                return False

        target_para = None
        shaded_paras, box_like_paras, placeholder_runs = [], [], []

        for para in cell.paragraphs:
            if _para_is_shaded(para):
                shaded_paras.append(para)

            # remove visible FORMTEXT tokens only
            for run in list(para.runs):
                if (run.text or "").strip().upper() == "FORMTEXT":
                    run.text = ""

            # collect placeholder runs (______ etc.)
            for run in para.runs:
                t = (run.text or "")
                if t and re.fullmatch(r"[_\-\u2014\.\s]+", t):
                    L = len(t.replace(" ", ""))
                    placeholder_runs.append((L, run))

            vis = "".join((r.text or "") for r in para.runs)
            if vis and not re.search(r"[A-Za-z0-9]", vis):
                box_like_paras.append(para)

        # 1) Prefer the longest underscore placeholder
        if placeholder_runs:
            placeholder_runs.sort(key=lambda x: x[0], reverse=True)
            L, run = placeholder_runs[0]
            run.text = val[: max(1, L)]
            return

        # 2) Shaded paragraph
        if shaded_paras:
            target_para = shaded_paras[0]
        # 3) Cell-level shading → first paragraph
        elif _cell_is_shaded(cell):
            target_para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        # 4) First “box-like” paragraph
        elif box_like_paras:
            target_para = box_like_paras[0]
        # 5) Fallback
        elif cell.paragraphs:
            target_para = cell.paragraphs[-1]
        else:
            target_para = cell.add_paragraph()

        empties = [r for r in target_para.runs if not (r.text or "").strip()]
        if empties:
            empties[-1].text = val
        else:
            target_para.add_run(val)

    # ------------------------
    # Pass 1: multi-pair grids
    # ------------------------
    import re as _re

    for tbl in doc.tables:
        try:
            ncols = len(tbl.rows[0].cells)
        except Exception:
            continue
        if ncols < 2:
            continue

        pair_headers = _guess_pair_headers(tbl)
        start_row = 1 if any(_first_cell_text(tbl.rows[0].cells[c]).strip()
                             for c in range(min(ncols, 2))) else 0

        for ri in range(start_row, len(tbl.rows)):
            row = tbl.rows[ri]
            for ci in range(0, ncols - 1, 2):
                try:
                    label_cell = row.cells[ci]
                    value_cell = row.cells[ci + 1]
                except Exception:
                    continue

                label_text = _first_cell_text(label_cell).strip()
                if not label_text:
                    label_text = _scan_up_label(tbl, ri, ci).strip()

                fld_norm = _norm_key(label_text)
                if not fld_norm:
                    continue

                sec_display = (pair_headers.get(ci, "") or "").strip()
                sec_norm = _norm_key(_strip_trailing_paren(sec_display))

                # derive index: column pair + number in header (e.g., "Contact 3")
                pair_index_hint = (ci // 2) + 1
                m = _re.search(r'(\d+)\b', sec_display)
                if m:
                    try:
                        pair_index_hint = int(m.group(1))
                    except Exception:
                        pass

                # --- choose value with robust fallbacks ---
                chosen_value = ""
                # 1) exact (section, field, index) – try header-derived index first
                for idx_try in (pair_index_hint, 1, 2, 3, 4):
                    chosen_value = values_exact.get((sec_norm, fld_norm, idx_try), "")
                    if chosen_value:
                        break
                # 2) sectionless fallback
                if not chosen_value:
                    for idx_try in (pair_index_hint, 1, 2, 3, 4):
                        chosen_value = values_sectionless.get(fld_norm, {}).get(idx_try, "")
                        if chosen_value:
                            break
                # 3) field-only fallback by index
                if not chosen_value:
                    for idx_try in (pair_index_hint, 1, 2, 3, 4):
                        chosen_value = values_by_field.get(fld_norm, {}).get(idx_try, "")
                        if chosen_value:
                            break
                # 4) FINAL GREEDY FALLBACK: take any available index for this field (smallest index)
                if not chosen_value:
                    idx_map = values_by_field.get(fld_norm, {})
                    if idx_map:
                        chosen_value = (idx_map.get(pair_index_hint, "")
                                        or idx_map.get(1, "")
                                        or idx_map.get(2, "")
                                        or idx_map.get(3, "")
                                        or next((v for _, v in sorted(idx_map.items(), key=lambda kv: kv[0])), ""))

                if chosen_value:
                    _write_into_form_area(value_cell, chosen_value)

    # ------------------------
    # Pass 2: simple 2-col grids
    # ------------------------
    for tbl in doc.tables:
        try:
            ncols = len(tbl.rows[0].cells)
        except Exception:
            continue
        if ncols < 2:
            continue

        sec_display = _first_cell_text(tbl.rows[0].cells[0]) if tbl.rows else ""
        sec_norm = _norm_key(_strip_trailing_paren(sec_display))

        for ri, row in enumerate(tbl.rows):
            if ri == 0:
                continue
            try:
                left = _first_cell_text(row.cells[0])
                right_cell = row.cells[1]
            except Exception:
                continue

            fld_norm = _norm_key(left)
            if not fld_norm:
                # scan up in column 0 for a label if current left is blank
                left_up = _scan_up_label(tbl, ri, 0)
                fld_norm = _norm_key(left_up)
                if not fld_norm:
                    continue

            # choose value with the same robust fallbacks
            chosen_value = ""
            for idx_try in (1, 2, 3, 4):
                chosen_value = values_exact.get((sec_norm, fld_norm, idx_try), "")
                if chosen_value:
                    break
            if not chosen_value:
                for idx_try in (1, 2, 3, 4):
                    chosen_value = values_sectionless.get(fld_norm, {}).get(idx_try, "")
                    if chosen_value:
                        break
            if not chosen_value:
                for idx_try in (1, 2, 3, 4):
                    chosen_value = values_by_field.get(fld_norm, {}).get(idx_try, "")
                    if chosen_value:
                        break
            # FINAL GREEDY FALLBACK
            if not chosen_value:
                idx_map = values_by_field.get(fld_norm, {})
                if idx_map:
                    chosen_value = (idx_map.get(1, "")
                                    or idx_map.get(2, "")
                                    or idx_map.get(3, "")
                                    or next((v for _, v in sorted(idx_map.items(), key=lambda kv: kv[0])), ""))

            if chosen_value:
                _write_into_form_area(right_cell, chosen_value)

    # ---- finish ----
    if dry_run:
        print("Dry-run complete (DOCX). No file written.")
        return 0

    doc.save(output_docx)
    print(f"📝 DOCX write complete → {output_docx}")
    return 1





# ---------------------------
# Text / section helpers
# ---------------------------
_COLON_CHARS = {":", "：", "﹕", "꞉", "˸", "፡", "︓"}

def _strip_colon_like(s: str) -> str:
    s = s.strip()
    while s and s[-1] in _COLON_CHARS:
        s = s[:-1].rstrip()
    return s

def _text_blocks(page: fitz.Page) -> List[Dict[str, Any]]:
    blocks = []
    for b in page.get_text("blocks"):
        x0, y0, x1, y1, text, *_ = b
        blocks.append({"x0": float(x0), "y0": float(y0),
                       "x1": float(x1), "y1": float(y1),
                       "text": (text or "").strip()})
    blocks.sort(key=lambda r: (round(r["y0"], 1), round(r["x0"], 1)))
    return blocks

def _page_lines_with_fonts(page: fitz.Page) -> List[Dict[str, Any]]:
    out = []
    d = page.get_text("dict")
    for b in d.get("blocks", []):
        for l in b.get("lines", []):
            texts, sizes, y_vals = [], [], []
            x0 = None; x1 = None
            for s in l.get("spans", []):
                t = (s.get("text") or "").strip()
                if not t:
                    continue
                texts.append(t)
                sizes.append(float(s.get("size", 0.0)))
                x0 = s["bbox"][0] if x0 is None else min(x0, s["bbox"][0])
                x1 = s["bbox"][2] if x1 is None else max(x1, s["bbox"][2])
                y_vals.append((s["bbox"][1] + s["bbox"][3]) / 2.0)
            if texts:
                out.append({
                    "text": " ".join(texts).strip(),
                    "y_mid": sum(y_vals) / len(y_vals) if y_vals else 0.0,
                    "x0": x0 or 0.0,
                    "x1": x1 or 0.0,
                    "max_size": max(sizes) if sizes else 0.0
                })
    return out

def _is_section_header_relaxed(text: str) -> bool:
    if not text:
        return False
    t = unicodedata.normalize("NFKC", text).strip()
    if len(t) > 180:
        return False
    ends_like_colon = (len(t) > 5 and t[-1] in _COLON_CHARS)
    lower_t = t.lower()
    contains_keywords = (
            "for completion by subscribers" in lower_t
            or "investor information" in lower_t
            or "subscription details" in lower_t
            or "wire information" in lower_t
            or "beneficiary bank" in lower_t
            or "intermediary bank" in lower_t
            or "erisa status" in lower_t
            or "signature" in lower_t
            or "for all subscribers" in lower_t
    )
    letters = [c for c in t if c.isalpha()]
    caps_ratio = (sum(1 for c in letters if c.isupper()) / max(1, len(letters))) if letters else 0.0
    shouty = caps_ratio >= 0.45
    return ends_like_colon or contains_keywords or (shouty and len(t.split()) <= 15)

def find_sections_on_page(page: fitz.Page,
                          expected_sections_norm: Optional[Set[str]] = None,
                          dry_run: bool = False) -> List[Dict[str, Any]]:
    import re, unicodedata
    expected_sections_norm = expected_sections_norm or set()
    lines = _page_lines_with_fonts(page)
    candidates: List[Dict[str, Any]] = []

    # ---- helpers (local; no hard-coded names) --------------------------------
    def _strip_trailing_paren(s: str) -> str:
        # drop ONE trailing (...) — common in DOCX table headers
        return re.sub(r"\s*\([^()]*\)\s*$", "", str(s or "")).strip()

    def _norm_key(s: str) -> str:
        # mirror your _normalize but safe here
        t = unicodedata.normalize("NFKC", str(s or "")).lower().strip()
        t = re.sub(r"\s+", " ", t)
        return re.sub(r"[^\w\s]", "", t)

    # ---- (0) geometry-aware pass using text dict (handles table headers) -----
    try:
        td = page.get_text("dict") or {}
        page_w = float(page.rect.x1) or 1.0
        for b in td.get("blocks", []) or []:
            for ln in b.get("lines", []) or []:
                spans = [s for s in (ln.get("spans") or []) if (s.get("text") or "").strip()]
                if not spans:
                    continue
                x0 = min(float(s["bbox"][0]) for s in spans)
                x1 = max(float(s["bbox"][2]) for s in spans)
                y0 = min(float(s["bbox"][1]) for s in spans)
                y1 = max(float(s["bbox"][3]) for s in spans)
                txt_raw = " ".join((s.get("text") or "").strip() for s in spans).strip()
                if not txt_raw:
                    continue

                core = _strip_trailing_paren(unicodedata.normalize("NFKC", txt_raw))
                if not core:
                    continue

                # textual filters (digit-free, not extremely long)
                if any(ch.isdigit() for ch in core):
                    continue
                words = core.split()
                if not words or len(words) > 16:
                    continue

                # geometry: treat left-aligned, wide spans as headers too
                span_frac = max(0.0, (x1 - x0) / max(1.0, page_w))
                titleish = (sum(1 for w in words if w[:1].isupper()) / max(1, len(words))) >= 0.5
                wide_enough = span_frac >= 0.45  # relaxed for DOCX table rows

                if wide_enough or titleish:
                    candidates.append({
                        "name": txt_raw,                      # keep original (with parenthetical) for display
                        "name_norm": _norm_key(core),         # normalize on core
                        "y1": (y0 + y1) / 2.0                 # mid-line Y
                    })
    except Exception:
        pass

    # ---- (1) your relaxed header pass (kept) ---------------------------------
    for ln in lines:
        t = ln["text"].strip()
        if not t:
            continue
        if "\n" in t:
            t = t.splitlines()[0].strip()
        if _is_section_header_relaxed(t):
            name = _strip_colon_like(unicodedata.normalize("NFKC", t))
            candidates.append({"name": name, "name_norm": _normalize(name), "y1": ln["y_mid"]})

    # ---- (2) fuzzy vs expected sections (kept) --------------------------------
    if expected_sections_norm:
        expected_list = list(expected_sections_norm)
        for ln in lines:
            norm_line = _normalize(_strip_colon_like(unicodedata.normalize("NFKC", ln["text"])))
            if len(norm_line) > 220:
                continue
            bm, sc = _best_match_scored(norm_line, expected_list)
            if bm and sc >= 0.82:
                name = _strip_colon_like(ln["text"])
                candidates.append({"name": name, "name_norm": norm_line, "y1": ln["y_mid"]})

    # ---- (3) size-based catch (kept) ------------------------------------------
    if lines:
        sizes = sorted([ln["max_size"] for ln in lines if ln["max_size"] > 0])
        if sizes:
            thresh = sizes[int(0.80 * (len(sizes) - 1))]
            for ln in lines:
                if ln["max_size"] >= thresh and len(ln["text"]) <= 150:
                    t = _strip_colon_like(ln["text"])
                    if t and _is_section_header_relaxed(t):
                        candidates.append({"name": t, "name_norm": _normalize(t), "y1": ln["y_mid"]})

    # ---- dedupe & order -------------------------------------------------------
    candidates.sort(key=lambda c: (round(c["y1"], 1), c["name_norm"]))
    dedup, seen = [], set()
    for c in candidates:
        key = (c["name_norm"], round(c["y1"], 1))
        if key in seen:
            continue
        seen.add(key)
        dedup.append(c)

    if dry_run and dedup:
        print("  sections found:")
        for s in dedup:
            print("   -", s["name"])
    return dedup


def nearest_section_name(sections: List[Dict[str, Any]], y_mid: float) -> Tuple[str, str]:
    above = [s for s in sections if s["y1"] <= y_mid + 1e-6]  # tiny tolerance
    if not above:
        return "", ""
    last = above[-1]
    return last["name"], last["name_norm"]


# ---------------------------
# Yes/No detection around a point
# ---------------------------
def _nearby_yes_no_option(page: fitz.Page, x: float, y: float) -> Optional[str]:
    """Return 'yes' or 'no' if a token is detected near (x,y)."""
    XRANGE = 140.0
    YRANGE = 24.0
    nearest, best_dx = None, 1e9
    for b in page.get_text("blocks"):
        x0, y0, x1, y1, text, *_ = b
        t = (text or "").strip()
        if not t or len(t) > 6:  # very short tokens only
            continue
        ymid = (y0 + y1) / 2.0
        if abs(ymid - y) > YRANGE:
            continue
        if (x0 - XRANGE) <= x <= (x1 + XRANGE):
            token = t.lower().strip("._:-)()]}[({")
            if token in {"yes", "no"}:
                dx = min(abs(x - x0), abs(x - x1))
                if dx < best_dx:
                    best_dx = dx
                    nearest = token
    return nearest



# ---------------------------
# Drawn underline detection (for non-form PDFs)
# ---------------------------
def _line_like_segments(
        page: fitz.Page,
        min_len=60,
        max_len=1200,
        max_slope=0.02,
        max_thick=2.5
):
    """
    Return horizontal-ish line segments on the page, tolerating variations in how
    PyMuPDF encodes drawing items. Handles:
      - ('l', (x0, y0), (x1, y1))
      - ('re', x, y, w, h, ...)  OR  ('re', (x, y, w, h))
      - ignores other commands safely
    """
    def _as_float_tuple(obj, n):
        """Try to coerce obj to an n-length tuple of floats; else None."""
        try:
            if isinstance(obj, (list, tuple)) and len(obj) >= n:
                return tuple(float(obj[i]) for i in range(n))
        except Exception:
            pass
        return None

    segs = []
    drawings = page.get_drawings() or []
    for d in drawings:
        items = d.get("items", []) or []
        for it in items:
            if not it:
                continue

            cmd = it[0]

            # ----- line segments -----
            if cmd == "l":
                # expected: ('l', (x0, y0), (x1, y1))
                if len(it) >= 3:
                    p0 = _as_float_tuple(it[1], 2)
                    p1 = _as_float_tuple(it[2], 2)
                    if p0 and p1:
                        x0, y0 = p0
                        x1, y1 = p1
                        dx, dy = (x1 - x0), (y1 - y0)
                        length = math.hypot(dx, dy)
                        # keep “horizontal-ish” lines within size bounds
                        if (min_len <= length <= max_len) and (abs(dy) / (abs(dx) + 1e-6) <= max_slope):
                            xL, xR = (x0, x1) if x0 <= x1 else (x1, x0)
                            yMid = (y0 + y1) / 2.0
                            segs.append({"x0": xL, "x1": xR, "y0": yMid, "y1": yMid, "len": length})
                continue  # move to next item

            # ----- rectangle strokes (underlines often drawn as very thin rects) -----
            if cmd == "re":
                # Accept either ('re', x, y, w, h, ...)  OR  ('re', (x, y, w, h), ...)
                x = y = w = h = None

                if len(it) >= 5:
                    # ('re', x, y, w, h, ...)
                    try:
                        x = float(it[1]); y = float(it[2]); w = float(it[3]); h = float(it[4])
                    except Exception:
                        x = y = w = h = None

                if x is None:
                    # maybe ('re', (x, y, w, h), ...)
                    if len(it) >= 2:
                        rect4 = _as_float_tuple(it[1], 4)
                        if rect4:
                            x, y, w, h = rect4

                if x is None or w is None or h is None:
                    # unsupported/odd shape – skip safely
                    continue

                if h <= max_thick and (min_len <= w <= max_len):
                    xL, xR = x, x + w
                    yMid = y + h / 2.0
                    segs.append({"x0": xL, "x1": xR, "y0": yMid, "y1": yMid, "len": w})
                continue  # move to next item

            # else: ignore other drawing commands safely and continue
            continue

    # sort + de-dup near-identical segments
    segs.sort(key=lambda s: (round(s["y0"], 1), round(s["x0"], 1)))
    merged = []
    for s in segs:
        if merged:
            m = merged[-1]
            if abs(s["y0"] - m["y0"]) < 0.8 and abs(s["x0"] - m["x0"]) < 2 and abs(s["x1"] - m["x1"]) < 2:
                continue
        merged.append(s)
    return merged


# ---------------------------
# Checkbox square detection (new)
# ---------------------------
def _square_checkboxes(page: fitz.Page,
                       min_side: float = 6.0,
                       max_side: float = 30.0,
                       max_aspect: float = 1.6) -> List[Dict[str, float]]:
    """
    Return small, nearly-square rectangles likely to be checkboxes.
    Now detects:
      - Thin/stroked rectangles ('re' items) of various encodings
      - Unicode checkbox glyphs in text spans (e.g., '☐', '□', '◻', '■')
      - Small squares reconstructed from 4 line segments
    """
    def _as_float_tuple(obj, n):
        try:
            if isinstance(obj, (list, tuple)) and len(obj) >= n:
                return tuple(float(obj[i]) for i in range(n))
        except Exception:
            pass
        return None

    boxes: List[Dict[str, float]] = []

    # ---- A) Vector rectangles ('re') with wider thresholds ----
    for d in page.get_drawings() or []:
        for it in d.get("items", []) or []:
            if not it or it[0] != "re":
                continue

            x = y = w = h = None

            # Case 1: flat scalars
            if len(it) >= 5:
                try:
                    x = float(it[1]); y = float(it[2]); w = float(it[3]); h = float(it[4])
                except Exception:
                    x = y = w = h = None

            # Case 2: Rect object
            if x is None and len(it) >= 2 and isinstance(it[1], fitz.Rect):
                r: fitz.Rect = it[1]  # type: ignore
                x, y, w, h = float(r.x0), float(r.y0), float(r.width), float(r.height)

            # Case 3: 4-tuple
            if x is None and len(it) >= 2:
                rect4 = _as_float_tuple(it[1], 4)
                if rect4:
                    x, y, w, h = rect4

            if x is None or w is None or h is None:
                continue

            side_min, side_max = min(w, h), max(w, h)
            if side_min <= 0:
                continue
            aspect = side_max / side_min
            if (min_side <= side_min <= max_side) and aspect <= max_aspect:
                boxes.append({
                    "x0": x, "y0": y, "x1": x + w, "y1": y + h,
                    "cx": x + w / 2.0, "cy": y + h / 2.0,
                    "w": w, "h": h
                })

    # ---- B) Unicode glyphs that look like empty/filled squares ----
    GLYPHS = {"☐","■","□","◻","◼","▢","❏","❐","❑"}
    td = page.get_text("dict") or {}
    for b in td.get("blocks", []):
        for l in b.get("lines", []):
            for s in l.get("spans", []):
                t = s.get("text") or ""
                if len(t.strip()) != 1:
                    continue
                ch = t.strip()
                if ch not in GLYPHS:
                    continue
                x0, y0, x1, y1 = map(float, s.get("bbox", (0, 0, 0, 0)))
                w = x1 - x0; h = y1 - y0
                side_min, side_max = min(w, h), max(w, h)
                if side_min <= 0:
                    continue
                aspect = side_max / side_min
                if (min_side <= side_min <= max_side) and aspect <= max_aspect:
                    # small padding to better match visual box
                    pad = 0.4
                    boxes.append({
                        "x0": x0 + pad, "y0": y0 + pad, "x1": x1 - pad, "y1": y1 - pad,
                        "cx": (x0 + x1) / 2.0, "cy": (y0 + y1) / 2.0,
                        "w": w, "h": h
                    })

    # ---- C) Squares traced by 4 line segments (no 're') ----
    # collect short near-horizontal and near-vertical line segments
    hseg, vseg = [], []
    for d in page.get_drawings() or []:
        for it in d.get("items", []) or []:
            if not it or it[0] != "l" or len(it) < 3:
                continue
            p0 = _as_float_tuple(it[1], 2)
            p1 = _as_float_tuple(it[2], 2)
            if not p0 or not p1:
                continue
            x0, y0 = p0; x1, y1 = p1
            dx, dy = (x1 - x0), (y1 - y0)
            L = math.hypot(dx, dy)
            if L < min_side * 0.7 or L > max_side * 1.6:
                continue
            slope = abs(dy) / (abs(dx) + 1e-6)
            if slope < 0.15:  # horizontal-ish
                y = (y0 + y1) / 2.0
                hseg.append((min(x0, x1), max(x0, x1), y))
            elif slope > 6.0:  # vertical-ish
                x = (x0 + x1) / 2.0
                vseg.append((x, min(y0, y1), max(y0, y1)))

    # try to pair them into rectangles
    def _near(a, b, tol=1.2): return abs(a - b) <= tol

    for hx0, hx1, hy in hseg:
        w = hx1 - hx0
        if w <= 0:
            continue
        for hx0b, hx1b, hyb in hseg:
            if hyb <= hy or not _near(hx0, hx0b, 2.0) or not _near(hx1, hx1b, 2.0):
                continue
            h2 = hyb - hy
            if h2 <= 0:
                continue
            side_min, side_max = min(w, h2), max(w, h2)
            if not (min_side <= side_min <= max_side):
                continue
            aspect = side_max / (side_min + 1e-6)
            if aspect > max_aspect:
                continue
            # need 2 verticals roughly at x≈hx0 and x≈hx1 spanning [hy,hyb]
            left_ok = any(_near(x, hx0, 2.0) and (yv0 <= hy + 2.0) and (yv1 >= hyb - 2.0) for x, yv0, yv1 in vseg)
            right_ok = any(_near(x, hx1, 2.0) and (yv0 <= hy + 2.0) and (yv1 >= hyb - 2.0) for x, yv0, yv1 in vseg)
            if left_ok and right_ok:
                boxes.append({
                    "x0": hx0, "y0": hy, "x1": hx1, "y1": hyb,
                    "cx": (hx0 + hx1) / 2.0, "cy": (hy + hyb) / 2.0,
                    "w": w, "h": h2
                })

    # ---- de-dupe near-identical boxes ----
    boxes.sort(key=lambda b: (round(b["y0"], 1), round(b["x0"], 1)))
    dedup: List[Dict[str, float]] = []
    for b in boxes:
        if dedup:
            m = dedup[-1]
            if (abs(b["x0"] - m["x0"]) < 1.2 and abs(b["y0"] - m["y0"]) < 1.2 and
                    abs(b["x1"] - m["x1"]) < 1.2 and abs(b["y1"] - m["y1"]) < 1.2):
                continue
        dedup.append(b)
    return dedup

def _dropdown_like_boxes(
        page: fitz.Page,
        min_w: float = 90.0,
        max_w: float = 1600.0,
        min_h: float = 10.0,
        max_h: float = 38.0,
        min_aspect: float = 4.0,
):
    """
    Detect long, shallow rectangles that look like drawn dropdowns (combo boxes)
    rather than real AcroForm CHOICE widgets.

    Heuristics:
      - vector rectangle ('re') with width >> height (aspect >= min_aspect)
      - height roughly in a typical text-field range
    """
    def _rect_ok(w, h):
        if w <= 0 or h <= 0:
            return False
        if not (min_w <= w <= max_w and min_h <= h <= max_h):
            return False
        return (w / h) >= min_aspect

    dropdowns = []

    try:
        for dr in page.get_drawings() or []:
            for it in dr.get("items", []) or []:
                if not it or it[0] != "re":
                    continue

                # Support ('re', fitz.Rect) OR ('re', x, y, w, h) OR ('re', (x,y,w,h))
                x = y = w = h = None
                try:
                    if len(it) >= 2 and isinstance(it[1], fitz.Rect):
                        r: fitz.Rect = it[1]  # type: ignore
                        x, y, w, h = float(r.x0), float(r.y0), float(r.width), float(r.height)
                    elif len(it) >= 5:
                        x, y, w, h = float(it[1]), float(it[2]), float(it[3]), float(it[4])
                    elif len(it) >= 2 and isinstance(it[1], (list, tuple)) and len(it[1]) >= 4:
                        x, y, w, h = [float(v) for v in it[1][:4]]
                except Exception:
                    x = y = w = h = None

                if x is None or w is None or h is None:
                    continue

                if _rect_ok(w, h):
                    dropdowns.append({
                        "x0": x, "y0": y, "x1": x + w, "y1": y + h,
                        "cx": x + w / 2.0, "cy": y + h / 2.0,
                        "w": w, "h": h
                    })
    except Exception:
        pass

    # De-dupe near-identical rects
    dropdowns.sort(key=lambda b: (round(b["y0"], 1), round(b["x0"], 1)))
    dedup = []
    for b in dropdowns:
        if dedup:
            m = dedup[-1]
            if (abs(b["x0"] - m["x0"]) < 1.2 and abs(b["y0"] - m["y0"]) < 1.2 and
                    abs(b["x1"] - m["x1"]) < 1.2 and abs(b["y1"] - m["y1"]) < 1.2):
                continue
        dedup.append(b)
    return dedup


# alias to tolerate either helper name in code
_checkbox_squares = _square_checkboxes

# ---------------------------
# Template builder (non-form PDFs)
# ---------------------------
def _nearest_label_for_block(ul_block: Dict[str, Any], neighbor_blocks: List[Dict[str, Any]], y_tol=22.0, back_scan=8) -> str:
    same_line = [b for b in neighbor_blocks
                 if abs(b["y0"] - ul_block["y0"]) < y_tol and b["x1"] <= ul_block["x0"] and b["text"]]
    if same_line:
        same_line.sort(key=lambda b: ul_block["x0"] - b["x1"])
        return same_line[0]["text"]

    idx = neighbor_blocks.index(ul_block) if ul_block in neighbor_blocks else len(neighbor_blocks) - 1
    cands = []
    for j in range(max(0, idx - back_scan), idx):
        b = neighbor_blocks[j]
        if not b["text"]:
            continue
        if 0 <= (ul_block["y0"] - b["y0"]) < (y_tol * 2.0):
            cands.append(b)
    if cands:
        cands.sort(key=lambda b: (abs(ul_block["y0"] - b["y0"]), abs(ul_block["x0"] - b["x1"])))
        return cands[0]["text"]
    return ""

def build_pdf_template(pdf_path: str,
                       template_json: str = "template_fields.json",
                       lookup_rows: Optional[List[Dict[str, Any]]] = None,
                       dry_run: bool = False):
    import re
    import json
    import fitz  # PyMuPDF
    from typing import Any, Dict, List, Optional, Tuple

    # --- helpers -------------------------------------------------------------
    def _section_core(name: str) -> str:
        s = (name or "").strip()
        s = re.sub(r"\s*\(.*?\)\s*$", "", s).strip()
        return s

    def _short_label(section_name: str, ordinal: int) -> str:
        core = _section_core(section_name) or "Checklist"
        return f"{core} – item {ordinal}"

    # robustly pick nearest label text from (score, block) tuples OR blocks/strings
    def _pick_nearest_text(cands):
        if not cands:
            return ""
        def score_key(item):
            if isinstance(item, (tuple, list)):
                nums = []
                for x in item:
                    if isinstance(x, (int, float)):
                        nums.append(x)
                    else:
                        break
                return tuple(nums) if nums else (0,)
            return (0,)
        cands.sort(key=score_key)
        item = cands[0]
        if isinstance(item, (tuple, list)):
            for el in reversed(item):
                if isinstance(el, dict) and "text" in el:
                    return (el.get("text") or "").strip()
                if isinstance(el, str):
                    return el.strip()
            return ""
        if isinstance(item, dict):
            return (item.get("text") or "").strip()
        if isinstance(item, str):
            return item.strip()
        return ""

    # >>>>>> added helpers (line-aware label selection) <<<<<<
    def _label_from_block(blk: Dict[str, Any], ref_y: float) -> str:
        """
        If a text block holds multiple lines, choose the line whose vertical
        position best matches the field's y (ref_y). Otherwise return block text.
        """
        if not blk:
            return ""
        t = (blk.get("text") or "").strip()
        if not t:
            return ""
        parts = [p.strip() for p in re.split(r"[\r\n]+", t) if p.strip()]
        if len(parts) == 1:
            return parts[0]

        y0 = float(blk.get("y0", ref_y))
        y1 = float(blk.get("y1", ref_y))
        if y1 <= y0:
            return parts[-1]

        rel = (ref_y - y0) / max(1e-6, (y1 - y0))
        if rel <= 0.33:
            return parts[0]
        elif rel >= 0.66:
            return parts[-1]
        else:
            return parts[-1] if (ref_y - (y0 + y1) / 2.0) >= 0 else parts[0]

    def _best_line_from_candidates(cands, ref_y: float) -> str:
        """
        Like _pick_nearest_text but returns the best single line from the chosen block,
        using ref_y to pick among multi-line labels.
        """
        if not cands:
            return ""
        def score_key(item):
            if isinstance(item, (tuple, list)):
                nums = []
                for x in item:
                    if isinstance(x, (int, float)):
                        nums.append(x)
                    else:
                        break
                return tuple(nums) if nums else (0,)
            return (0,)
        cands.sort(key=score_key)
        item = cands[0]

        blk = None
        if isinstance(item, (tuple, list)):
            for el in reversed(item):
                if isinstance(el, dict) and "text" in el:
                    blk = el
                    break
        elif isinstance(item, dict) and "text" in item:
            blk = item

        if blk is None:
            return _pick_nearest_text(cands)

        return _label_from_block(blk, ref_y)
    # <<<<<< end added helpers >>>>>>

    # >>>>>> NEW: horizontal-overlap helper (used only for "above" matches) <<<<<<<
    def _horiz_overlap(a0: float, a1: float, b0: float, b1: float,
                       min_px: float = 6.0, min_frac: float = 0.10) -> bool:
        """
        True if [a0,a1] and [b0,b1] overlap horizontally by:
          • at least min_px pixels, AND
          • at least min_frac of the smaller width.
        This fixes centered column headers being ignored.
        """
        if a1 < a0: a0, a1 = a1, a0
        if b1 < b0: b0, b1 = b1, b0
        ov = min(a1, b1) - max(a0, b0)
        if ov <= 0:
            return False
        min_w = max(1.0, min(a1 - a0, b1 - b0))
        return (ov >= min_px) and ((ov / min_w) >= min_frac)
    # <<<<<< end NEW helper >>>>>>

    # >>>>>> NEW: column header resolver by span clusters <<<<<<<
    def _column_header_by_spans(page: fitz.Page,
                                ref_y: float,
                                ref_x: float,
                                x0: Optional[float] = None,
                                x1: Optional[float] = None,
                                y_band: float = 48.0,
                                gap_px: float = 35.0) -> str:
        """
        Find the closest line *above* ref_y, split that line's spans into column clusters
        using a large x-gap (gap_px), then return the cluster whose x-range overlaps [x0,x1]
        (if provided) or whose center is nearest to ref_x.
        """
        try:
            td = page.get_text("dict") or {}
        except Exception:
            return ""

        # 1) closest line above
        best_spans = None
        best_dy = 1e9
        for b in td.get("blocks", []) or []:
            for ln in b.get("lines", []) or []:
                spans = [s for s in (ln.get("spans") or []) if (s.get("text") or "").strip()]
                if not spans:
                    continue
                ly0 = min(float(s["bbox"][1]) for s in spans)
                ly1 = max(float(s["bbox"][3]) for s in spans)
                if ly1 >= ref_y:
                    continue
                dy = ref_y - ly1
                if dy <= y_band and dy < best_dy:
                    best_spans = spans
                    best_dy = dy
        if not best_spans:
            return ""

        # 2) cluster spans by big gaps
        best_spans = sorted(best_spans, key=lambda s: (float(s["bbox"][0]), float(s["bbox"][1])))
        clusters = []
        cur = []
        last_x1 = None
        for sp in best_spans:
            sx0, sy0, sx1, sy1 = map(float, sp["bbox"])
            if last_x1 is None or (sx0 - last_x1) <= gap_px:
                cur.append(sp)
            else:
                if cur:
                    clusters.append(cur)
                cur = [sp]
            last_x1 = sx1
        if cur:
            clusters.append(cur)

        cols = []
        for cl in clusters:
            cx0 = min(float(s["bbox"][0]) for s in cl)
            cx1 = max(float(s["bbox"][2]) for s in cl)
            txt = " ".join((s.get("text") or "").strip() for s in cl if (s.get("text") or "").strip()).strip()
            if txt:
                cols.append((cx0, cx1, txt))
        if len(cols) < 2:
            return ""

        # 3) choose by overlap if x-range known, else nearest to ref_x
        if x0 is not None and x1 is not None:
            best_txt = ""
            best_ov = 0.0
            for cx0, cx1, txt in cols:
                ov = min(x1, cx1) - max(x0, cx0)
                if ov > best_ov:
                    best_ov = ov
                    best_txt = txt
            if best_txt:
                return best_txt

        best_txt = ""
        best_dist = 1e9
        for cx0, cx1, txt in cols:
            c = (cx0 + cx1) / 2.0
            d = abs(c - ref_x)
            if d < best_dist:
                best_dist = d
                best_txt = txt
        return best_txt
    # <<<<<< end NEW helper >>>>>>

    def _right_zone_has_dropdown_affordance(page: fitz.Page, rect: fitz.Rect, arrow_zone_w: float = 30.0) -> bool:
        rz = fitz.Rect(float(rect.x1 - arrow_zone_w), float(rect.y0 - 6), float(rect.x1 + 6), float(rect.y1 + 6))
        ARROWS = {"▼", "▾", "▿", "▸", "▹", "▻", "⯆", "⯈"}
        td = page.get_text("dict") or {}
        for b in td.get("blocks", []) or []:
            for ln in b.get("lines", []) or []:
                for sp in ln.get("spans", []) or []:
                    t = (sp.get("text") or "").strip()
                    if len(t) == 1 and t in ARROWS:
                        x0, y0, x1, y1 = map(float, sp.get("bbox", (0, 0, 0, 0)))
                        if fitz.Rect(x0, y0, x1, y1).intersects(rz):
                            return True
        try:
            for (xref, *_rest) in page.get_images(full=True) or []:
                for r in page.get_image_rects(xref) or []:
                    if r.intersects(rz):
                        if r.width <= 24 and (r.height <= (rect.height * 1.6)):
                            return True
        except Exception:
            pass
        lines = []
        try:
            for g in page.get_drawings() or []:
                for it in g.get("items", []) or []:
                    if not it or it[0] != "l" or len(it) < 3:
                        continue
                    p0, p1 = it[1], it[2]
                    x0 = float(getattr(p0, "x", p0[0] if isinstance(p0, (list, tuple)) else 0.0))
                    y0 = float(getattr(p0, "y", p0[1] if isinstance(p0, (list, tuple)) else 0.0))
                    x1 = float(getattr(p1, "x", p1[0] if isinstance(p1, (list, tuple)) else 0.0))
                    y1 = float(getattr(p1, "y", p1[1] if isinstance(p1, (list, tuple)) else 0.0))
                    segrect = fitz.Rect(min(x0, x1), min(y0, y1), max(x0, x1), max(y0, y1))
                    if not segrect.intersects(rz):
                        continue
                    dx, dy = (x1 - x0), (y1 - y0)
                    if min(abs(dx), abs(dy)) < 1.0:
                        continue
                    slope = abs(dy / (dx + 1e-6))
                    if 0.3 <= slope <= 3.5:
                        lines.append((x0, y0, x1, y1))
        except Exception:
            pass
        for i in range(len(lines)):
            x0a, y0a, x1a, y1a = lines[i]
            for j in range(i + 1, len(lines)):
                x0b, y0b, x1b, y1b = lines[j]
                ptsA = [(x0a, y0a), (x1a, y1a)]
                ptsB = [(x0b, y0b), (x1b, y1b)]
                share = any(abs(ax - bx) <= 3.0 and abs(ay - by) <= 3.0 for ax, ay in ptsA for bx, by in ptsB)
                if not share:
                    continue
                s1 = (y1a - y0a) * (x1a - x0a)
                s2 = (y1b - y0b) * (x1b - x0b)
                if s1 != 0 and s2 != 0 and ((s1 > 0 and s2 < 0) or (s1 < 0 and s2 > 0)):
                    return True
        return False

    def _drawn_dropdowns(page: fitz.Page,
                         min_width: float = 80.0,
                         min_aspect: float = 3.0,
                         min_h: float = 10.0,
                         max_h: float = 30.0,
                         arrow_zone_w: float = 28.0) -> List[Dict[str, float]]:
        found: List[Dict[str, float]] = []
        rects: List[fitz.Rect] = []
        try:
            for g in page.get_drawings() or []:
                for it in g.get("items", []) or []:
                    if not it or it[0] != "re":
                        continue
                    x = y = w = h = None
                    try:
                        if len(it) >= 5:
                            x = float(it[1]); y = float(it[2]); w = float(it[3]); h = float(it[4])
                        elif len(it) >= 2 and isinstance(it[1], fitz.Rect):
                            r = it[1]; x, y, w, h = r.x0, r.y0, r.width, r.height
                        elif len(it) >= 2 and isinstance(it[1], (list, tuple)) and len(it[1]) >= 4:
                            x, y, w, h = map(float, it[1][:4])
                    except Exception:
                        x = y = w = h = None
                    if x is None or w is None or h is None:
                        continue
                    if w < min_width or h < min_h or h > max_h:
                        continue
                    if (w / max(h, 1e-6)) < min_aspect:
                        continue
                    rects.append(fitz.Rect(float(x), float(y), float(x + w), float(y + h)))
        except Exception:
            pass
        if not rects:
            return found

        ARROWS = {"▼", "▾", "▿", "▸", "▹", "▻", "⯆", "⯈"}
        td = page.get_text("dict") or {}

        def has_arrow_glyph(r: fitz.Rect) -> bool:
            rz = fitz.Rect(r.x1 - arrow_zone_w, r.y0, r.x1 + 6, r.y1)
            for b in td.get("blocks", []) or []:
                for ln in b.get("lines", []) or []:
                    for sp in ln.get("spans", []) or []:
                        t = (sp.get("text") or "").strip()
                        if len(t) != 1 or t not in ARROWS:
                            continue
                        bx0, by0, bx1, by1 = map(float, sp.get("bbox", (0, 0, 0, 0)))
                        if fitz.Rect(bx0, by0, bx1, by1).intersects(rz):
                            return True
            return False

        def has_vector_chevron(r: fitz.Rect) -> bool:
            rz = fitz.Rect(r.x1 - arrow_zone_w, r.y0, r.x1, r.y1)
            lines = []
            try:
                for g in page.get_drawings() or []:
                    for it in g.get("items", []) or []:
                        if not it or it[0] != "l" or len(it) < 3:
                            continue
                        p0, p1 = it[1], it[2]
                        x0 = float(getattr(p0, "x", p0[0] if isinstance(p0, (list, tuple)) else 0.0))
                        y0 = float(getattr(p0, "y", p0[1] if isinstance(p0, (list, tuple)) else 0.0))
                        x1 = float(getattr(p1, "x", p1[0] if isinstance(p1, (list, tuple)) else 0.0))
                        y1 = float(getattr(p1, "y", p1[1] if isinstance(p1, (list, tuple)) else 0.0))
                        seg = fitz.Rect(min(x0, x1), min(y0, y1), max(x0, x1), max(y0, y1))
                        if not seg.intersects(rz):
                            continue
                        dx, dy = (x1 - x0), (y1 - y0)
                        if min(abs(dx), abs(dy)) < 1.0:
                            continue
                        slope = abs(dy / (dx + 1e-6))
                        if 0.3 <= slope <= 3.5:
                            lines.append((x0, y0, x1, y1))
            except Exception:
                pass
            for i in range(len(lines)):
                x0a, y0a, x1a, y1a = lines[i]
                for j in range(i + 1, len(lines)):
                    x0b, y0b, x1b, y1b = lines[j]
                    pts = [(x0a, y0a), (x1a, y1a)]
                    pts2 = [(x0b, y0b), (x1b, y1b)]
                    ok = any(abs(ax - bx) <= 3.0 and abs(ay - by) <= 3.0 for ax, ay in pts for bx, by in pts2)
                    if not ok:
                        continue
                    s1 = (y1a - y0a) * (x1a - x0a)
                    s2 = (y1b - y0b) * (x1b - x0b)
                    if s1 and s2 and ((s1 > 0 and s2 < 0) or (s1 < 0 and s2 > 0)):
                        return True
            return False

        for r in rects:
            if has_arrow_glyph(r) or has_vector_chevron(r):
                found.append({
                    "x0": float(r.x0), "y0": float(r.y0),
                    "x1": float(r.x1), "y1": float(r.y1),
                    "cx": float((r.x0 + r.x1) / 2.0),
                    "cy": float((r.y0 + r.y1) / 2.0),
                })
        return found

    def _tiny_vector_squares(page: fitz.Page,
                             size_min: float = 2.0,
                             size_max: float = 12.5,
                             squareness_tol: float = 0.55):
        boxes = []
        try:
            drawings = page.get_drawings() or []
        except Exception:
            return boxes
        for g in drawings:
            r = g.get("rect", None)
            if r is None:
                xs, ys = [], []
                for item in g.get("items", []) or []:
                    pts = item[1] if len(item) > 1 else None
                    if not pts:
                        continue
                    for p in pts:
                        xs.append(float(getattr(p, "x", p[0] if isinstance(p, (tuple, list)) else 0.0)))
                        ys.append(float(getattr(p, "y", p[1] if isinstance(p, (tuple, list)) else 0.0)))
                if xs and ys:
                    r = fitz.Rect(min(xs), min(ys), max(xs), max(ys))
            if r is None:
                continue
            w = float(r.x1 - r.x0); h = float(r.y1 - r.y0)
            if w <= 0 or h <= 0:
                continue
            if not (size_min <= w <= size_max and size_min <= h <= size_max):
                continue
            ar = w / h if h else 99.0
            if abs(ar - 1.0) > squareness_tol:
                continue
            boxes.append({"x0": float(r.x0), "y0": float(r.y0),
                          "x1": float(r.x1), "y1": float(r.y1),
                          "cx": float((r.x0 + r.x1) / 2.0),
                          "cy": float((r.y0 + r.y1) / 2.0)})
        return boxes

    BOX_CHARS = {"☐","■","□","◻","◼","▢","❏","❐","❑","❒"}
    def _glyph_line_checkboxes(page_dict,
                               min_side=2.4, max_side=14.0,
                               ar_low=0.40, ar_high=2.00,
                               min_gap=1.0, max_gap=140.0):
        boxes = []
        for b in page_dict.get("blocks", []) or []:
            for ln in b.get("lines", []) or []:
                spans = (ln.get("spans") or [])
                spans = sorted(spans, key=lambda s: (s["bbox"][0], s["bbox"][1]))
                first_i = None
                for i, sp in enumerate(spans):
                    if (sp.get("text") or "").strip():
                        first_i = i
                        break
                if first_i is None:
                    continue
                sp = spans[first_i]
                txt = (sp.get("text") or "").strip()
                x0, y0, x1, y1 = sp["bbox"]
                w = float(x1 - x0); h = float(y1 - y0)
                if not (min_side <= w <= max_side and min_side <= h <= max_side):
                    continue
                ar = (w / h) if h else 99.0
                fnt = (sp.get("font") or "").lower()
                looks_like_box = (len(txt) == 1 and (txt in BOX_CHARS)) or \
                                 ("dingbat" in fnt or "wingd" in fnt or "symbol" in fnt) or \
                                 (len(txt) == 1 and ar_low <= ar <= ar_high)
                if not looks_like_box:
                    continue
                right_text_ok = False
                for j in range(first_i + 1, len(spans)):
                    t2 = (spans[j].get("text") or "").strip()
                    if not t2:
                        continue
                    gap = float(spans[j]["bbox"][0]) - float(x1)
                    if min_gap <= gap <= max_gap:
                        right_text_ok = True
                    break
                if not right_text_ok:
                    continue
                boxes.append({"x0": float(x0), "y0": float(y0),
                              "x1": float(x1), "y1": float(y1),
                              "cx": float((x0 + x1) / 2.0),
                              "cy": float((y0 + y1) / 2.0)})
        return boxes

    def _large_square_boxes(page: fitz.Page,
                            min_side: float = 18.0,
                            max_side: float = 60.0,
                            ar_low: float = 0.80, ar_high: float = 1.30):
        out = []
        try:
            drs = page.get_drawings()
        except Exception:
            drs = []
        for dr in drs or []:
            for it in dr.get("items", []):
                if not it or it[0] != "re":
                    continue
                rect = it[1]
                x0, y0, x1, y1 = float(rect.x0), float(rect.y0), float(rect.x1), float(rect.y1)
                w = abs(x1 - x0); h = abs(y1 - y0)
                if (min_side <= w <= max_side) and (min_side <= h <= max_side):
                    ar = (w / h) if h else 99.0
                    if ar_low <= ar <= ar_high:
                        out.append({"x0": x0, "y0": y0, "x1": x1, "y1": y1,
                                    "cx": (x0 + x1) / 2.0, "cy": (y0 + y1) / 2.0})
        try:
            for (xref, *_rest) in page.get_images(full=True) or []:
                for r in page.get_image_rects(xref) or []:
                    x0, y0, x1, y1 = float(r.x0), float(r.y0), float(r.x1), float(r.y1)
                    w = abs(x1 - x0); h = abs(y1 - y0)
                    if (min_side <= w <= max_side) and (min_side <= h <= max_side):
                        ar = (w / h) if h else 99.0
                        if ar_low <= ar <= ar_high:
                            out.append({"x0": x0, "y0": y0, "x1": x1, "y1": y1,
                                        "cx": (x0 + x1) / 2.0, "cy": (y0 + y1) / 2.0})
        except Exception:
            pass
        return out

    def _rects_overlap(a, b, pad=1.6):
        ax0, ay0, ax1, ay1 = a
        bx0, by0, bx1, by1 = b
        ax0 -= pad; ay0 -= pad; ax1 += pad; ay1 += pad
        bx0 -= pad; by0 -= pad; bx1 += pad; by1 += pad
        return not (ax1 < bx0 or bx1 < ax0 or ay1 < by0 or by1 < ay0)

    expected = expected_section_set(lookup_rows or [])
    doc = fitz.open(pdf_path)
    tpl = {"pdf": pdf_path, "fields": []}

    underline_re = re.compile(r"_{3,}")
    counters: Dict[Tuple[int, str, str], int] = {}

    for pno, page in enumerate(doc, start=1):
        if dry_run:
            print(f"p{pno}:")
            # diagnostics
            try:
                td = page.get_text("dict") or {}
                n_blocks = len(td.get("blocks", []) or [])
            except Exception:
                n_blocks = -1
            try:
                drawings = page.get_drawings() or []
                n_drawings = len(drawings)
            except Exception:
                drawings = []
                n_drawings = -1
            try:
                imgs = page.get_images(full=True) or []
                n_images = len(imgs)
            except Exception:
                n_images = -1
            try:
                _widgets = list(page.widgets() or []) if hasattr(page, "widgets") else []
            except TypeError:
                _widgets = list(page.widgets or []) if hasattr(page, "widgets") else []
            except Exception:
                _widgets = []
            n_widgets = len(_widgets)
            print(f"   diag: text_blocks={n_blocks} drawings={n_drawings} images={n_images} widgets={n_widgets}")
            if n_widgets > 0:
                print("   diag: widget dump:")
                for w in _widgets:
                    ft = getattr(w, "field_type", None)
                    fts = (getattr(w, "field_type_string", "") or getattr(w, "field_type_str", "") or "")
                    fn  = getattr(w, "field_name", "")
                    rc  = getattr(w, "rect", None)
                    print(f"      • type={ft}/{fts!r} name={fn!r} rect={rc}")

        sections = find_sections_on_page(page, expected_sections_norm=expected, dry_run=dry_run)
        blocks = _text_blocks(page)
        y_tol = 18.0

        # ---------- 1) typed underscores ----------
        d = page.get_text("dict")
        lines_on_page: List[Dict[str, Any]] = []
        for b in d.get("blocks", []):
            for ln in b.get("lines", []):
                spans = ln.get("spans", []) or []
                texts = []
                x0 = None; x1 = None
                y_mids = []
                for sp in spans:
                    t = (sp.get("text") or "").replace("\xa0", " ").strip()
                    if not t:
                        continue
                    texts.append(t)
                    x0 = sp["bbox"][0] if x0 is None else min(x0, sp["bbox"][0])
                    x1 = sp["bbox"][2] if x1 is None else max(x1, sp["bbox"][2])
                    y_mids.append((sp["bbox"][1] + sp["bbox"][3]) / 2.0)
                if texts:
                    lines_on_page.append({
                        "text": " ".join(texts),
                        "x0": float(x0 or 0.0),
                        "x1": float(x1 or 0.0),
                        "y_mid": float(sum(y_mids) / len(y_mids) if y_mids else 0.0),
                    })

        for li, ln in enumerate(lines_on_page):
            txt = ln["text"]
            if not txt or not underline_re.search(txt):
                continue
            y_window = 2.0 * y_tol
            x_band_pad = 28.0
            cands = []
            for lj, other in enumerate(lines_on_page):
                if lj == li:
                    continue
                t2 = other["text"]
                if not t2 or underline_re.search(t2):
                    continue
                dy = abs(other["y_mid"] - ln["y_mid"])
                if dy > y_window:
                    continue
                start_x = ln["x0"]
                if other["x1"] >= start_x - x_band_pad:
                    cands.append((dy, max(0.0, abs(other["x0"] - start_x)), other))
            label_text = _pick_nearest_text(cands) if cands else ""
            if not label_text:
                near_blocks = [blk for blk in blocks
                               if blk["text"]
                               and abs(((blk["y0"] + blk["y1"]) / 2.0) - ln["y_mid"]) < (1.6 * y_tol)
                               and blk["x1"] <= ln["x0"] + 6]
                near_blocks.sort(key=lambda blk: ln["x0"] - blk["x1"])
                label_text = _label_from_block(near_blocks[0], ln["y_mid"]) if near_blocks else "Field"

            insert_x   = float(ln["x0"] + 6.0)
            baseline_y = float(ln["y_mid"])
            section_name, section_norm = nearest_section_name(sections, baseline_y)
            label_norm = alias_normal(_normalize(label_text))
            key = (pno, section_norm, label_norm)
            counters[key] = counters.get(key, 0) + 1
            idx = counters[key]

            tpl["fields"].append({
                "page": pno,
                "label": label_text,
                "label_short": label_text,
                "label_full": label_text,
                "label_norm": label_norm,
                "anchor_x": insert_x,
                "anchor_y": baseline_y,
                "line_box": [ln["x0"], ln["y_mid"], ln["x1"], ln["y_mid"]],
                "placement": "start",
                "section": section_name,
                "section_norm": section_norm,
                "index": idx,
            })
            if dry_run:
                print(f"   • field[{idx}] (typed underscores, line) → '{label_text}' @y≈{baseline_y:.1f} (sec: {section_name})")

        # ---------- 2) drawn underlines ----------
        segs = _line_like_segments(page)
        for j, seg in enumerate(segs):
            label_text = "Field"
            same_row = [blk for blk in blocks
                        if blk["text"]
                        and abs(((blk["y0"] + blk["y1"]) / 2.0) - seg["y0"]) < y_tol
                        and blk["x1"] <= seg["x0"] + 4]
            if same_row:
                same_row.sort(key=lambda blk: seg["x0"] - blk["x1"])
                label_text = _label_from_block(same_row[0], seg["y0"])
            else:
                # use horizontal overlap to catch centered headers
                above = [blk for blk in blocks
                         if blk["text"]
                         and (0 <= (seg["y0"] - blk["y1"]) < 2 * y_tol)
                         and _horiz_overlap(seg["x0"], seg["x1"], blk["x0"], blk["x1"])]
                if above:
                    above.sort(key=lambda blk: (seg["y0"] - blk["y1"], seg["x0"] - blk["x1"]))
                    label_text = _label_from_block(above[0], seg["y0"])
                else:
                    label_text = f"unknown_drawn_{pno}_{j}"

            # NEW: column-aware fallback for grid columns
            col_hdr = _column_header_by_spans(page, seg["y0"], (seg["x0"] + seg["x1"]) / 2.0, seg["x0"], seg["x1"])
            if col_hdr:
                label_text = col_hdr

            _label_use = (label_text or "Field").strip()
            insert_x = float(seg["x0"] + 6.0)
            y_mid    = float(seg["y0"])
            section_name, section_norm = nearest_section_name(sections, y_mid)

            line_w = float(seg["x1"] - seg["x0"])
            if line_w >= 90.0:
                underline_rect = fitz.Rect(float(seg["x0"]), float(seg["y0"] - 8),
                                           float(seg["x1"]), float(seg["y1"] + 8))
                if _right_zone_has_dropdown_affordance(page, underline_rect, arrow_zone_w=30.0):
                    label_norm = alias_normal(_normalize(_label_use))
                    key = (pno, section_norm, label_norm)
                    counters[key] = counters.get(key, 0) + 1
                    idx = counters[key]
                    tpl["fields"].append({
                        "page": pno,
                        "label": _label_use,
                        "label_short": _label_use,
                        "label_full": _label_use,
                        "label_norm": label_norm,
                        "anchor_x": float((seg["x0"] + seg["x1"]) / 2.0),
                        "anchor_y": y_mid,
                        "box_rect": [float(seg["x0"]), float(seg["y0"] - 8), float(seg["x1"]), float(seg["y1"] + 8)],
                        "line_box": [float(seg["x0"]), float(seg["y0"]), float(seg["x1"]), float(seg["y0"])],
                        "placement": "acro_choice",
                        "choices": [],
                        "section": section_name,
                        "section_norm": section_norm,
                        "index": idx,
                    })
                    if dry_run:
                        print(f"   • field[{idx}] (UNDERLINE→dropdown) → '{_label_use}' @y≈{y_mid:.1f} (sec: {section_name})")
                    continue

            label_norm = alias_normal(_normalize(_label_use))
            key = (pno, section_norm, label_norm)
            counters[key] = counters.get(key, 0) + 1
            idx = counters[key]
            tpl["fields"].append({
                "page": pno,
                "label": _label_use,
                "label_short": _label_use,
                "label_full": _label_use,
                "label_norm": label_norm,
                "anchor_x": insert_x,
                "anchor_y": y_mid,
                "line_box": [float(seg["x0"]), float(seg["y0"]), float(seg["x1"]), float(seg["y0"])],
                "placement": "center",
                "section": section_name,
                "section_norm": section_norm,
                "index": idx,
            })
            if dry_run:
                print(f"   • field[{idx}] (drawn line) → '{_label_use}' @y≈{y_mid:.1f} (sec: {section_name})")

        # ---------- 2.5) drawn dropdowns ----------
        dd_rects = _drawn_dropdowns(page)
        for j, r in enumerate(dd_rects, start=1):
            same_row = [blk for blk in blocks
                        if blk["text"]
                        and abs(((blk["y0"] + blk["y1"]) / 2.0) - r["cy"]) < y_tol
                        and blk["x1"] <= r["x0"] + 6]
            if same_row:
                same_row.sort(key=lambda blk: r["x0"] - blk["x1"])
                label_text = _label_from_block(same_row[0], r["cy"])
            else:
                # use horizontal overlap
                above = [blk for blk in blocks
                         if blk["text"] and (0 <= (r["cy"] - blk["y1"]) < 2 * y_tol)
                         and _horiz_overlap(r["x0"], r["x1"], blk["x0"], blk["x1"])]
                if above:
                    above.sort(key=lambda blk: (r["cy"] - blk["y1"], r["x0"] - blk["x1"]))
                    label_text = _label_from_block(above[0], r["cy"])
                else:
                    label_text = "Field"

            # NEW: column-aware fallback
            col_hdr = _column_header_by_spans(page, r["cy"], r["cx"], r["x0"], r["x1"])
            if col_hdr:
                label_text = col_hdr

            section_name, section_norm = nearest_section_name(sections, r["cy"])
            label_norm = alias_normal(_normalize(label_text))
            key = (pno, section_norm, label_norm)
            counters[key] = counters.get(key, 0) + 1
            idx = counters[key]
            tpl["fields"].append({
                "page": pno,
                "label": label_text,
                "label_short": label_text,
                "label_full": label_text,
                "label_norm": label_norm,
                "anchor_x": r["cx"],
                "anchor_y": r["cy"],
                "box_rect": [r["x0"], r["y0"], r["x1"], r["y1"]],
                "line_box": [r["x0"], r["cy"], r["x1"], r["cy"]],
                "placement": "acro_choice",
                "choices": [],
                "section": section_name,
                "section_norm": section_norm,
                "index": idx,
            })
            if dry_run:
                print(f"   • field[{idx}] (DRAWN dropdown) → '{label_text}' @y≈{r['cy']:.1f} (sec: {section_name})")

        # ---------- 2.6) fallback: long shallow rectangles ----------
        dl_rects = _dropdown_like_boxes(page)
        if dl_rects:
            existing_keys = set()
            for f in tpl["fields"]:
                if f.get("placement") == "acro_choice" and f.get("page") == pno and f.get("box_rect"):
                    x0, y0, x1, y1 = map(float, f["box_rect"])
                    existing_keys.add((round(x0,1), round(y0,1), round(x1,1), round(y1,1)))

            for j, r in enumerate(dl_rects, start=1):
                key4 = (round(r["x0"],1), round(r["y0"],1), round(r["x1"],1), round(r["y1"],1))
                if key4 in existing_keys:
                    continue
                same_row = [blk for blk in blocks
                            if blk["text"]
                            and abs(((blk["y0"] + blk["y1"]) / 2.0) - r["cy"]) < y_tol
                            and blk["x1"] <= r["x0"] + 6]
                if same_row:
                    same_row.sort(key=lambda blk: r["x0"] - blk["x1"])
                    label_text = _label_from_block(same_row[0], r["cy"])
                else:
                    # use horizontal overlap
                    above = [blk for blk in blocks
                             if blk["text"] and (0 <= (r["cy"] - blk["y1"]) < 2 * y_tol)
                             and _horiz_overlap(r["x0"], r["x1"], blk["x0"], blk["x1"])]
                    if above:
                        above.sort(key=lambda blk: (r["cy"] - blk["y1"], r["x0"] - blk["x1"]))
                        label_text = _label_from_block(above[0], r["cy"])
                    else:
                        label_text = "Field"

                # NEW: column-aware fallback
                col_hdr = _column_header_by_spans(page, r["cy"], r["cx"], r["x0"], r["x1"])
                if col_hdr:
                    label_text = col_hdr

                section_name, section_norm = nearest_section_name(sections, r["cy"])
                label_norm = alias_normal(_normalize(label_text))
                key = (pno, section_norm, label_norm)
                counters[key] = counters.get(key, 0) + 1
                idx = counters[key]
                tpl["fields"].append({
                    "page": pno,
                    "label": label_text,
                    "label_short": label_text,
                    "label_full": label_text,
                    "label_norm": label_norm,
                    "anchor_x": r["cx"],
                    "anchor_y": r["cy"],
                    "box_rect": [r["x0"], r["y0"], r["x1"], r["y1"]],
                    "line_box": [r["x0"], r["cy"], r["x1"], r["cy"]],
                    "placement": "acro_choice",
                    "choices": [],
                    "section": section_name,
                    "section_norm": section_norm,
                    "index": idx,
                })
                if dry_run:
                    print(f"   • field[{idx}] (FALLBACK dropdown) → '{label_text}' @y≈{r['cy']:.1f} (sec: {section_name})")

        # ---------- 2.7) colon/right-gap heuristic (DOCX-export fallback) ----------
        # If a line ends with a colon and there's a large empty horizontal region to its right,
        # synthesize a text field using that gap.
        try:
            page_w = float(page.rect.x1)
            left_margin  = 18.0
            right_margin = 24.0
            min_gap_px   = 48.0   # require at least this much whitespace to the right
            min_box_w    = 90.0   # only create a field if the box is at least this wide

            # Build quick index of existing line_boxes to avoid duplicates
            existing_line_boxes = []
            for f in tpl["fields"]:
                if f.get("page") == pno and f.get("line_box"):
                    x0, y0, x1, y1 = map(float, f["line_box"])
                    existing_line_boxes.append((x0, y0, x1, y1))

            def _too_close_to_existing(y_mid, x0, x1, y_slop=9.0, x_overlap_min=36.0):
                for ex0, ey, ex1, _ in existing_line_boxes:
                    if abs(ey - y_mid) <= y_slop:
                        ov = min(ex1, x1) - max(ex0, x0)
                        if ov >= x_overlap_min:
                            return True
                return False

            colon_like = []
            for blk in blocks:
                t = (blk.get("text") or "").strip()
                if not t:
                    continue
                # Only 1 line of text from the block (use the line near the block).
                lab = _label_from_block(blk, (blk["y0"] + blk["y1"]) / 2.0)
                if not lab:
                    continue
                # Must look like a label: ends with a colon (fullwidth etc. handled earlier)
                if not re.search(r"[:：﹕꞉˸፡︓]\s*$", lab):
                    continue
                # Skip very long headings; focus on short/medium labels
                if len(lab) > 120:
                    continue

                y_mid = (blk["y0"] + blk["y1"]) / 2.0

                # Find nearest right-hand neighbor on same row
                rightmates = [
                    (max(0.0, float(other["x0"]) - float(blk["x1"])), other)
                    for other in blocks
                    if other is not blk
                       and other.get("text")
                       and abs(((other["y0"] + other["y1"]) / 2.0) - y_mid) < y_tol
                       and float(other["x0"]) >= float(blk["x1"]) - 2.0
                ]
                rightmates.sort(key=lambda it: it[0])

                # Candidate field box across the right gap
                if rightmates and rightmates[0][0] >= min_gap_px:
                    x0f = float(blk["x1"]) + 6.0
                    x1f = float(rightmates[0][1]["x0"]) - 6.0
                else:
                    # No neighbor; use to page margin if there is enough space
                    gap = page_w - float(blk["x1"]) - right_margin
                    if gap < min_gap_px:
                        continue
                    x0f = float(blk["x1"]) + 6.0
                    x1f = page_w - right_margin

                if (x1f - x0f) < min_box_w:
                    continue
                if _too_close_to_existing(y_mid, x0f, x1f):
                    continue

                # All good: add a synthetic text field centered on this line
                section_name, section_norm = nearest_section_name(sections, y_mid)
                label_text = lab.rstrip(":：﹕꞉˸፡︓").rstrip()
                label_norm = alias_normal(_normalize(label_text))
                key = (pno, section_norm, label_norm)
                counters[key] = counters.get(key, 0) + 1
                idx = counters[key]

                tpl["fields"].append({
                    "page": pno,
                    "label": label_text,
                    "label_short": label_text,
                    "label_full": label_text,
                    "label_norm": label_norm,
                    "anchor_x": float(x0f + (x1f - x0f) * 0.5),
                    "anchor_y": float(y_mid),
                    "line_box": [float(x0f), float(y_mid), float(x1f), float(y_mid)],
                    "placement": "center",
                    "section": section_name,
                    "section_norm": section_norm,
                    "index": idx,
                })
                if dry_run:
                    print(f"   • field[{idx}] (colon-gap) → '{label_text}' @y≈{y_mid:.1f} (x≈{x0f:.1f}…{x1f:.1f}) (sec: {section_name})")

        except Exception as _e:
            if dry_run:
                print(f"   ⚠️ colon-gap heuristic skipped due to error: {_e}")

        # ---------- 2.8) word-level colon/right-gap heuristic (DOCX-export fallback v2) ----------
        # Similar to 2.7 but operates at the word level to catch cases where the whole line
        # is one text block or when block-splitting misses the label.
        try:
            from collections import defaultdict

            page_w = float(page.rect.x1)
            right_margin = 24.0
            min_gap_px   = 48.0    # require at least this much whitespace to the right
            min_box_w    = 90.0    # only create a field if the box is at least this wide

            # Build quick index of existing line_boxes to avoid duplicates
            existing_line_boxes = []
            for f in tpl["fields"]:
                if f.get("page") == pno and f.get("line_box"):
                    x0, y0, x1, y1 = map(float, f["line_box"])
                    existing_line_boxes.append((x0, y0, x1, y1))

            def _too_close_to_existing(y_mid, x0, x1, y_slop=9.0, x_overlap_min=36.0):
                for ex0, ey, ex1, _ in existing_line_boxes:
                    if abs(ey - y_mid) <= y_slop:
                        ov = min(ex1, x1) - max(ex0, x0)
                        if ov >= x_overlap_min:
                            return True
                return False

            words = page.get_text("words") or []
            # words format: [x0, y0, x1, y1, "text", block_no, line_no, word_no]
            words.sort(key=lambda w: (round(float(w[1]), 1), round(float(w[0]), 1)))

            # group by (block_no, line_no)
            groups = defaultdict(list)
            for x0, y0, x1, y1, wd, bno, lno, wno in words:
                groups[(int(bno), int(lno))].append((float(x0), float(y0), float(x1), float(y1), str(wd)))

            for (_blk, _ln), ws in groups.items():
                ws.sort(key=lambda t: (t[0], t[1]))
                # Rebuild line text to check for colon
                line_text = " ".join(w for *_, w in ws).strip()
                if not line_text or len(line_text) > 180:
                    continue

                # Must contain at least one colon-like character
                if not re.search(r"[:：﹕꞉˸፡︓]\s*$", line_text):
                    continue

                # y-mid for the line
                y_mid = (min(t[1] for t in ws) + max(t[3] for t in ws)) / 2.0

                # Find the x1 of the last token that ends with a colon
                label_end_x = None
                for x0, y0, x1, y1, wtxt in ws:
                    if re.search(r"[:：﹕꞉˸፡︓]\s*$", wtxt):
                        label_end_x = float(x1)
                if label_end_x is None:
                    continue

                # Find first word to the right (if any)
                next_right_x0 = None
                for x0, y0, x1, y1, wtxt in ws:
                    if float(x0) > label_end_x:
                        next_right_x0 = float(x0)
                        break

                # Build candidate box based on gap
                if next_right_x0 is not None:
                    gap = next_right_x0 - label_end_x
                    if gap < min_gap_px:
                        continue
                    x0f = label_end_x + 6.0
                    x1f = next_right_x0 - 6.0
                else:
                    gap = page_w - right_margin - label_end_x
                    if gap < min_gap_px:
                        continue
                    x0f = label_end_x + 6.0
                    x1f = page_w - right_margin

                if (x1f - x0f) < min_box_w:
                    continue
                if _too_close_to_existing(y_mid, x0f, x1f):
                    continue

                # Label = everything up to the last colon, with colon stripped
                # (handle fullwidth/alt colons too)
                label_text = re.sub(r"[:：﹕꞉˸፡︓]\s*$", "", line_text).rstrip()
                if not label_text or len(label_text) > 160:
                    continue

                section_name, section_norm = nearest_section_name(sections, y_mid)
                label_norm = alias_normal(_normalize(label_text))
                key = (pno, section_norm, label_norm)
                counters[key] = counters.get(key, 0) + 1
                idx = counters[key]

                tpl["fields"].append({
                    "page": pno,
                    "label": label_text,
                    "label_short": label_text,
                    "label_full": label_text,
                    "label_norm": label_norm,
                    "anchor_x": float(x0f + (x1f - x0f) * 0.5),
                    "anchor_y": float(y_mid),
                    "line_box": [float(x0f), float(y_mid), float(x1f), float(y_mid)],
                    "placement": "center",
                    "section": section_name,
                    "section_norm": section_norm,
                    "index": idx,
                })
                if dry_run:
                    print(f"   • field[{idx}] (word-colon-gap) → '{label_text}' @y≈{y_mid:.1f} (x≈{x0f:.1f}…{x1f:.1f}) (sec: {section_name})")

        except Exception as _e:
            if dry_run:
                print(f"   ⚠️ word-level colon-gap heuristic skipped due to error: {_e}")

        # ---------- 2.9) DOCX generic label→right-gap fallback (no colon required) ----------
        # Build fields when a line looks like a short label with a large empty area to the right.
        # This helps for Word forms without underlines / widgets / trailing colons.
        try:
            from collections import defaultdict

            # quick index to avoid duplicates with existing line_boxes
            existing_line_boxes = []
            for f in tpl["fields"]:
                if f.get("page") == pno and f.get("line_box"):
                    x0, y0, x1, y1 = map(float, f["line_box"])
                    existing_line_boxes.append((x0, y0, x1, y1))

            def _too_close_to_existing(y_mid, x0, x1, y_slop=9.0, x_overlap_min=36.0):
                for ex0, ey, ex1, _ in existing_line_boxes:
                    if abs(ey - y_mid) <= y_slop:
                        ov = min(ex1, x1) - max(ex0, x0)
                        if ov >= x_overlap_min:
                            return True
                return False

            page_w = float(page.rect.x1)
            right_margin = 24.0
            min_gap_px   = 70.0     # require at least this much whitespace on the right
            min_box_w    = 110.0    # only create a field if the “fillable” box is wide enough

            # Collect words and group by (block_no, line_no)
            words = page.get_text("words") or []
            words.sort(key=lambda w: (round(float(w[1]), 1), round(float(w[0]), 1)))
            groups = defaultdict(list)
            for x0, y0, x1, y1, wd, bno, lno, wno in words:
                groups[(int(bno), int(lno))].append((float(x0), float(y0), float(x1), float(y1), str(wd)))

            # Basic section header filter using existing logic
            def _looks_like_header(t: str) -> bool:
                return _is_section_header_relaxed(t)

            # light label heuristic: short-ish, not shouty, not boilerplate
            STOPWORDS = {"application", "form", "investor", "registration", "signature"}
            def _labelish(t: str) -> bool:
                t0 = unicodedata.normalize("NFKC", t).strip()
                if not t0:
                    return False
                if len(t0) > 48:           # keep “label” short
                    return False
                low = t0.lower()
                if any(sw in low for sw in STOPWORDS):
                    return False
                # avoid ALL-CAPS headers
                letters = [c for c in t0 if c.isalpha()]
                if letters:
                    caps_ratio = sum(1 for c in letters if c.isupper()) / max(1, len(letters))
                    if caps_ratio >= 0.70:
                        return False
                # must contain at least one alpha and at most ~7 words
                if not any(c.isalpha() for c in t0):
                    return False
                if len(t0.split()) > 7:
                    return False
                return True

            for (_blk, _ln), ws in groups.items():
                ws.sort(key=lambda t: (t[0], t[1]))
                line_text = " ".join(w for *_, w in ws).strip()
                if not line_text:
                    continue
                if _looks_like_header(line_text):
                    continue

                # y mid of the line
                y_mid = (min(t[1] for t in ws) + max(t[3] for t in ws)) / 2.0

                # Rightmost word on the line (treat its x1 as label end)
                last_x1 = max(t[2] for t in ws)

                # Compute empty right-side region
                gap = page_w - right_margin - last_x1
                if gap < min_gap_px:
                    continue

                # Candidate label text (strip trailing colon-like characters if present)
                label_text = re.sub(r"[:：﹕꞉˸፡︓]\s*$", "", line_text).rstrip()
                if not _labelish(label_text):
                    continue

                x0f = float(last_x1 + 6.0)
                x1f = float(page_w - right_margin)
                if (x1f - x0f) < min_box_w:
                    continue
                if _too_close_to_existing(y_mid, x0f, x1f):
                    continue

                # avoid colliding with a text block immediately to the right on (nearly) same line
                # (if there is a word whose x0 is within the fill box, skip)
                collision = any((x0f - 4.0) <= t[0] <= (x1f + 4.0) and abs(((t[1]+t[3])/2.0) - y_mid) < 9.0 for t in ws)
                if collision:
                    continue

                section_name, section_norm = nearest_section_name(sections, y_mid)
                label_norm = alias_normal(_normalize(label_text))
                key = (pno, section_norm, label_norm)
                counters[key] = counters.get(key, 0) + 1
                idx = counters[key]

                tpl["fields"].append({
                    "page": pno,
                    "label": label_text,
                    "label_short": label_text,
                    "label_full": label_text,
                    "label_norm": label_norm,
                    "anchor_x": float(x0f + (x1f - x0f) * 0.5),
                    "anchor_y": float(y_mid),
                    "line_box": [float(x0f), float(y_mid), float(x1f), float(y_mid)],
                    "placement": "center",
                    "section": section_name,
                    "section_norm": section_norm,
                    "index": idx,
                })
                if dry_run:
                    print(f"   • field[{idx}] (docx-right-gap) → '{label_text}' @y≈{y_mid:.1f} (x≈{x0f:.1f}…{x1f:.1f}) (sec: {section_name})")
        except Exception as _e:
            if dry_run:
                print(f"   ⚠️ docx-right-gap heuristic skipped due to error: {_e}")


        # ---------- 3) checkboxes ----------
        vector_boxes = _square_checkboxes(page) if '_square_checkboxes' in globals() else []
        large_boxes  = _large_square_boxes(page)
        tiny_vector_boxes = _tiny_vector_squares(page)
        glyph_boxes  = _glyph_line_checkboxes(page.get_text("dict"))
        merged_boxes = []
        def _add_if_not_overlapping(bx):
            grect = (bx["x0"], bx["y0"], bx["x1"], bx["y1"])
            for mb in merged_boxes:
                if _rects_overlap((mb["x0"], mb["y0"], mb["x1"], mb["y1"]), grect, pad=1.6):
                    return
            merged_boxes.append(bx)
        for src in (vector_boxes, large_boxes, tiny_vector_boxes, glyph_boxes):
            for b in src:
                _add_if_not_overlapping(b)

        for k, bx in enumerate(merged_boxes, start=1):
            cx = bx.get("cx", (bx["x0"] + bx["x1"]) / 2.0)
            cy = bx.get("cy", (bx["y0"] + bx["y1"]) / 2.0)
            right_cands = [
                (abs(((blk["y0"] + blk["y1"]) / 2.0) - cy) + max(0.0, blk["x0"] - bx["x1"]), blk)
                for blk in blocks
                if blk["text"] and abs(((blk["y0"] + blk["y1"]) / 2.0) - cy) < y_tol and blk["x0"] >= bx["x1"] - 2
            ]
            if right_cands:
                bullet_text = _best_line_from_candidates(right_cands, cy)
            else:
                left_cands = [
                    (abs(((blk["y0"] + blk["y1"]) / 2.0) - cy) + max(0.0, bx["x0"] - blk["x1"]), blk)
                    for blk in blocks
                    if blk["text"] and abs(((blk["y0"] + blk["y1"]) / 2.0) - cy) < y_tol and blk["x1"] <= bx["x0"] + 2
                ]
                bullet_text = _best_line_from_candidates(left_cands, cy)

            section_name, section_norm = nearest_section_name(sections, cy)
            short_key  = _short_label(section_name, k)
            pretty     = f"{short_key}: {bullet_text}".strip(": ")
            label_norm = alias_normal(_normalize(pretty))
            key = (pno, section_norm, label_norm)
            counters[key] = counters.get(key, 0) + 1
            idx = counters[key]
            tpl["fields"].append({
                "page": pno,
                "label": pretty,
                "label_short": short_key,
                "label_full": bullet_text,
                "label_norm": label_norm,
                "anchor_x": cx,
                "anchor_y": cy,
                "box_rect": [bx["x0"], bx["y0"], bx["x1"], bx["y1"]],
                "line_box": [bx["x0"], bx["y0"], bx["x1"], bx["y1"]],
                "placement": "checkbox",
                "section": section_name,
                "section_norm": section_norm,
                "index": idx,
            })
            if dry_run:
                print(f"   • checkbox[{idx}] (drawn/glyph) → '{pretty}' @y≈{cy:.1f} (sec: {section_name})")

        # AcroForm widgets
        try:
            widgets = list(page.widgets() or [])
        except TypeError:
            widgets = list(page.widgets or [])
        except Exception:
            widgets = []
        cb_widgets, text_widgets, choice_widgets = [], [], []
        _CHK = getattr(fitz, "PDF_WIDGET_TYPE_CHECKBOX", 2)
        _RAD = getattr(fitz, "PDF_WIDGET_TYPE_RADIOBUTTON", 3)
        _TXT = getattr(fitz, "PDF_WIDGET_TYPE_TEXT", 1)
        _CBX = getattr(fitz, "PDF_WIDGET_TYPE_COMBOBOX", 7)
        _LBX = getattr(fitz, "PDF_WIDGET_TYPE_LISTBOX", 6)
        _CHO = getattr(fitz, "PDF_WIDGET_TYPE_CHOICE", 6)
        for w in widgets:
            ft = getattr(w, "field_type", None)
            ft_str = (getattr(w, "field_type_string", "") or getattr(w, "field_type_str", "") or "").lower()
            name_l = (getattr(w, "field_name", "") or "").lower()
            is_checkbox = ft in (_CHK, _RAD)
            is_text = ft == _TXT
            is_choice = ft in (_CBX, _LBX, _CHO)
            if not is_choice and any(k in ft_str for k in ("combo", "list", "choice")):
                is_choice = True
            if not is_choice and (hasattr(w, "choice_values") or hasattr(w, "options") or hasattr(w, "items")):
                is_choice = True
            if not is_choice and any(k in name_l for k in ("combo", "list", "select", "dropdown")):
                is_choice = True
            try:
                if is_checkbox:
                    cb_widgets.append(w)
                elif is_choice:
                    choice_widgets.append(w)
                elif is_text:
                    text_widgets.append(w)
                else:
                    text_widgets.append(w)
            except Exception:
                text_widgets.append(w)

        cb_widgets.sort(key=lambda ww: (round(ww.rect.y0, 2), round(ww.rect.x0, 2)))
        for m, w in enumerate(cb_widgets, start=1):
            r  = w.rect
            cx = (r.x0 + r.x1) / 2.0
            cy = (r.y0 + r.y1) / 2.0
            right_cands = [
                (abs(((blk["y0"] + blk["y1"]) / 2.0) - cy) + max(0.0, blk["x0"] - r.x1), blk)
                for blk in blocks
                if blk["text"] and abs(((blk["y0"] + blk["y1"]) / 2.0) - cy) < y_tol and blk["x0"] >= r.x1 - 2
            ]
            if right_cands:
                bullet_text = _best_line_from_candidates(right_cands, cy)
            else:
                left_cands = [
                    (abs(((blk["y0"] + blk["y1"]) / 2.0) - cy) + max(0.0, r.x0 - blk["x1"]), blk)
                    for blk in blocks
                    if blk["text"] and abs(((blk["y0"] + blk["y1"]) / 2.0) - cy) < y_tol and blk["x1"] <= r.x0 + 2
                ]
                bullet_text = _best_line_from_candidates(left_cands, cy)

            section_name, section_norm = nearest_section_name(sections, cy)
            short_key  = _short_label(section_name, m)
            pretty     = f"{short_key}: {bullet_text}".strip(": ")
            label_norm = alias_normal(_normalize(pretty))
            key = (pno, section_norm, label_norm)
            counters[key] = counters.get(key, 0) + 1
            idx = counters[key]
            tpl["fields"].append({
                "page": pno,
                "label": pretty,
                "label_short": short_key,
                "label_full": bullet_text,
                "label_norm": label_norm,
                "anchor_x": cx,
                "anchor_y": cy,
                "box_rect": [float(r.x0), float(r.y0), float(r.x1), float(r.y1)],
                "line_box": [float(r.x0), float(r.y0), float(r.x1), float(r.y1)],
                "placement": "checkbox",
                "section": section_name,
                "section_norm": section_norm,
                "index": idx,
            })
            if dry_run:
                print(f"   • checkbox[{idx}] (AcroForm) → '{pretty}' @y≈{cy:.1f} (sec: {section_name})")

        # ---------- 4.5) AcroForm TEXT widgets ----------
        def _rects_overlap_field_line(r):
            for fd in tpl["fields"]:
                if fd["page"] != pno:
                    continue
                if fd.get("placement") not in ("start", "center"):
                    continue
                lb = fd.get("line_box")
                if not lb:
                    continue
                if _rects_overlap((float(lb[0]), float(lb[1]) - 2, float(lb[2]), float(lb[3]) + 2),
                                  (float(r.x0), float(r.y0), float(r.x1), float(r.y1)), pad=0.0):
                    return True
            return False

        text_widgets.sort(key=lambda ww: (round(ww.rect.y0, 2), round(ww.rect.x0, 2)))
        for n, w in enumerate(text_widgets, start=1):
            r = w.rect
            if _rects_overlap_field_line(r):
                continue
            cx = (r.x0 + r.x1) / 2.0
            cy = (r.y0 + r.y1) / 2.0
            same_row = [
                (abs(((blk["y0"] + blk["y1"]) / 2.0) - cy) + max(0.0, r.x0 - blk["x1"]), blk)
                for blk in blocks
                if blk["text"] and abs(((blk["y0"] + blk["y1"]) / 2.0) - cy) < y_tol and blk["x1"] <= r.x0 + 6
            ]
            if same_row:
                label_text = _best_line_from_candidates(same_row, cy)
            else:
                # horizontal overlap
                above = [
                    ((cy - blk["y1"]), blk)
                    for blk in blocks
                    if blk["text"] and (0 <= (cy - blk["y1"]) < 2 * y_tol)
                       and _horiz_overlap(r.x0, r.x1, blk["x0"], blk["x1"])
                ]
                label_text = _best_line_from_candidates(above, cy) or (w.field_name or "Field").strip()

            # NEW: column-aware fallback
            col_hdr = _column_header_by_spans(page, cy, cx, float(r.x0), float(r.x1))
            if col_hdr:
                label_text = col_hdr

            section_name, section_norm = nearest_section_name(sections, cy)
            label_norm = alias_normal(_normalize(label_text))
            key = (pno, section_norm, label_norm)
            counters[key] = counters.get(key, 0) + 1
            idx = counters[key]
            tpl["fields"].append({
                "page": pno,
                "label": label_text,
                "label_short": label_text,
                "label_full": label_text,
                "label_norm": label_norm,
                "anchor_x": cx,
                "anchor_y": cy,
                "box_rect": [float(r.x0), float(r.y0), float(r.x1), float(r.y1)],
                "line_box": [float(r.x0), float(r.y0), float(r.x1), float(r.y1)],
                "placement": "acro_text",
                "section": section_name,
                "section_norm": section_norm,
                "index": idx,
            })
            if dry_run:
                print(f"   • field[{idx}] (AcroForm TEXT) → '{label_text}' @y≈{cy:.1f} (sec: {section_name})")

        # ---------- 4.6) AcroForm CHOICE widgets ----------
        choice_widgets.sort(key=lambda ww: (round(ww.rect.y0, 2), round(ww.rect.x0, 2)))
        for q, w in enumerate(choice_widgets, start=1):
            r  = w.rect
            cx = (r.x0 + r.x1) / 2.0
            cy = (r.y0 + r.y1) / 2.0
            same_row = [
                (abs(((blk["y0"] + blk["y1"]) / 2.0) - cy) + max(0.0, r.x0 - blk["x1"]), blk)
                for blk in blocks
                if blk["text"] and abs(((blk["y0"] + blk["y1"]) / 2.0) - cy) < y_tol and blk["x1"] <= r.x0 + 6
            ]
            if same_row:
                label_text = _best_line_from_candidates(same_row, cy)
            else:
                # horizontal overlap
                above = [
                    ((cy - blk["y1"]), blk)
                    for blk in blocks
                    if blk["text"] and (0 <= (cy - blk["y1"]) < 2 * y_tol)
                       and _horiz_overlap(r.x0, r.x1, blk["x0"], blk["x1"])
                ]
                label_text = _best_line_from_candidates(above, cy) or (w.field_name or "Field").strip()

            # NEW: column-aware fallback
            col_hdr = _column_header_by_spans(page, cy, cx, float(r.x0), float(r.x1))
            if col_hdr:
                label_text = col_hdr

            section_name, section_norm = nearest_section_name(sections, cy)
            label_norm = alias_normal(_normalize(label_text))
            key = (pno, section_norm, label_norm)
            counters[key] = counters.get(key, 0) + 1
            idx = counters[key]

            choices = []
            try:
                raw = getattr(w, "choice_values", None) or getattr(w, "options", None) or getattr(w, "items", None)
                if raw:
                    for opt in list(raw):
                        if isinstance(opt, (list, tuple)) and opt:
                            choices.append(str(opt[0]))
                        else:
                            choices.append(str(opt))
            except Exception:
                pass

            tpl["fields"].append({
                "page": pno,
                "label": label_text,
                "label_short": label_text,
                "label_full": label_text,
                "label_norm": label_norm,
                "anchor_x": cx,
                "anchor_y": cy,
                "box_rect": [float(r.x0), float(r.y0), float(r.x1), float(r.y1)],
                "line_box": [float(r.x0), float(r.y0), float(r.x1), float(r.y1)],
                "placement": "acro_choice",
                "choices": choices,
                "section": section_name,
                "section_norm": section_norm,
                "index": idx,
            })
            if dry_run:
                extra = f" (choices={len(choices)})" if choices else ""
                print(f"   • field[{idx}] (AcroForm CHOICE) → '{label_text}' @y≈{cy:.1f}{extra} (sec: {section_name})")

    with open(template_json, "w", encoding="utf-8") as f:

        # --- DOCX: promote header-like rows to sections (generic, no hard-coded names) ---
        def _norm_text(s: str) -> str:
            import unicodedata, re
            s = unicodedata.normalize("NFKC", str(s or ""))
            s = re.sub(r"\s+", " ", s).strip()
            return s

        def _norm_key(s: str) -> str:
            import unicodedata, re, string
            s = unicodedata.normalize("NFKC", str(s or "")).lower()
            s = re.sub(r"\s+", " ", s).strip()
            return s.translate(str.maketrans("", "", string.punctuation))

        def _order_key(fd):
            # stable visual order: page -> y -> x -> index
            p = int(fd.get("page", 1) or 1)
            y = float(fd.get("anchor_y", 0.0) or 0.0)
            x = float(fd.get("anchor_x", 0.0) or 0.0)
            i = int(fd.get("index", 10**9) or 10**9)
            return (p, y, x, i)

        def _looks_like_header_generic(f, page_min_x, page_max_x) -> bool:
            """
            Generic header detector:
              • short, digit-free, punctuation-light text (≤ ~7 words)
              • title-ish capitalization (majority of words capitalized OR many caps)
              • visually spans most of the row and roughly centered
            """
            txt = _norm_text(f.get("label", ""))
            if not txt:
                return False

            # textual filters (no hard-coded names)
            if any(ch.isdigit() for ch in txt):
                return False
            words = [w for w in txt.split() if w]
            if len(words) == 0 or len(words) > 7:
                return False

            # avoid obvious field-ish endings (colon => treat as a header line though)
            ends_colon = txt.endswith(":")
            # if it ends with a colon, that's a very strong signal already
            strong_text_signal = ends_colon

            # capitalization heuristics
            caps_start_ratio = sum(1 for w in words if w[:1].isupper()) / max(1, len(words))
            letters = [c for c in txt if c.isalpha()]
            caps_ratio = (sum(1 for c in letters if c.isupper()) / max(1, len(letters))) if letters else 0.0
            titleish = (caps_start_ratio >= 0.55) or (caps_ratio >= 0.55)

            # geometric heuristics (use line_box if present, else anchor_x as a tiny box)
            x0 = x1 = None
            if f.get("line_box"):
                xb = f["line_box"]
                x0, x1 = float(xb[0]), float(xb[2])
            elif f.get("box_rect"):
                xb = f["box_rect"]
                x0, x1 = float(xb[0]), float(xb[2])
            else:
                # synthesize a small span around anchor_x so we can still compute center
                ax = float(f.get("anchor_x", 0.0) or 0.0)
                x0, x1 = ax - 1.0, ax + 1.0

            page_w = max(1.0, float(page_max_x - page_min_x))
            span_w = max(0.0, float(x1 - x0))
            span_frac = span_w / page_w

            # center-ish?
            center = (abs(((x0 + x1) / 2.0) - (page_min_x + page_w / 2.0)) <= 0.25 * page_w)

            # header-ish if it spans a good portion and is centered
            strong_geom_signal = (span_frac >= 0.55 and center)

            # combine signals
            return (titleish and strong_geom_signal) or (strong_text_signal and center)

        _is_docx_source = str(tpl.get("pdf", "")).lower().endswith((".docx", ".doc"))
        _is_docx_source = str(tpl.get("pdf", "")).lower().endswith((".docx", ".doc"))
        if _is_docx_source and tpl.get("fields"):
            # 1) sort fields in visual order
            fields_sorted = sorted(list(tpl["fields"]), key=_order_key)

            # 2) per-page horizontal bounds
            from collections import defaultdict
            bounds = defaultdict(lambda: [float("+inf"), float("-inf")])  # page -> [min_x, max_x]
            for f in fields_sorted:
                p = int(f.get("page", 1) or 1)
                if f.get("line_box"):
                    x0, x1 = float(f["line_box"][0]), float(f["line_box"][2])
                elif f.get("box_rect"):
                    x0, x1 = float(f["box_rect"][0]), float(f["box_rect"][2])
                else:
                    ax = float(f.get("anchor_x", 0.0) or 0.0)
                    x0, x1 = ax - 1.0, ax + 1.0
                bounds[p][0] = min(bounds[p][0], x0)
                bounds[p][1] = max(bounds[p][1], x1)

            # 3) scan: detect headers, KEEP them as rows, and propagate section to following fields
            current_sec_name = ""
            current_sec_norm = ""
            keep = []

            # separate counter for section headers so they don’t collide with normal fields
            section_counters = {}

            for fd in fields_sorted:
                p = int(fd.get("page", 1) or 1)
                min_x, max_x = bounds[p]

                if _looks_like_header_generic(fd, min_x, max_x):
                    # header text -> section
                    sec_name = _norm_text(fd.get("label", "")).rstrip(":").strip()
                    sec_norm = _norm_key(sec_name) if sec_name else ""

                    current_sec_name = sec_name
                    current_sec_norm = sec_norm

                    # KEEP header as its own row in the template
                    sh_key = (p, sec_norm)
                    section_counters[sh_key] = section_counters.get(sh_key, 0) + 1
                    header_idx = section_counters[sh_key]

                    header_row = dict(fd)  # shallow copy
                    header_row["section"] = sec_name
                    header_row["section_norm"] = sec_norm
                    header_row["label"] = sec_name  # clean label
                    header_row["label_short"] = sec_name
                    header_row["label_full"] = sec_name
                    header_row["label_norm"] = alias_normal(_normalize(sec_name))
                    header_row["placement"] = "section_header"
                    header_row["is_section"] = True
                    header_row["index"] = header_idx

                    keep.append(header_row)
                    continue

                # normal field: inherit last seen section if missing
                if not _norm_key(fd.get("section_norm", "")) and current_sec_norm:
                    fd["section"] = current_sec_name
                    fd["section_norm"] = current_sec_norm

                keep.append(fd)

            # 4) replace with full list (headers INCLUDED, sections propagated)
            tpl["fields"] = keep


        # ... your template post-processing code above ...
        # (Make sure any loops use e.g. `for fd in fields_sorted:` to avoid `f` lingering.)

        # --- write template JSON safely ---
        out_path = str(template_json)  # template_json is the path passed from field_map_generate
        import os, json
        os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)

        with open(out_path, "w", encoding="utf-8") as fh:   # <- use fh, not f
            json.dump(tpl, fh, indent=2, ensure_ascii=False)

        return tpl






# ---------------------------
# Resolver: choose best row by Page / Section / Index / Field
# ---------------------------
def resolve_value(rows: List[Dict[str, Any]],
                  field_label: str,
                  page: Optional[int],
                  section_norm: str,
                  occurrence_index: int,
                  min_field_fuzzy: float = 0.82,
                  return_row: bool = False,
                  strict_index: bool = True,
                  require_page_match: bool = False,       # NEW
                  require_section_match: bool = False):    # NEW
    """
    Select value for a given (Field label, Page, Section, occurrence_index).

    strict_index=True:
      - If ANY candidate rows for this Field specify an Index, require Index == occurrence_index.
      - If NO candidate rows have Index, ignore Index and select by Section/Page.

    require_page_match=True:
      - If 'page' is provided, only accept rows whose Page == page. If none exist, return None.

    require_section_match=True:
      - If 'section_norm' is provided, only accept rows whose section_norm == section_norm. If none exist, return None.
    """
    field_norm = alias_normal(_normalize(field_label))

    # 1) Candidate by exact field; else fuzzy
    field_candidates = [r for r in rows if r["field_norm"] == field_norm]
    if not field_candidates:
        all_fields = [r["field_norm"] for r in rows]
        bm, sc = _best_match_scored(field_norm, all_fields)
        if bm and sc >= min_field_fuzzy:
            field_candidates = [r for r in rows if r["field_norm"] == bm]
        else:
            return (None, None) if return_row else None

    candidates = list(field_candidates)

    # 2) Section handling
    if section_norm:
        sec_exact = [r for r in candidates if r.get("section_norm") == section_norm]
        if require_section_match:
            # hard require section match
            candidates = sec_exact
            if not candidates:
                return (None, None) if return_row else None
        elif sec_exact:
            # prefer (soft)
            candidates = sec_exact

    # 3) Page handling
    if page is not None:
        pg_exact = [r for r in candidates if r.get("Page") is not None and int(r["Page"]) == int(page)]
        if require_page_match:
            # hard require page match
            candidates = pg_exact
            if not candidates:
                return (None, None) if return_row else None
        elif pg_exact:
            # prefer (soft)
            candidates = pg_exact

    # 4) Index handling
    if strict_index:
        with_idx = [r for r in candidates if r.get("Index") is not None]
        if with_idx:
            exact_idx = [r for r in with_idx if int(r["Index"]) == int(occurrence_index)]
            if not exact_idx:
                return (None, None) if return_row else None
            candidates = exact_idx

    # 5) Tie-breaker
    def score_row(r: Dict[str, Any]) -> int:
        s = 0
        if r.get("Page") is not None and page is not None and int(r["Page"]) == int(page):
            s += 8
        if r.get("section_norm") and section_norm and r["section_norm"] == section_norm:
            s += 6
        if r.get("Index") is not None and int(r["Index"]) == int(occurrence_index):
            s += 3
        return s

    best = max(candidates, key=score_row, default=None)
    if return_row:
        return (best["Value"], best) if best else (None, None)
    return best["Value"] if best else None

# ---------------------------
# AcroForm fill (stable order + field-based Index + Yes/No)
# ---------------------------
def _get_widgets(page: fitz.Page):
    try:
        it = page.widgets() or []
    except TypeError:
        it = page.widgets or []
    return list(it)

# --- helper: collect explicit indices for a given field/page/(section) ---
def _explicit_indices_for(rows, field_norm: str, page: int, section_norm: str) -> List[int]:
    """Return sorted list of explicit Index values available for this (field, page, section?)."""
    cand = [r for r in rows if r.get("field_norm") == field_norm]
    if page is not None:
        cand = [r for r in cand if r.get("Page") is not None and int(r["Page"]) == int(page)]
    if section_norm:
        cand = [r for r in cand if r.get("section_norm") == section_norm]
    cand = [r for r in cand if r.get("Index") is not None]
    idxs = sorted({int(r["Index"]) for r in cand})
    return idxs

def fill_acroform_with_context(input_pdf: str,
                               output_pdf: str,
                               lookup_rows: List[Dict[str, Any]],
                               dry_run: bool = False):
    expected = expected_section_set(lookup_rows)
    doc = fitz.open(input_pdf)
    changed: List[Tuple[str, str, str, int, str, int]] = []
    widgets_exist = False

    for pno in range(len(doc)):
        page = doc[pno]
        if dry_run:
            print(f"p{pno+1}:")
        sections = find_sections_on_page(page, expected_sections_norm=expected, dry_run=dry_run)

        # --- AUGMENT: detect letter-only section headers like (A), A), A., A ---
        # This catches cases where (B)/(C) are printed as plain spans and missed by the relaxed header logic.
        try:
            d_for_sections = page.get_text("dict")
        except Exception:
            d_for_sections = None

        if d_for_sections:
            letter_pat = re.compile(r"^\s*\(?([A-Za-z])\)?\.?\s*$")
            letter_hits = []
            for b_ in d_for_sections.get("blocks", []) or []:
                for ln_ in b_.get("lines", []) or []:
                    for sp_ in (ln_.get("spans") or []):
                        t_ = (sp_.get("text") or "").strip()
                        if not t_ or len(t_) > 6:  # super short tokens only
                            continue
                        m_ = letter_pat.match(t_)
                        if not m_:
                            continue
                        ch = m_.group(1).upper()
                        # You can restrict the allowed letters here if desired:
                        # if ch not in {"A","B","C","D","E"}: continue
                        ymid_ = (float(sp_["bbox"][1]) + float(sp_["bbox"][3])) / 2.0
                        name_ = f"({ch})"
                        letter_hits.append({"name": name_, "name_norm": _normalize(name_), "y1": ymid_})

            if letter_hits:
                # merge with existing 'sections' preserving order; de-dup on (name_norm, y1~)
                merged = (sections or []) + letter_hits
                merged.sort(key=lambda c: (round(c["y1"], 1), c["name_norm"]))
                dedup = []
                seen_keys = set()
                for c in merged:
                    key_ = (c["name_norm"], round(c["y1"], 1))
                    if key_ in seen_keys:
                        continue
                    seen_keys.add(key_)
                    dedup.append(c)
                sections = dedup

                if dry_run and sections:
                    print("  sections (augmented):")
                    for s_ in sections:
                        print("   -", s_["name"])


        blocks = _text_blocks(page)

        widgets = list(page.widgets() or [])
        widgets.sort(key=lambda w: (round(w.rect.y0, 2), round(w.rect.x0, 2)))
        widgets_exist = widgets_exist or bool(widgets)

        # counters / bookkeeping
        occ_counters: Dict[Tuple[int, str, str], int] = {}
        used_indices: Dict[Tuple[int, str, str], Set[int]] = {}

        for w in widgets:
            try:
                name = (w.field_name or "").strip()
                # do NOT skip pre-filled – we may intentionally overwrite
                try:
                    cur_existing = w.field_value if w.field_value is not None else ""
                except Exception:
                    cur_existing = ""

                r    = w.rect
                midy = (r.y0 + r.y1) / 2.0

                # nearest-left label guess
                y_tol = 18.0
                lefts = [blk for blk in blocks
                         if blk["text"]
                         and abs(((blk["y0"] + blk["y1"]) / 2.0) - midy) < y_tol
                         and blk["x1"] <= r.x0 + 4]
                if lefts:
                    lefts.sort(key=lambda blk: r.x0 - blk["x1"])
                    label_guess = lefts[0]["text"]
                else:
                    above = [blk for blk in blocks
                             if blk["text"]
                             and (0 <= (midy - blk["y1"]) < 2 * y_tol)
                             and blk["x0"] <= r.x0]
                    if above:
                        above.sort(key=lambda blk: (midy - blk["y1"], r.x0 - blk["x1"]))
                        label_guess = above[0]["text"]
                    else:
                        label_guess = name or ""

                # section nearest to the widget
                section_name, detected_section_norm = nearest_section_name(sections, midy)

                # ---- Phase 1: discover FIELD bucket (ignore Index) ----
                _, picked_row = resolve_value(
                    lookup_rows,
                    label_guess or name,
                    page=pno + 1,
                    section_norm=detected_section_norm,
                    occurrence_index=1,
                    return_row=True,
                    strict_index=False,
                    )
                if not picked_row:
                    continue

                field_norm          = picked_row["field_norm"]
                excel_section_norm  = picked_row.get("section_norm") or ""
                effective_section   = excel_section_norm or detected_section_norm

                key = (pno + 1, effective_section, field_norm)

                # ---- Choose occurrence index: explicit first, else fallback counter ----
                explicit = _explicit_indices_for(lookup_rows, field_norm, pno + 1, effective_section)
                if explicit:
                    used = used_indices.setdefault(key, set())
                    next_idx = None
                    for idx_val in explicit:
                        if idx_val not in used:
                            next_idx = idx_val
                            break
                    if next_idx is None:
                        occ_counters[key] = occ_counters.get(key, 0) + 1
                        next_idx = occ_counters[key]
                    used.add(next_idx)
                    idx = next_idx
                else:
                    occ_counters[key] = occ_counters.get(key, 0) + 1
                    idx = occ_counters[key]

                # ---- Phase 2: resolve with progressive fallback ----
                value = resolve_value(
                    lookup_rows,
                    label_guess or name,
                    page=pno + 1,
                    section_norm=effective_section,
                    occurrence_index=idx,
                    strict_index=True,
                    )
                if value is None:
                    # 2a) ignore Index
                    value = resolve_value(
                        lookup_rows, label_guess or name,
                        page=pno + 1, section_norm=effective_section,
                        occurrence_index=idx, strict_index=False
                    )
                if value is None and effective_section:
                    # 2b) ignore Section too
                    value = resolve_value(
                        lookup_rows, label_guess or name,
                        page=pno + 1, section_norm="",  # relax section
                        occurrence_index=idx, strict_index=False
                    )
                if value is None:
                    # 2c) last resort: ignore Page & Section
                    value = resolve_value(
                        lookup_rows, label_guess or name,
                        page=None, section_norm="",
                        occurrence_index=idx, strict_index=False
                    )
                if value is None:
                    continue

                # ---- Yes/No handling (ERISA etc.) ----
                option = _nearby_yes_no_option(page, r.x0, (r.y0 + r.y1) / 2.0)
                if not option:
                    nm = (name or "").lower()
                    if "yes" in nm:
                        option = "yes"
                    elif "no" in nm:
                        option = "no"

                to_write = str(value)
                if option in {"yes", "no"}:
                    yn = _truthy(value)
                    if yn is None:
                        if dry_run:
                            print(f"[DRY] skip unclear yes/no for '{name}'")
                        continue
                    match = (yn and option == "yes") or ((yn is False) and option == "no")
                    if not match:
                        if dry_run:
                            print(f"[DRY] skip non-matching option '{name}' for value '{value}'")
                        continue

                    # Acrobat-proof visual X inside the widget — do not write the word "Yes"
                    if dry_run:
                        print(f"[DRY] X inside widget '{name or '(unnamed)'}' at "
                              f"({r.x0:.1f},{r.y0:.1f},{r.x1:.1f},{r.y1:.1f})")
                    else:
                        pad = max(0.8, 0.18 * min(r.width, r.height))
                        width = max(1.6, 0.22 * min(r.width, r.height))
                        p1 = fitz.Point(r.x0 + pad, r.y0 + pad)
                        p2 = fitz.Point(r.x1 - pad, r.y1 - pad)
                        p3 = fitz.Point(r.x0 + pad, r.y1 - pad)
                        p4 = fitz.Point(r.x1 - pad, r.y0 + pad)
                        page.draw_line(p1, p2, width=width, color=(0, 0, 0), overlay=True)
                        page.draw_line(p3, p4, width=width, color=(0, 0, 0), overlay=True)
                        changed.append((name or "(unnamed)", label_guess, "X", pno + 1, section_name, idx))
                    continue

                # Non yes/no flow: write value (allow overwrite unless identical)
                if dry_run:
                    print(f"[DRY] fill '{name or '(unnamed)'}' ← '{label_guess or name}' "
                          f"(sec='{section_name}' idx={idx}): '{to_write}'")
                else:
                    try:
                        cur_text = cur_existing
                        if isinstance(cur_text, bytes):
                            cur_text = cur_text.decode("utf-8", "ignore")
                    except Exception:
                        cur_text = ""
                    if str(cur_text) != to_write:
                        w.field_value = to_write
                        w.update()
                        changed.append((name or "(unnamed)", label_guess, to_write, pno + 1, section_name, idx))

            except Exception as e:
                print(f"⚠️ Widget update error on page {pno+1}: {e}")

        if widgets:
            try:
                page.apply_redactions()
            except Exception:
                pass

    if dry_run:
        print("Dry-run complete (AcroForm). No file written.")
        doc.close()
        return changed, widgets_exist

    doc.save(output_pdf, incremental=False)
    doc.close()
    if changed:
        print(f"📝 AcroForm write complete → {output_pdf}")
    else:
        print(f"ℹ️ AcroForm path wrote a copy → {output_pdf}")
    return changed, widgets_exist

def _why_skip(label: str, idx: int, reason: str, page_no: int, dry_run: bool):
    if dry_run:
        print(f"[DRY][SKIP] p{page_no} '{label}' (idx={idx}) -> {reason}")

def _find_any_widget_overlapping(
        page: fitz.Page,
        box: Tuple[float, float, float, float]
):
    """
    Return the first widget whose rect intersects `box`.
    Prefers checkbox / radio widgets if multiple overlap.
    """
    # Get widgets robustly across PyMuPDF versions
    try:
        widgets = list(page.widgets() or [])
    except TypeError:
        widgets = list(page.widgets or [])
    except Exception:
        widgets = []

    if not widgets:
        return None

    target = fitz.Rect(*box)

    # 1) Prefer checkbox / radio widgets
    for w in widgets:
        try:
            if getattr(w, "field_type", None) in (
                    getattr(fitz, "PDF_WIDGET_TYPE_CHECKBOX", None),
                    getattr(fitz, "PDF_WIDGET_TYPE_RADIOBUTTON", None),
            ):
                if fitz.Rect(w.rect).intersects(target):
                    return w
        except Exception:
            pass

    # 2) Otherwise return any intersecting widget
    for w in widgets:
        try:
            if fitz.Rect(w.rect).intersects(target):
                return w
        except Exception:
            pass

    return None


# ---------------------------
# Overlay fill (template) with field-based Index + Yes/No + checkboxes
# ---------------------------

def fill_from_template(pdf_path: str,
                       template_json: str,
                       lookup_rows: List[Dict[str, Any]],
                       out_pdf: str,
                       center_on_line: bool = True,
                       font_size: float = 10.5,
                       min_field_fuzzy: float = 0.82,
                       dry_run: bool = False):
    import json
    with open(template_json, "r", encoding="utf-8") as f:
        tpl = json.load(f)

    if dry_run:
        from collections import Counter
        page_counts = Counter(fd.get("page") for fd in tpl.get("fields", []))
        total = len(tpl.get("fields", []))
        print(f"🧩 Template loaded: {total} fields")
        for p in sorted(page_counts):
            print(f"   • page {p}: {page_counts[p]} fields")

    doc = fitz.open(pdf_path)
    filled = 0

    # ---------- tiny helpers ----------
    def _rect_from_fdef(fdef):
        if fdef.get("box_rect"):
            x0, y0, x1, y1 = map(float, fdef["box_rect"])
        elif fdef.get("line_box"):
            x0, y0, x1, y1 = map(float, fdef["line_box"])
        else:
            x = float(fdef["anchor_x"]); y = float(fdef["anchor_y"])
            x0, y0, x1, y1 = x - 6, y - 6, x + 6, y + 6
        return fitz.Rect(x0, y0, x1, y1)

    def _find_any_widget_overlapping(page: fitz.Page, box: Tuple[float, float, float, float]):
        try:
            widgets = list(page.widgets() or [])
        except TypeError:
            widgets = list(page.widgets or [])
        if not widgets:
            return None
        target = fitz.Rect(*box)
        for w in widgets:
            try:
                if fitz.Rect(w.rect).intersects(target):
                    return w
            except Exception:
                pass
        return None

    def _ensure_checkbox_widget(page: fitz.Page,
                                rect: Tuple[float, float, float, float],
                                field_name: str,
                                tooltip: str = ""):
        try:
            widget_dict = {
                "type": fitz.PDF_WIDGET_TYPE_CHECKBOX,
                "rect": fitz.Rect(*rect),
                "field_name": field_name,
                "field_value": "Off",
                "tooltip": tooltip or field_name,
                "text_color": (0, 0, 0),
                "border_color": None,
                "fill_color": None,
                "readonly": False,
                "required": False,
                "rotate": 0,
            }
            w = page.add_widget(widget_dict)
            try:
                w.set_flags(fitz.ANNOT_PRINT)
            except Exception:
                pass
            return w
        except Exception:
            return None

    def _rect_key(x0, y0, x1, y1, grid: float = 2.0):
        return (int(round(x0 / grid)), int(round(y0 / grid)),
                int(round(x1 / grid)), int(round(y1 / grid)))

    def _draw_center_X(page: fitz.Page, rect: fitz.Rect):
        pad = max(0.8, 0.18 * min(rect.width, rect.height))
        width = max(1.6, 0.22 * min(rect.width, rect.height))
        p1 = fitz.Point(rect.x0 + pad, rect.y0 + pad)
        p2 = fitz.Point(rect.x1 - pad, rect.y1 - pad)
        p3 = fitz.Point(rect.x0 + pad, rect.y1 - pad)
        p4 = fitz.Point(rect.x1 - pad, rect.y0 + pad)
        page.draw_line(p1, p2, width=width, color=(0, 0, 0), overlay=True)
        page.draw_line(p3, p4, width=width, color=(0, 0, 0), overlay=True)

    # --- find vertical cell boundaries (DOCX table snap) ---
    def _cell_span_from_verticals(page: fitz.Page,
                                  y: float,
                                  x_pref: float,
                                  y_tol: float = 10.0,
                                  max_gap: float = 540.0,
                                  min_gap: float = 28.0):
        try:
            drawings = page.get_drawings() or []
        except Exception:
            return None
        verts = []
        for g in drawings:
            for it in g.get("items", []) or []:
                if not it or it[0] != "l" or len(it) < 3:
                    continue
                p0, p1 = it[1], it[2]
                try:
                    x0 = float(getattr(p0, "x", p0[0] if isinstance(p0, (list, tuple)) else 0.0))
                    y0 = float(getattr(p0, "y", p0[1] if isinstance(p0, (list, tuple)) else 0.0))
                    x1 = float(getattr(p1, "x", p1[0] if isinstance(p1, (list, tuple)) else 0.0))
                    y1 = float(getattr(p1, "y", p1[1] if isinstance(p1, (list, tuple)) else 0.0))
                except Exception:
                    continue
                dx, dy = (x1 - x0), (y1 - y0)
                if abs(dx) > 2.0:
                    continue
                ymin, ymax = sorted([y0, y1])
                if ymin - y_tol <= y <= ymax + y_tol:
                    verts.append(((x0 + x1) / 2.0, ymin, ymax))
        if not verts:
            return None
        verts.sort(key=lambda v: v[0])
        xs = [v[0] for v in verts]
        if x_pref <= xs[0] or x_pref >= xs[-1]:
            return None
        left_idx = right_idx = None
        for i in range(1, len(xs)):
            if xs[i-1] <= x_pref <= xs[i]:
                left_idx, right_idx = i-1, i
                break
        if left_idx is None:
            return None
        xL, xR = xs[left_idx], xs[right_idx]
        gap = xR - xL
        if gap < min_gap or gap > max_gap:
            return None
        return (xL + 2.0, xR - 2.0)

    # STRICT lookup just for checkboxes
    def _strict_checkbox_value(label_short: str, page_no: int, section_norm: str, occurrence_index: int):
        norm = alias_normal(_normalize(label_short))
        cands = [r for r in lookup_rows if r.get("field_norm") == norm]
        if not cands:
            return None
        if section_norm:
            sec = [r for r in cands if r.get("section_norm") == section_norm]
            if sec:
                cands = sec
        if page_no is not None:
            pg = [r for r in cands if r.get("Page") is not None and int(r["Page"]) == int(page_no)]
            if pg:
                cands = pg
        with_idx = [r for r in cands if r.get("Index") is not None]
        if with_idx:
            exact = [r for r in with_idx if int(r["Index"]) == int(occurrence_index)]
            if exact:
                cands = exact
            else:
                return None
        def _score(r):
            s = 0
            if r.get("section_norm") == section_norm and section_norm:
                s += 5
            if r.get("Page") is not None and int(r["Page"]) == int(page_no):
                s += 4
            if r.get("Index") is not None and int(r["Index"]) == int(occurrence_index):
                s += 2
            return s
        best = max(cands, key=_score, default=None)
        return best["Value"] if best else None

    # bookkeeping
    ticked_regions: Dict[int, Set[Tuple[int, int, int, int]]] = {}
    occ_counters: Dict[Tuple[int, str, str], int] = {}
    used_indices: Dict[Tuple[int, str, str], Set[int]] = {}
    written_once: Set[Tuple[int, str, str, int]] = set()

    for fdef in tpl.get("fields", []):
        label = fdef.get("label", "") or ""
        if not label or label.startswith("unknown_"):
            continue

        page_idx = max(0, int(fdef["page"]) - 1)
        if page_idx >= len(doc):
            if dry_run:
                print(f"[DRY] skip '{label}' – template refers to missing page {page_idx+1}")
            continue

        detected_section_norm = fdef.get("section_norm", "")
        placement = (fdef.get("placement") or "").lower()
        page = doc[page_idx]

        # ---- Phase 1: discover bucket (ignore Index) ----
        raw_label   = label
        label_short = fdef.get("label_short", "")
        if placement == "checkbox" and label_short:
            search_keys = [label_short]
            fuzzy_for_this = 0.999
        else:
            search_keys: List[str] = []
            if label_short:
                search_keys.append(label_short)
            if ":" in raw_label:
                search_keys.append(raw_label.split(":", 1)[0].strip())
            search_keys.append(raw_label)
            fuzzy_for_this = (0.68 if placement == "checkbox" else min_field_fuzzy)

        picked_row = None
        picked_key = None
        for key_try in search_keys:
            _, pr = resolve_value(
                lookup_rows, key_try,
                page=page_idx + 1, section_norm=detected_section_norm,
                occurrence_index=1, min_field_fuzzy=fuzzy_for_this,
                return_row=True, strict_index=False
            )
            if pr:
                picked_row = pr
                picked_key = key_try
                break
        if not picked_row:
            if dry_run:
                kind = "checkbox" if placement == "checkbox" else "field"
                print(f"[DRY] p{page_idx+1} {kind} '{label}' → no match for {search_keys}")
            continue

        field_norm         = picked_row["field_norm"]
        excel_section_norm = picked_row.get("section_norm") or ""
        excel_page         = picked_row.get("Page", None)
        effective_section  = excel_section_norm or detected_section_norm
        bucket_key         = (page_idx + 1, effective_section, field_norm)

        # Occurrence index (explicit first)
        explicit = _explicit_indices_for(lookup_rows, field_norm, page_idx + 1, effective_section)
        if explicit:
            used = used_indices.setdefault(bucket_key, set())
            next_idx = None
            for idx_val in explicit:
                if idx_val not in used:
                    next_idx = idx_val
                    break
            if next_idx is None:
                occ_counters[bucket_key] = occ_counters.get(bucket_key, 0) + 1
                next_idx = occ_counters[bucket_key]
            used.add(next_idx); idx = next_idx
        else:
            occ_counters[bucket_key] = occ_counters.get(bucket_key, 0) + 1
            idx = occ_counters[bucket_key]

        logical_key = (page_idx + 1, effective_section, field_norm, idx)
        if logical_key in written_once:
            if dry_run:
                print(f"[DRY] p{page_idx+1} '{label}' (idx={idx}) → skip (already written this occurrence)")
            continue

        # ---- Phase 2: resolve with ordered, scoped matching (no global fallback) ----
        # Prefer (page & section) → section-only → page-only.
        # We DO NOT fall back to "anywhere in document".
        excel_has_page    = (excel_page is not None)
        excel_has_section = bool(excel_section_norm)

        def _resolve(require_page: bool, require_section: bool,
                     strict_idx: bool = True, pg=None, sec=None):
            return resolve_value(
                lookup_rows, picked_key or label,
                page=(pg if pg is not None else page_idx + 1),
                section_norm=(sec if sec is not None else effective_section),
                occurrence_index=idx,
                min_field_fuzzy=fuzzy_for_this,
                strict_index=strict_idx,
                require_page_match=require_page,
                require_section_match=require_section
            )

        value = None
        if excel_has_page and excel_has_section:
            value = _resolve(True, True, strict_idx=True)
            if value is None:
                value = _resolve(False, True, strict_idx=True)   # section-only
            if value is None:
                value = _resolve(True, False, strict_idx=True)   # page-only
        elif excel_has_section and not excel_has_page:
            value = _resolve(False, True, strict_idx=True)       # section-only
            if value is None:
                value = _resolve(True, False, strict_idx=True)   # page-only
        elif excel_has_page and not excel_has_section:
            value = _resolve(True, False, strict_idx=True)       # page-only
            if value is None:
                value = _resolve(False, True, strict_idx=True)   # section-only
        else:
            value = _resolve(True, False, strict_idx=True)       # page-only
            if value is None:
                value = _resolve(False, True, strict_idx=True)   # section-only

        if value is None:
            if dry_run:
                print(f"[DRY] p{page_idx+1} '{label}' (idx={idx}) → no value (scoped search only)")
            continue

        # =========================
        # CHECKBOX
        # =========================
        if placement == "checkbox":
            yn = _truthy(value)
            if yn is not True and yn is not False:
                if dry_run:
                    print(f"[DRY] p{page_idx+1} checkbox '{label_short or label}' -> skip (unclear '{value}')")
                continue
            r = _rect_from_fdef(fdef)
            rk = _rect_key(r.x0, r.y0, r.x1, r.y1)
            pgset = ticked_regions.setdefault(page_idx, set())
            if rk in pgset:
                continue
            rcx = (r.x0 + r.x1) / 2.0
            rcy = (r.y0 + r.y1) / 2.0
            opt_here = _nearby_yes_no_option(page, rcx, rcy)
            if opt_here in {"yes", "no"}:
                if (yn and opt_here != "yes") or ((yn is False) and opt_here != "no"):
                    if dry_run:
                        print(f"[DRY] p{page_idx+1} checkbox '{label_short or label}' -> skip (token mismatch)")
                    continue
            if dry_run:
                print(f"[DRY] p{page_idx+1} checkbox '{label_short or label}' -> TICK @ ({rcx:.1f},{rcy:.1f})")
                pgset.add(rk); written_once.add(logical_key)
                continue
            w = _find_any_widget_overlapping(page, (r.x0, r.y0, r.x1, r.y1))
            if w is not None:
                try:
                    if getattr(w, "field_value", "") != "Yes" and yn is True:
                        w.field_value = "Yes"; w.update()
                    filled += 1; pgset.add(rk); written_once.add(logical_key)
                    continue
                except Exception:
                    pass
            if yn is True:
                _draw_center_X(page, r)
                filled += 1; pgset.add(rk); written_once.add(logical_key)
            continue

        # =========================
        # NORMAL TEXT
        # =========================
        x = float(fdef["anchor_x"])
        y = float(fdef["anchor_y"])
        # default span
        if fdef.get("line_box"):
            x0_lb, y0_lb, x1_lb, y1_lb = map(float, fdef["line_box"])
            ux0 = max(x0_lb, x); ux1 = x1_lb
            base_y0, base_y1 = y0_lb, y1_lb
        elif fdef.get("box_rect"):
            x0_lb, y0_lb, x1_lb, y1_lb = map(float, fdef["box_rect"])
            ux0, ux1 = x0_lb, x1_lb
            base_y0, base_y1 = y0_lb, y1_lb
        else:
            ux0, ux1 = x - 6, x + 6
            base_y0, base_y1 = y - 6, y + 6

        # DOCX table snap
        snap = _cell_span_from_verticals(page, (base_y0 + base_y1) / 2.0, x_pref=x)
        if snap:
            cell_x0, cell_x1 = snap
            if (cell_x1 - cell_x0) >= 20 and not (cell_x1 < ux0 or cell_x0 > ux1):
                ux0 = max(ux0, cell_x0); ux1 = min(ux1, cell_x1)

        max_span = 520.0
        if (ux1 - ux0) > max_span:
            mid = (ux0 + ux1) / 2.0
            ux0 = mid - (max_span / 2.0)
            ux1 = mid + (max_span / 2.0)

        approx_char_w = 4.8
        text_w = max(1.0, len(str(value)) * approx_char_w)

        if center_on_line:
            draw_x = ux0 + max(0.0, (ux1 - ux0 - text_w)) / 2.0
            draw_y = (base_y0 + base_y1) / 2.0
        else:
            draw_x, draw_y = max(ux0, x), (base_y0 + base_y1) / 2.0

        if dry_run:
            print(f"[DRY] p{page_idx+1} '{label}' (idx={idx}) → '{value}' at ({draw_x:.1f},{draw_y:.1f}) [span {ux0:.1f}-{ux1:.1f}]")
            written_once.add(logical_key)
        else:
            rect = fitz.Rect(draw_x, draw_y - font_size, draw_x + 1200, draw_y + 2 * font_size)
            page.insert_textbox(rect, str(value), fontsize=font_size, align=fitz.TEXT_ALIGN_LEFT)
            filled += 1
            written_once.add(logical_key)

    if dry_run:
        print("Dry-run complete (overlay). No file written.")
        doc.close()
        return filled

    try:
        doc.set_need_appearances(True)
    except Exception:
        pass

    doc.save(out_pdf)
    doc.close()
    print(f"🎉 Overlay filled PDF saved to {out_pdf} (values placed: {filled})")
    return filled





# ---------------------------
# Coordinates exporter (debug)
# ---------------------------
def export_pdf_coordinates(pdf_path: str, csv_path: str = "pdf_coordinates.csv"):
    with _as_pdf(pdf_path) as _pdf_path:
        doc = fitz.open(_pdf_path)
        rows = []
        for page_num, page in enumerate(doc, start=1):
            for block in page.get_text("blocks"):
                x, y, x1, y1, text, *_ = block
                rows.append([page_num, int(x), int(y), int(x1), int(y1), (text or "").strip()])
        with open(csv_path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(["Page", "X0", "Y0", "X1", "Y1", "Text"])
            writer.writerows(rows)
        print(f"📄 Export complete: {csv_path}")
        doc.close()

def prefill_any(input_path: str,
                output_path: str,
                lookup_path: str,
                template_json: str = "template_fields.json",
                build_template_if_missing: bool = True,
                dry_run: bool = False,
                rebuild_template: bool = False):
    """
    Dispatch to DOCX or PDF path based on input/output extensions.
    """
    in_ext = os.path.splitext(input_path)[1].lower()
    out_ext = os.path.splitext(output_path)[1].lower()

    lookup_rows = read_lookup_rows(lookup_path)

    if in_ext in (".docx", ".doc") and out_ext in (".docx", ".doc"):
        # Native DOCX
        return prefill_docx(input_path, output_path, lookup_rows, dry_run=dry_run)

    # Fallback to your existing PDF flow (unchanged)
    return prefill_pdf(
        input_pdf=input_path,
        output_pdf=output_path if out_ext == ".pdf" else "out.pdf",
        lookup_path=lookup_path,
        template_json=template_json,
        build_template_if_missing=build_template_if_missing,
        dry_run=dry_run,
        rebuild_template=rebuild_template,
    )


# ---------------------------
# Orchestrator
# ---------------------------
def prefill_pdf(input_pdf: str,
                output_pdf: str,
                lookup_path: str,
                template_json: str = "template_fields.json",
                build_template_if_missing: bool = True,
                dry_run: bool = False,
                rebuild_template: bool = False):
    lookup_rows = read_lookup_rows(lookup_path)

    # Convert .doc/.docx to PDF if needed, then run existing logic unchanged.
    with _as_pdf(input_pdf) as _pdf_path:

        # 1) AcroForms
        print("🔎 Inspecting AcroForm widgets…")
        changed, widgets_exist = fill_acroform_with_context(
            input_pdf=_pdf_path,
            output_pdf=output_pdf,
            lookup_rows=lookup_rows,
            dry_run=dry_run
        )

        # In DRY mode we *also* run overlay for visibility.
        if not dry_run and widgets_exist and changed:
            print("✅ AcroForm path succeeded; overlay will ALSO run to draw vector checks.")
            # no return: we still run the overlay so vector checkboxes (like p4) get drawn
        else:
            if dry_run:
                if widgets_exist and changed:
                    print("ℹ️ DRY-RUN: AcroForm had matches; overlay will also run for reporting…")
                else:
                    print("ℹ️ DRY-RUN: No AcroForm changes; proceeding to overlay…")
            else:
                print("ℹ️ AcroForm did not fill everything; proceeding to overlay…")

        # 2) Overlay (template)
        print("ℹ️ Proceeding with underline/vector overlay template…")
        if rebuild_template or (not os.path.exists(template_json) and build_template_if_missing):
            print("🧩 Building template…(forced)" if rebuild_template else "🧩 Building template…")
            build_pdf_template(_pdf_path, template_json, lookup_rows=lookup_rows, dry_run=dry_run)
        else:
            print(f"📄 Using existing template: {template_json}")

        fill_from_template(
            pdf_path=_pdf_path,
            template_json=template_json,
            lookup_rows=lookup_rows,
            out_pdf=output_pdf,
            center_on_line=True,
            font_size=10.5,
            min_field_fuzzy=0.82,
            dry_run=dry_run
        )


# -------------------------------
# DOCX parsing helpers (no-PDF)
# -------------------------------
def _is_word_path(path: str) -> bool:
    p = str(path).lower()
    return p.endswith(".docx") or p.endswith(".doc")

def _require_python_docx():
    try:
        import docx  # noqa
    except Exception:
        raise RuntimeError(
            "DOCX parsing requires the 'python-docx' package. "
            "Install it with: pip install python-docx"
        )

def _docx_iter_pages(doc):
    """
    Yield (page_no, block) across the document.
    A 'block' is either ('para', paragraph) or ('table', table).
    We increment 'page_no' on explicit page breaks and on section breaks.
    """
    from docx.enum.text import WD_BREAK

    page_no = 1

    # Helper to scan page breaks in paragraph runs
    def _para_has_page_break(p):
        for r in p.runs:
            br_elems = r._r.xpath("./w:br")
            for br in br_elems:
                # if w:br has @w:type="page"
                t = br.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type")
                if t == "page":
                    return True
            # some producers mark it as an explicit break on the run
            if r._r.xpath("./w:lastRenderedPageBreak"):
                return True
        return False

    # If the document has sections, a new section can imply a new page (depending on type).
    def _section_new_page(sec):
        # If there is explicit type "nextPage", "oddPage", "evenPage" -> count as new page.
        # python-docx doesn't expose the break type directly, so we detect sectPr child @w:type.
        sp = sec._sectPr
        if sp is None:
            return False
        type_el = sp.xpath("./w:type")
        if not type_el:
            return False
        tval = type_el[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
        return tval in ("nextPage", "oddPage", "evenPage")

    # Walk the document in *source order* (paragraphs and tables interleaved)
    # python-docx doesn’t give us a direct interleaved iterator, so we iterate the body XML.
    body = doc._element.body
    for child in body.iterchildren():
        tag = child.tag
        if tag.endswith("}p"):
            p = docx.text.paragraph.Paragraph(child, doc)
            # detect section end before emitting next block? (Word pages are rendered,
            # but we approximate: first yield, then if there's a page break -> increment)
            yield (page_no, ("para", p))
            if _para_has_page_break(p):
                page_no += 1
        elif tag.endswith("}tbl"):
            t = docx.table.Table(child, doc)
            yield (page_no, ("table", t))
        elif tag.endswith("}sectPr"):
            # section properties outside paragraphs (rare), count as a new page if type says so
            # and bump page number for *next* content
            sp = child
            type_el = sp.xpath("./w:type")
            if type_el:
                tval = type_el[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
                if tval in ("nextPage", "oddPage", "evenPage"):
                    page_no += 1
        else:
            # other elements ignored
            pass

def _docx_guess_sections(line_text: str) -> Optional[str]:
    """
    Heuristic: treat all-caps lines, title-cased headings, or lines ending with 'FORM'
    etc. as a section header; return normalized section or None.
    """
    s = (line_text or "").strip()
    if not s:
        return None
    # Simple headers: ALL CAPS (≥ 2 words) or Ends with FORM / APPLICATION
    import re
    caps_words = sum(1 for w in s.split() if w.isupper())
    if caps_words >= 2 or re.search(r"(FORM|APPLICATION|SIGNATURE|REGISTRATION)\b", s.upper()):
        return s
    # Bold-only detection is not reliable in python-docx (style varies).
    return None

_CHECKBOX_TOKENS = ("☑", "☒", "☐", "■", "[ ]", "( )")

def _docx_detect_checkbox(text: str) -> bool:
    t = (text or "")
    if any(tok in t for tok in _CHECKBOX_TOKENS):
        return True
    # also allow 'Yes / No' patterns bracketed:
    if "Yes" in t and "No" in t and any(ch in t for ch in "[]()"):
        return True
    return False

def _docx_field_candidates_from_para(p):
    """
    Very light heuristic:
      - If a line has a colon or looks like 'Label .....' (with tab/underline),
        we treat the left part as a field label (split on colon).
      - Otherwise, if it matches common keys we've seen (e.g., 'Street 1', 'First name', etc.) we keep as is.
    """
    txt = p.text.strip()
    if not txt:
        return []

    cands = []
    # Split on colon if present (Label: __________)
    if ":" in txt:
        left = txt.split(":", 1)[0].strip()
        if left:
            cands.append(left)
    else:
        # If line is like "Field ......" with tabs or underscores, keep left token
        if "\t" in txt:
            cands.append(txt.split("\t", 1)[0].strip())
        else:
            # common keys (customize or allow all single short lines)
            tokens = txt.split()
            if len(tokens) <= 6:  # keep short headings as possible labels
                cands.append(txt)

    # Remove over-generic strings
    cands = [c for c in cands if c and len(c) >= 2]
    return list(dict.fromkeys(cands))  # unique, keep order

def _docx_field_rows_from_table(tbl):
    """
    Extract label-ish content from the first column (or header row) of tables.
    Returns a list of strings (candidate field labels).
    """
    rows = []
    try:
        for r in tbl.rows:
            try:
                if len(r.cells) >= 1:
                    raw = r.cells[0].text.strip()
                    if raw:
                        rows.append(raw)
            except Exception:
                continue
    except Exception:
        pass
    return rows

def _normalize_and_shorten(label: str) -> Tuple[str, str]:
    # Reuse normalize helpers already present in your file if available.
    # Fallback small normalizer here to avoid breaking PDF code.
    def _basic_norm(s):
        import re, unicodedata
        s = unicodedata.normalize("NFKD", s)
        s = re.sub(r"[^\w\s\-\/()]", "", s, flags=re.UNICODE).strip().lower()
        s = re.sub(r"\s+", " ", s)
        return s
    lab = label.strip()
    norm = _basic_norm(lab)
    short = lab.split(":", 1)[0].strip()
    return norm, short if short else lab

def parse_docx_structure(docx_path: str,
                         lookup_rows: Optional[List[Dict[str, Any]]] = None,
                         dry_run: bool = True) -> List[Dict[str, Any]]:
    """
    Parse a DOCX and produce a list of dicts with:
      { 'page': int, 'section': str, 'section_norm': str,
        'label': str, 'label_short': str, 'placement': 'text'|'checkbox',
        'Index': int, 'Value': Optional[str] (resolved if lookup_rows provided) }

    We do not compute coordinates (PDF-only feature).
    """
    _require_python_docx()
    import docx
    document = docx.Document(docx_path)

    fields = []
    current_section = ""
    current_section_norm = ""
    page_no = 1
    bucket_occ = {}  # (page, section_norm, field_norm) -> next index

    # we iterate in source order with page numbers from explicit breaks
    for pg, (kind, obj) in _docx_iter_pages(document):
        page_no = pg

        if kind == "para":
            p = obj
            txt = (p.text or "").strip()
            if not txt:
                continue

            # Section detection (heuristic)
            sec_guess = _docx_guess_sections(txt)
            if sec_guess:
                current_section = sec_guess
                current_section_norm = _normalize(sec_guess) if ' _normalize' in globals() else sec_guess.strip().lower()
                continue

            # Checkboxes?
            is_cb = _docx_detect_checkbox(txt)

            # Field candidates from para
            cand_labels = _docx_field_candidates_from_para(p)
            for lab in cand_labels:
                field_norm, label_short = _normalize_and_shorten(lab)
                key = (page_no, current_section_norm, field_norm)
                bucket_occ[key] = bucket_occ.get(key, 0) + 1
                idx = bucket_occ[key]
                value = None
                if lookup_rows:
                    # use your existing resolve_value if present; otherwise do exact label match
                    try:
                        v = resolve_value(
                            lookup_rows, lab,
                            page=None,
                            section_norm=current_section_norm,
                            occurrence_index=idx,
                            min_field_fuzzy=0.82,
                            strict_index=True,
                            require_page_match=False,
                            require_section_match=True
                        )
                        value = v
                    except Exception:
                        # fallback: exact match on normalized label
                        ln = alias_normal(_normalize(lab)) if 'alias_normal' in globals() and '_normalize' in globals() else field_norm
                        for r in lookup_rows:
                            if (r.get("field_norm") == ln and
                                    (not r.get("Page") or int(r.get("Page")) == page_no) and
                                    (not r.get("section_norm") or r.get("section_norm") == current_section_norm) and
                                    (not r.get("Index") or int(r.get("Index")) == idx)):
                                value = r.get("Value")
                                break

                fields.append({
                    "page": page_no,
                    "section": current_section,
                    "section_norm": current_section_norm,
                    "label": lab,
                    "label_short": label_short,
                    "placement": "checkbox" if is_cb else "text",
                    "Index": idx,
                    "Value": value,
                })

        elif kind == "table":
            tbl = obj
            labels = _docx_field_rows_from_table(tbl)
            for lab in labels:
                if not lab.strip():
                    continue
                is_cb = _docx_detect_checkbox(lab)
                field_norm, label_short = _normalize_and_shorten(lab)
                key = (page_no, current_section_norm, field_norm)
                bucket_occ[key] = bucket_occ.get(key, 0) + 1
                idx = bucket_occ[key]
                value = None
                if lookup_rows:
                    try:
                        value = resolve_value(
                            lookup_rows, lab,
                            page=None,
                            section_norm=current_section_norm,
                            occurrence_index=idx,
                            min_field_fuzzy=0.82,
                            strict_index=True,
                            require_page_match=False,
                            require_section_match=True
                        )
                    except Exception:
                        ln = alias_normal(_normalize(lab)) if 'alias_normal' in globals() and '_normalize' in globals() else field_norm
                        for r in lookup_rows:
                            if (r.get("field_norm") == ln and
                                    (not r.get("Page") or int(r.get("Page")) == page_no) and
                                    (not r.get("section_norm") or r.get("section_norm") == current_section_norm) and
                                    (not r.get("Index") or int(r.get("Index")) == idx)):
                                value = r.get("Value")
                                break

                fields.append({
                    "page": page_no,
                    "section": current_section,
                    "section_norm": current_section_norm,
                    "label": lab,
                    "label_short": label_short,
                    "placement": "checkbox" if is_cb else "text",
                    "Index": idx,
                    "Value": value,
                })

    if dry_run:
        # Mimic your existing dry-run console output style
        from collections import defaultdict
        by_page = defaultdict(list)
        for f in fields:
            by_page[f["page"]].append(f)
        for p in sorted(by_page):
            print(f"p{p}:")
            secs = [f["section"] for f in by_page[p] if f["section"]]
            secs_unique = list(dict.fromkeys(secs))
            if secs_unique:
                print("  sections found:")
                for s in secs_unique:
                    print(f"   - {s}")
            for f in by_page[p]:
                nm = f["label"]
                print(f"   • {('checkbox' if f['placement']=='checkbox' else 'field')} "
                      f"→ '{nm}' (sec: {f['section'] or ''}) idx={f['Index']}")
        print("Dry-run complete (DOCX). No file written.")

    return fields

def run_docx_mode(input_docx: str,
                  lookup_rows: Optional[List[Dict[str, Any]]],
                  dry_run: bool = True,
                  export_json: Optional[str] = None):
    """
    Parse DOCX and print/list Section / Page / Field / Index / Value.
    Does NOT modify the DOCX, does NOT convert to PDF.
    Leaves all PDF logic untouched.
    """
    fields = parse_docx_structure(input_docx, lookup_rows=lookup_rows, dry_run=dry_run)

    # Optional: export a JSON (no coordinates) that mirrors your template shape enough for inspection
    if export_json:
        out = {"docx": True, "fields": []}
        for f in fields:
            out["fields"].append({
                "page": f["page"],
                "section": f["section"],
                "section_norm": f["section_norm"],
                "label": f["label"],
                "label_short": f["label_short"],
                "placement": f["placement"],
                "Index": f["Index"],
                # no coordinates for DOCX path
            })
        import json
        with open(export_json, "w", encoding="utf-8") as fp:
            json.dump(out, fp, ensure_ascii=False, indent=2)
        print(f"🧩 DOCX parse exported → {export_json}")

    return fields

def export_docx_json(input_path: str, out_json: str, lookup_path: str, dry_run: bool = False):
    """
    Build a template JSON for a DOC/DOCX (or PDF) by converting to PDF via _as_pdf
    and reusing the existing template builder. Works for both Word and PDF inputs.
    """
    rows = read_lookup_rows(lookup_path)
    with _as_pdf(input_path) as _pdf_path:
        build_pdf_template(_pdf_path, out_json, lookup_rows=rows, dry_run=dry_run)
    print(f"🧩 DOCX/Word template JSON exported to {out_json}")


# ---------------------------
# CLI
# ---------------------------
def main():
    ap = argparse.ArgumentParser(description="PDF/DOCX prefill")
    ap.add_argument("--input", required=True, help="Input PDF or DOCX path")
    ap.add_argument("--output", required=True, help="Output PDF or DOCX path")
    ap.add_argument("--lookup", default="lookup_table.xlsx", help="Excel/CSV with Field,Value[,Section,Page,Index]")
    ap.add_argument("--template", default="template_fields.json", help="Template JSON (PDF overlay only)")
    ap.add_argument("--dry-run", action="store_true", help="Print what would be filled; no write")
    ap.add_argument("--export-coords", action="store_true", help="(PDF) also export pdf_coordinates.csv (debug)")
    ap.add_argument("--rebuild-template", action="store_true", help="(PDF) force rebuild of the template JSON")
    # NEW: export a template JSON for DOCX/PDF without filling
    ap.add_argument("--export-docx-json", metavar="PATH",
                    help="Build template JSON from the input (doc/docx/pdf) and save to PATH; then exit")
    args = ap.parse_args()

    if args.export_coords:
        # Only meaningful for PDFs; if input is DOCX, this will convert if needed
        try:
            export_pdf_coordinates(args.input)
        except Exception as e:
            print(f"⚠️ export-coords skipped: {e}")

    prefill_any(
        input_path=args.input,
        output_path=args.output,
        lookup_path=args.lookup,
        template_json=args.template,
        build_template_if_missing=True,
        dry_run=args.dry_run,
        rebuild_template=args.rebuild_template,
    )

if __name__ == "__main__":
    main()

