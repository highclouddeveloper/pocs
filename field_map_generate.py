import os
import sys
import json
import argparse
import traceback
import inspect
import pandas as pd
import re

# --- Inline checkbox splitting helpers (generic, pattern-based) ---
_BOX_TOKEN_RE = re.compile(r'(\[\s*[xX✓]?\s*\]|[□☐☒])')
_INLINE_SEP_RE = re.compile(r'(?:\t+|\s{2,}|\s+/\s+|\s+\|\s+|\s+OR\s+)', re.IGNORECASE)
# Strict checkbox detection: only mark checkboxes when a real checkbox control is found in the DOCX.
STRICT_CHECKBOX_DETECTION = True

# ---------- PDF row emission helpers (new) ----------

_RECIP_SEC_RE = re.compile(r"\brecipient\s*(\d+)\b", re.IGNORECASE)

def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", str(s or "").strip())

def _short_label_like(s: str, max_words=10, max_len=140) -> bool:
    t = _norm(s)
    if not t or len(t) > max_len:
        return False
    # avoid giant sentences with punctuation
    if t.count(".") >= 2:
        return False
    return len(t.split()) <= max_words

def _token_len(s: str) -> int:
    return len(re.findall(r"[A-Za-z0-9]", s or ""))

def _infer_recipient_index(section: str, field: str, index: int) -> int:
    """
    Prefer index inferred from 'Recipient N' in Section or Field.
    """
    for src in (section, field):
        m = _RECIP_SEC_RE.search(src or "")
        if m:
            try:
                n = int(m.group(1))
                return max(1, min(4, n))
            except Exception:
                pass
    return index or 1

def _is_pdf_fillable_fdef(fdef: dict) -> tuple[bool, str]:
    """
    Return (should_emit, type_tag) where type_tag ∈ {'checkbox','underline','table','dropdown','text',''}
    Uses fdef['type'], fdef['field_type'], fdef['placement'], fdef['choices'].
    """
    placement = _norm(fdef.get("placement", "")).lower()
    kind      = _norm(fdef.get("type", "") or fdef.get("field_type", "")).lower()
    label     = _norm(fdef.get("label_short") or fdef.get("label") or fdef.get("label_full") or "")

    # Widgets
    if "check" in kind or "acro_check" in placement or "acro-checkbox" in placement:
        return True, "checkbox"
    if "choice" in kind or "combo" in kind or "list" in kind or "acro_choice" in placement:
        return True, "dropdown"
    if ("text" in kind and "check" not in kind) or "acro_text" in placement:
        return (_short_label_like(label), "text")

    # Inline/visual checkbox hints
    if fdef.get("choices"):
        return True, "checkbox"
    if re.search(r"(?:\[\s*\]|\[\s*[xX✓]\s*\]|[□☐☑☒])", label):
        return True, "checkbox"

    # Underlines / blanks / rules
    if any(k in placement for k in ("underline", "line_rule", "line", "blank", "rule")):
        return (_short_label_like(label), "underline")

    # Table/grids
    if any(k in placement for k in ("grid_header", "table_header", "grid", "table")):
        return (_short_label_like(label, max_words=6, max_len=60), "table")

    # Rect boxes near labels → treat as text if label is short-ish
    if any(k in placement for k in ("rect_box", "box", "square")) and _short_label_like(label, 8, 80):
        return True, "text"

    return False, ""


def _final_split_compound_checkbox_fields(df_in: pd.DataFrame) -> pd.DataFrame:
    """
    After all other expansions, split any single 'Field' that is really a
    compound checkbox line (e.g., 'Initial Subscription Additional Subscription')
    into separate rows, one per option, keeping Section/Page and setting Index=1.
    """
    additions = []
    drop_idx = []

    for i, r in df_in.iterrows():
        field_text = str(r.get("Field", "")).strip()
        chunks = _split_compound_checkboxish(field_text)
        if not chunks:
            continue

        drop_idx.append(i)
        for ch in chunks:
            additions.append({
                "Section": r.get("Section", ""),
                "Page": r.get("Page", ""),
                "Field": ch,
                "Index": 1,
                "Value": "",
                "Choices": (r.get("Choices", "") or "checkbox")
            })

    if not additions:
        return df_in

    df_out = df_in.drop(index=drop_idx).reset_index(drop=True)
    df_out = pd.concat([df_out, pd.DataFrame(additions, columns=df_out.columns)], ignore_index=True)
    return df_out


def _split_compound_checkboxish(label: str) -> list:
    """
    Split lines like 'Initial Subscription Additional Subscription' into
    ['Initial Subscription', 'Additional Subscription'].

    Defensive heuristics; now tolerant to multiple/nbsp spaces.
    """
    t = (label or "").strip()
    if not t:
        return []
    if ":" in t:
        return []
    if len(t) > 60:
        return []
    if re.search(r"\d", t):  # avoid account numbers etc.
        return []

    # normalize whitespace (collapses multiple spaces / NBSP to single space)
    t_norm = re.sub(r"\s+", " ", t.replace("\u00A0", " ")).strip()
    words = t_norm.split()
    if len(words) < 3:  # need at least two chunks of 1–4 words each
        return []

    def _is_titleish(w: str) -> bool:
        # Accept TitleCase (Bank, State) or ALLCAPS (USA)
        return (w[:1].isupper() and (w[1:].islower() or w[1:] == "")) or w.isupper()

    if not all(_is_titleish(w) for w in words):
        return []

    # Greedy grouping into 1–4-word chunks
    chunks, cur = [], []
    for w in words:
        cur.append(w)
        if 2 <= len(cur) <= 4:
            chunks.append(" ".join(cur))
            cur = []
    if cur:
        if chunks and len(cur) <= 2:
            chunks[-1] = f"{chunks[-1]} {' '.join(cur)}"
            cur = []
        else:
            return []

    # Need 2–4 chunks, and each chunk 1–4 words
    if len(chunks) < 2 or len(chunks) > 4:
        return []
    if not all(1 <= len(c.split()) <= 4 for c in chunks):
        return []

    # Relaxed guard: match after whitespace normalization OR ensure chunks appear in order
    combined = " ".join(chunks)
    if combined != t_norm:
        # in-order containment fallback
        pos, ok = 0, True
        key = t_norm.lower()
        for c in chunks:
            ck = c.lower()
            j = key.find(ck, pos)
            if j < 0:
                ok = False
                break
            pos = j + len(ck)
        if not ok:
            return []

    return chunks



def _looks_short_option(s: str, max_words: int = 5) -> bool:
    return len((s or "").split()) <= max_words


def _is_title_case_phrase(s: str) -> bool:
    # Accept phrases where most words start uppercase (Title-Case-ish).
    words = (s or "").split()
    if not words:
        return False
    caps = sum(1 for w in words if w[:1].isupper())
    return caps >= max(1, int(0.6 * len(words)))


def _extract_inline_checkbox_options(text: str):
    """
    Extract inline options from a single-line label such as:
      '□ Initial Subscription    □ Additional Subscription'
      '[ ] Initial Subscription   [ ] Additional Subscription'
      'Initial Subscription    Additional Subscription' (2+ spaces)
      'Initial Subscription OR Additional Subscription'
      'Initial Subscription Additional Subscription'  <-- (title-case split)
    Returns list[str] or None.
    """
    if not text:
        return None

    # 0) Normalize a cheap test string
    t_norm = re.sub(r'\s+', ' ', text).strip()

    # 1) Split on explicit checkbox tokens (□, ☐, ☒, [ ], [x], etc.)
    t = _BOX_TOKEN_RE.sub(' ::: ', text)
    t = re.sub(r'\s+', ' ', t).strip()
    parts = [p.strip(' -:') for p in t.split(':::') if p and p.strip()]
    parts = [p for p in parts if _looks_short_option(p)]
    if len(parts) >= 2 and len(set(p.lower() for p in parts)) >= 2:
        return parts

    # 2) Fallback: split on tabs, 2+ spaces, '/', '|', or ' OR '
    parts = [p.strip(' -:') for p in _INLINE_SEP_RE.split(text) if p and p.strip()]
    parts = [p for p in parts if _looks_short_option(p)]
    if 2 <= len(parts) <= 8 and len(set(p.lower() for p in parts)) >= 2:
        return parts

    # 3) Title-case chunk heuristic (handles "Initial Subscription Additional Subscription")
    if ':' not in t_norm:
        words = t_norm.split()
        if 2 <= len(words) <= 8 and _is_title_case_phrase(t_norm):
            # Equal halves first
            if len(words) % 2 == 0:
                mid = len(words) // 2
                left = " ".join(words[:mid]).strip(' -:')
                right = " ".join(words[mid:]).strip(' -:')
                if _looks_short_option(left) and _looks_short_option(right) and left.lower() != right.lower():
                    return [left, right]

            # Otherwise, split before a capitalized word
            for i in range(2, len(words) - 1):
                if words[i][0].isupper():
                    left = " ".join(words[:i]).strip(' -:')
                    right = " ".join(words[i:]).strip(' -:')
                    if (_looks_short_option(left) and _looks_short_option(right)
                            and _is_title_case_phrase(left) and _is_title_case_phrase(right)
                            and left.lower() != right.lower()):
                        return [left, right]

    return None


# -----------------------------------------------------------------------------
# Try to find build template no matter where you place this helper.
# -----------------------------------------------------------------------------
_build_pdf_template_external = None
_build_pdf_template_import_error = None
_build_pdf_template_import_tb = None

# Optional DOCX builder import
_docx_builder = None
try:
    # Most repos expose this as build_docs_template
    from docx_prefill import build_docs_template as _docx_builder  # type: ignore
except Exception:
    try:
        # Back-compat: some repos call it build_docx_template
        from docx_prefill import build_docx_template as _docx_builder  # type: ignore
    except Exception:
        _docx_builder = None

# Optional PDF builder import (legacy/default)
try:
    from pdf_prefill import build_pdf_template as _build_pdf_template_external  # type: ignore
except Exception as _e:
    _build_pdf_template_external = None
    _build_pdf_template_import_error = _e
    _build_pdf_template_import_tb = traceback.format_exc()


def _describe_builder(fn) -> str:
    try:
        mod = getattr(fn, "__module__", None)
        src = inspect.getsourcefile(fn) or inspect.getfile(fn)
        return f"{mod or '<unknown module>'} :: {src}"
    except Exception:
        return "<uninspectable builder>"


def _pick_builder_by_input(input_path: str, verbose: bool = False):
    """
    Decide which builder to call based on the input file extension.
      - *.docx -> prefer docx_prefill.build_docx_template if available
      - otherwise -> pdf_prefill.build_pdf_template
    Also allows a local `build_pdf_template` or `build_docx_template` to override via globals().
    """
    local_build_pdf = globals().get("build_pdf_template")
    local_build_docx = globals().get("build_docx_template")

    ext = (os.path.splitext(input_path)[1] or "").lower()

    if ext == ".docx":
        if callable(local_build_docx):
            if verbose:
                print(f"🔧 Using local DOCX builder: {_describe_builder(local_build_docx)}")
            return local_build_docx, "docx"
        if callable(_docx_builder):
            if verbose:
                print(f"🔧 Using docx builder: {_describe_builder(_docx_builder)}")
            return _docx_builder, "docx"
        if callable(local_build_pdf):
            if verbose:
                print(f"🔧 Using local PDF builder (fallback for DOCX): {_describe_builder(local_build_pdf)}")
            return local_build_pdf, "pdf"
        if callable(_build_pdf_template_external):
            if verbose:
                print(f"🔧 Using imported PDF builder (fallback for DOCX): {_describe_builder(_build_pdf_template_external)}")
            return _build_pdf_template_external, "pdf"

    if callable(local_build_pdf):
        if verbose:
            print(f"🔧 Using local PDF builder: {_describe_builder(local_build_pdf)}")
        return local_build_pdf, "pdf"
    if callable(_build_pdf_template_external):
        if verbose:
            print(f"🔧 Using imported PDF builder: {_describe_builder(_build_pdf_template_external)}")
        return _build_pdf_template_external, "pdf"

    detail = ""
    if _build_pdf_template_import_error is not None:
        sys.stderr.write("----- pdf_prefill import traceback -----\n")
        if _build_pdf_template_import_tb:
            sys.stderr.write(_build_pdf_template_import_tb + "\n")
        else:
            sys.stderr.write(repr(_build_pdf_template_import_error) + "\n")
        sys.stderr.write("----- end traceback -----\n")
        detail = (
            f"\n(Import error: {type(_build_pdf_template_import_error).__name__}: "
            f"{_build_pdf_template_import_error})"
        )
    raise RuntimeError(
        "No builder found. Define build_docx_template/build_pdf_template in this file "
        "OR ensure docx_prefill/pdf_prefill are importable." + detail
    )


def _ensure_parent_dir(path: str):
    d = os.path.dirname(os.path.abspath(path))
    if d and not os.path.exists(d):
        os.makedirs(d, exist_ok=True)


def _field_title_from_fdef(fdef: dict) -> str:
    """Prefer a concise key when available, otherwise fallback to label."""
    title = (fdef.get("label_short") or fdef.get("label") or "").strip()
    if ":" in title:
        left = title.split(":", 1)[0].strip()
        if len(left) >= 6:
            title = left
    return title or (fdef.get("label_full") or "").strip()

# --- NEW: helpers to detect inline checkbox groups (no hard-coding of names)
_INLINE_SPLIT_RE = re.compile(r"(?:\t+|\s{2,}|(?:\s+OR\s+))", re.IGNORECASE)
_BOX_PREFIX_RE = re.compile(r"^\s*(?:[□☐☒]\s*|\[\s?[xX✓]?\s?\]\s*)")


def _split_inline_checkbox_options(display_text: str):
    """
    If display_text looks like a single line containing multiple short options
    separated by tabs / 2+ spaces / ' OR ', split into options.
    Returns list[str] or None if not applicable.
    """
    if not display_text or not display_text.strip():
        return None
    t = _BOX_PREFIX_RE.sub("", display_text).strip()
    parts = [p.strip() for p in _INLINE_SPLIT_RE.split(t) if p and p.strip()]
    parts = [p for p in parts if _looks_short_option(p, max_words=4)]
    if 2 <= len(parts) <= 6 and len(set(p.lower() for p in parts)) >= 2:
        return parts
    return None

from collections import defaultdict
_RECIP_RE = re.compile(r"\brecipient\s*\d+\b", re.IGNORECASE)

# running item counter PER section (so each RECIPIENT N gets item 1..K)
_section_item_counter = defaultdict(int)


def export_lookup_template_from_json(
        template_json: str,
        out_path: str = "lookup_template.xlsx",
        docx_path: str = None
) -> str:
    """
    Produce an Excel (or CSV) with columns:
    Section | Page | Field | Index | Value | Choices

    If docx_path is a DOCX, we also scan for 'matrix' tables whose header row
    contains multiple field labels and there are multiple blank rows beneath.
    For each header field found in the template, we auto-append Index=2..N rows
    so you can enter values for each data row in the document.
    """
    if not os.path.exists(template_json):
        raise FileNotFoundError(f"Template JSON not found: {template_json}")

    import json
    import pandas as pd
    import re
    from collections import OrderedDict, defaultdict

    with open(template_json, "r", encoding="utf-8") as f:
        tpl = json.load(f)

    rows = []
    seq = 0  # used to make a stable unique fallback field name when needed

    # --- NEW (scoped to this function): helpers for recipients & inline options
    _RECIP_RE = re.compile(r"\brecipient\s*(\d+)\b", re.IGNORECASE)

    # rely on your global _INLINE_SPLIT_RE if present; else provide a conservative default
    _INLINE_SPLIT_DEFAULT = re.compile(r"(?:\t+|\s{2,}|\s+\bOR\b\s+|[;•/])", re.IGNORECASE)
    try:
        _INLINE_SPLIT_RE  # noqa: F401
    except NameError:    # fallback only if it's not defined elsewhere
        _INLINE_SPLIT_RE = _INLINE_SPLIT_DEFAULT  # type: ignore

    def _recipient_index_from(section: str, default_idx: int) -> int:
        m = _RECIP_RE.search(section or "")
        if not m:
            return default_idx
        try:
            return int(m.group(1))
        except Exception:
            return default_idx

    def _split_inline_options_text(s: str) -> list[str]:
        t = (s or "").strip()
        if not t:
            return []
        parts = [p.strip(" -:") for p in _INLINE_SPLIT_RE.split(t) if p.strip()]
        # keep only short-ish items (≤ 8 words; ≥ 3 token chars)
        def _tok_len(x: str) -> int:
            return len(re.findall(r"[A-Za-z0-9]", x))
        parts = [p for p in parts if 1 <= len(p.split()) <= 8 and _tok_len(p) >= 3]
        # require at least 2 distinct parts to treat it as a group
        seen, out = set(), []
        for p in parts:
            k = p.lower()
            if k not in seen:
                seen.add(k)
                out.append(p)
        if len(out) >= 2 and len(out) <= 10:
            return out
        return []

    for fdef in (tpl.get("fields") or []):
        seq += 1

        raw_label = (fdef.get("label") or "").strip()
        section = (fdef.get("section") or "").strip()

        page_val = fdef.get("page", None)
        try:
            page = int(page_val) if page_val not in (None, "") else ""
            if isinstance(page, int) and page <= 0:
                page = ""
        except Exception:
            page = ""

        try:
            index = int(fdef.get("index", 1) or 1)
        except Exception:
            index = 1

        title_from_def = _field_title_from_fdef(fdef).strip()
        looks_unknown = (raw_label.lower().startswith("unknown_") if raw_label else True)
        if not title_from_def or title_from_def.lower().startswith("unknown_") or looks_unknown:
            sec_display = section if section else "No Section"
            title_from_def = f"Field #{seq} (Sec: {sec_display}, Idx: {index})"

        # Existing behavior: only pre-fill the "Choices" cell for AcroForm choice fields.
        choices_cell = ""
        if (fdef.get("placement") or "").lower() == "acro_choice":
            ch = fdef.get("choices") or []
            if isinstance(ch, list) and ch:
                choices_cell = " | ".join(str(x) for x in ch)

        # ---- NEW: Expand checkbox-like choice groups into separate rows for recipients & general checkboxes
        # 1) If this is a real choice list (AcroForm), expand to one row per option.
        if (fdef.get("placement") or "").lower() == "acro_choice":
            choice_list = fdef.get("choices") or []
            if isinstance(choice_list, list) and choice_list:
                if _RECIP_RE.search(section or ""):
                    # RECIPIENT N → use actual option labels and Index=N
                    recip_idx = _recipient_index_from(section, index)
                    for opt in choice_list:
                        opt_label = str(opt).strip()
                        if not opt_label:
                            continue
                        rows.append({
                            "Section": section,
                            "Page": page,
                            "Field": opt_label,   # ← actual label, not "item k"
                            "Index": recip_idx,   # ← recipient number
                            "Value": "",
                            "Choices": "checkbox",
                        })
                else:
                    # Non-recipient: keep each option as its own checkbox row
                    for opt in choice_list:
                        opt_label = str(opt).strip()
                        if not opt_label:
                            continue
                        rows.append({
                            "Section": section,
                            "Page": page,
                            "Field": opt_label,
                            "Index": 1,
                            "Value": "",
                            "Choices": "checkbox",
                        })
                continue  # skip the default single-row append


        # 2) If no explicit choices, try to split inline compact options in the label text.
        if not choices_cell:
            split_source = (fdef.get("label_full") or fdef.get("label") or title_from_def or "").strip()
            if not split_source.endswith(":"):
                opts = _split_inline_options_text(split_source)
                if opts:
                    if _RECIP_RE.search(section or ""):
                        # RECIPIENT N → use actual option labels and Index=N
                        recip_idx = _recipient_index_from(section, index)
                        for opt in opts:
                            opt_label = str(opt).strip()
                            if not opt_label:
                                continue
                            rows.append({
                                "Section": section,
                                "Page": page,
                                "Field": opt_label,   # ← actual label, not "item k"
                                "Index": recip_idx,   # ← recipient number
                                "Value": "",
                                "Choices": "checkbox",
                            })
                    else:
                        for opt in opts:
                            opt_label = str(opt).strip()
                            if not opt_label:
                                continue
                            rows.append({
                                "Section": section,
                                "Page": page,
                                "Field": opt_label,
                                "Index": 1,
                                "Value": "",
                                "Choices": "checkbox",
                            })
                    continue  # expanded; skip default


        # Default: original single-row emission (text/underline/table/etc.)
        rows.append({
            "Section": section,
            "Page": page,
            "Field": title_from_def,
            "Index": index,
            "Value": "",
            "Choices": choices_cell,
        })

    # Build initial DF
    cols = ["Section", "Page", "Field", "Index", "Value", "Choices"]
    df = pd.DataFrame(rows, columns=cols).fillna("")

    # --- Cleanup: drop obvious non-fillable narrative lines on PDF pages ---
    # Keep checkbox rows and short labels; drop long sentences / page boilerplate.
    def _too_narrative(row) -> bool:
        if str(row.get("Choices","")).strip().lower() == "checkbox":
            return False  # keep checkboxes
        field = str(row.get("Field","")).strip()
        if not field:
            return True
        # very long or too many words → likely narrative
        if len(field) > 120 or len(field.split()) > 20:
            return True
        # boilerplate fragments that slipped in as "fields"
        bad_snippets = (
            "subscription agreement", "page ", "the investor may also contact",
            "copy documents for any particular recipient", "attach additional pages if needed"
        )
        f_low = field.lower()
        if any(s in f_low for s in bad_snippets):
            return True
        return False

    df = df[~df.apply(_too_narrative, axis=1)].reset_index(drop=True)


    # --- NEW: normalize Recipient N short items into numbered checkbox rows (safety net)
    # If a RECIPIENT section still has multiple short, label-like rows (e.g., “Performance Indications”),
    # collapse them into standardized:
    #   RECIPIENT N – item 1..K   (Index = N, Choices = checkbox)
    # --- NEW: normalize Recipient N short items into numbered checkbox rows (safety net)
    # If a RECIPIENT section still has multiple short, label-like rows (e.g., “Performance Indications”),
    # collapse them into standardized:
    #   RECIPIENT N – item 1..K   (Index = N, Choices = checkbox)
    def _is_short_option(label: str) -> bool:
        t = (label or "").strip()
        if not t or t.endswith(":"):
            return False
        # accept the long "Annual (audited)/quarterly (unaudited) financial statements" style explicitly
        if re.search(r"\bannual\s*\(audited\)\s*/\s*quarterly\s*\(unaudited\)\s*financial\s*statements\b",
                     t, flags=re.IGNORECASE):
            return True
        # be more permissive in general so 4th item is kept
        if len(t) > 140:
            return False
        if len(t.split()) > 16:
            return False
        lower = t.lower()
        # avoid obvious data-entry fields
        if lower in {"name", "e-mail", "email", "telephone", "facsimile", "fax", "address"}:
            return False
        return True


    new_rows = []
    drop_idx = []
    for (sec, pg), grp in df.groupby(["Section", "Page"], dropna=False):
        if not _RECIP_RE.search(sec or ""):
            continue
        recip_idx = _recipient_index_from(sec or "", 1)
        # collect short options in original order
        candidates = []
        for i, r in grp.iterrows():
            fld = str(r["Field"]).strip()
            if _is_short_option(fld):
                candidates.append((i, fld))
        if len(candidates) >= 2:
            for k, (irow, _label) in enumerate(candidates, start=1):
                drop_idx.append(irow)
                new_rows.append({
                    "Section": sec,
                    "Page": pg,
                    "Field": f"{sec} – item {k}",
                    "Index": recip_idx,
                    "Value": "",
                    "Choices": "checkbox",
                })

    if new_rows:
        df = df.drop(index=drop_idx).reset_index(drop=True)
        df = pd.concat([df, pd.DataFrame(new_rows, columns=df.columns)], ignore_index=True)

    # ---------- OPTIONAL: expand matrix tables by scanning the DOCX ----------
    def _clean_label(s: str) -> str:
        s = (s or "").strip()
        return re.sub(r"[:：﹕꞉˸]\s*$", "", s).strip()

    def _first_cell_text(cell) -> str:
        return " ".join(p.text for p in cell.paragraphs).strip()

    def _is_boxlike(cell) -> bool:
        txt = "".join(p.text or "" for p in cell.paragraphs).strip()
        if not txt:
            return True
        vis = re.sub(r"[\s_\-\u2014\.\u2002\u2003\u2007\u2009\u00A0]+", "", txt)
        return not re.search(r"[A-Za-z0-9]", vis)

    def _expand_matrix_rows_with_docx(df_in: pd.DataFrame, docx_file: str) -> pd.DataFrame:
        try:
            from docx import Document
        except Exception:
            return df_in

        if not (docx_file and os.path.exists(docx_file) and docx_file.lower().endswith(".docx")):
            return df_in

        doc = Document(docx_file)
        additions = []

        for tbl in doc.tables:
            if not tbl.rows:
                continue
            header = tbl.rows[0]
            ncols = len(header.cells)
            if ncols < 2 or len(tbl.rows) < 2:
                continue

            headers = [_clean_label(_first_cell_text(c)) for c in header.cells]
            nonempty = [h for h in headers if h]
            if len(nonempty) < 2:
                continue

            row1 = tbl.rows[1]
            box_cnt = sum(1 for c in row1.cells if _is_boxlike(c))
            if box_cnt < max(2, int(0.6 * ncols)):
                continue  # not a matrix grid

            total_rows = max(1, len(tbl.rows) - 1)  # data rows
            if total_rows <= 1:
                continue

            for h in range(ncols):
                field_label = headers[h]
                if not field_label:
                    continue

                base_mask = (
                        (df_in["Field"].str.strip().str.lower() == field_label.strip().lower())
                        & (df_in["Index"].astype(int) == 1)
                )
                if not base_mask.any():
                    continue

                base_row = df_in[base_mask].iloc[0]
                for idx in range(2, total_rows + 1):
                    exists = (
                            (df_in["Field"].str.strip().str.lower() == field_label.strip().lower())
                            & (df_in["Index"].astype(int) == idx)
                    ).any()
                    if exists:
                        continue
                    additions.append({
                        "Section": base_row["Section"],
                        "Page": base_row["Page"],
                        "Field": field_label,
                        "Index": idx,
                        "Value": "",
                        "Choices": base_row.get("Choices", ""),
                    })

        if additions:
            df_out = pd.concat([df_in, pd.DataFrame(additions, columns=df_in.columns)], ignore_index=True)
            return df_out
        return df_in

    # ---- Checkbox splitting that relies on true glyphs in the DOCX (unchanged) ----
    def _split_inline_checkboxes_with_docx(df_in: pd.DataFrame, docx_file: str) -> pd.DataFrame:
        """
        Split a single combined Field into one row per option when the DOCX shows:
          1) real checkbox glyphs (□, ☐, ☑, ☒, [ ], [x]) next to short labels, or
          2) a 'please tick one' block with (a)/(b)/(c) options (often joined by 'OR').

        If no existing DF row matches to split, we append rows (Choices='checkbox')
        so they appear in the lookup sheet and can be filled later.
        """
        try:
            from docx import Document
        except Exception:
            return df_in

        if not (docx_file and os.path.exists(docx_file) and docx_file.lower().endswith(".docx")):
            return df_in

        import re as _re
        import unicodedata as _ud
        import string as _st

        doc = Document(docx_file)

        def _norm(s: str) -> str:
            s = _ud.normalize("NFKC", str(s or "")).strip()
            s = _re.sub(r"\s+", " ", s)
            return s

        def _norm_key(s: str) -> str:
            t = _norm(s).lower()
            return t.translate(str.maketrans("", "", _st.punctuation))

        def _iter_all_paragraphs(d):
            for p in d.paragraphs:
                yield p
            for t in d.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p

        def _looks_short_label(s: str, max_words: int = 12, max_len: int = 120) -> bool:
            s = (s or "").strip(" -:")
            if not s or len(s) > max_len:
                return False
            words = s.split()
            return 1 <= len(words) <= max_words

        BOX_TOKEN = r"(?:\[\s*[xX✓]?\s*\]|[□☐☑☒])"
        LABEL = r"([A-Z][\w\.\-/&()' ]{0,40}?)(?<!:)"
        LABEL_BOX_RE = _re.compile(LABEL + r"\s*" + BOX_TOKEN)
        TICK_TRIGGER_RE = _re.compile(r"\bplease\s+tick\s+one\b", _re.IGNORECASE)
        ENUM_BUL_RE     = _re.compile(r"^\s*(?:\((?:[a-eA-E]|i{1,3}|iv|v)\)|[-–•])\s+(.*)$")
        OR_SPLIT_RE     = _re.compile(r"\s+\bOR\b\s+", _re.IGNORECASE)

        glyph_groups, ctx_groups = [], []
        paras = [p for p in _iter_all_paragraphs(doc)]

        for p in paras:
            txt = _norm(p.text)
            if not txt:
                continue
            labels_here = [m.group(1).strip(" -:") for m in LABEL_BOX_RE.finditer(txt)]
            labels_here = [lb for lb in labels_here if _looks_short_label(lb, max_words=8, max_len=60)]
            if len(labels_here) >= 2:
                glyph_groups.append(labels_here)

        for i, p in enumerate(paras):
            line = _norm(p.text)
            if not line:
                continue
            if TICK_TRIGGER_RE.search(line):
                opts: list[str] = []
                j = i + 1
                while j < len(paras) and len(opts) < 6:
                    t = _norm(paras[j].text)
                    if not t:
                        j += 1
                        continue
                    m = ENUM_BUL_RE.match(t)
                    if m:
                        label = m.group(1).strip()
                        label = _re.sub(r"\s*\(tick[^)]*\)\s*$", "", label, flags=_re.IGNORECASE).strip()
                        parts = [s.strip() for s in OR_SPLIT_RE.split(label) if s.strip()]
                        for part in parts:
                            if _looks_short_label(part, max_words=24, max_len=140):
                                opts.append(part)
                    else:
                        parts = [s.strip() for s in OR_SPLIT_RE.split(t) if s.strip()]
                        for part in parts:
                            if _looks_short_label(part, max_words=24, max_len=140):
                                opts.append(part)
                    if len(t) > 200 and opts:
                        break
                    j += 1

                uniq, seen = [], set()
                for o in opts:
                    k = _norm_key(o)
                    if k and k not in seen:
                        seen.add(k)
                        uniq.append(o)
                if len(uniq) >= 2:
                    ctx_groups.append(uniq[:6])

        df = df_in.copy()

        def _apply_split_by_labels(df0: pd.DataFrame, labels: list[str]) -> tuple[pd.DataFrame, bool]:
            combined = _norm_key(" ".join(labels))
            hit_idx = []
            for idx, r in df0.iterrows():
                f_raw = _norm(str(r["Field"]))
                f_key = _norm_key(f_raw)
                if f_key == combined:
                    hit_idx.append(idx)
                    continue
                pos, ok = 0, True
                for lb in labels:
                    lk = _norm_key(lb)
                    j = f_key.find(lk, pos)
                    if j < 0:
                        ok = False
                        break
                    pos = j + len(lk)
                if ok:
                    hit_idx.append(idx)

            if not hit_idx:
                return df0, False

            new_rows = []
            for idx in hit_idx:
                base = df0.loc[idx]
                for lb in labels:
                    new_rows.append({
                        "Section": base["Section"],
                        "Page": base["Page"],
                        "Field": lb,
                        "Index": 1,
                        "Value": "",
                        "Choices": "checkbox",
                    })
            out = df0.drop(index=hit_idx).reset_index(drop=True)
            out = pd.concat([out, pd.DataFrame(new_rows, columns=df0.columns)], ignore_index=True)
            return out, True

        did_any_split = False
        for g in glyph_groups:
            df, did = _apply_split_by_labels(df, g)
            did_any_split = did_any_split or did

        for g in ctx_groups:
            df, did = _apply_split_by_labels(df, g)
            did_any_split = did_any_split or did

            if not did:
                pending = []
                for lb in g:
                    exists = (df["Field"].str.strip().str.lower() == lb.strip().lower()).any()
                    if not exists:
                        pending.append({
                            "Section": "",
                            "Page": "",
                            "Field": lb,
                            "Index": 1,
                            "Value": "",
                            "Choices": "checkbox",
                        })
                if pending:
                    df = pd.concat([df, pd.DataFrame(pending, columns=df.columns)], ignore_index=True)

        return df

    def _mark_checkbox_rows(df_in: pd.DataFrame) -> pd.DataFrame:
        """
        When STRICT_CHECKBOX_DETECTION is True, do NOT guess checkboxes from text alone.
        Only rows produced by the real-checkbox splitter will carry Choices='checkbox'.
        """
        try:
            STRICT_CHECKBOX = bool(globals().get("STRICT_CHECKBOX_DETECTION", False))
        except Exception:
            STRICT_CHECKBOX = False

        if STRICT_CHECKBOX:
            return df_in

        df = df_in.copy()
        def _is_checkboxish(s: str) -> bool:
            t = (str(s or "").strip())
            if not t or ":" in t or len(t) > 60:
                return False
            words = t.split()
            if len(words) > 6:
                return False
            caps = sum(1 for w in words if w[:1].isalpha() and w[:1].upper() == w[:1])
            return caps >= max(1, int(0.6 * len(words)))
        mask = df["Field"].apply(_is_checkboxish)
        df.loc[mask, "Choices"] = df.loc[mask, "Choices"].replace({"": "checkbox"}, regex=False)
        return df

    # Apply DOCX-aware expansions/splits (if a DOCX path was provided)
    if docx_path:
        df = _expand_matrix_rows_with_docx(df, docx_path)
        df = _split_inline_checkboxes_with_docx(df, docx_path)
        df = _final_split_compound_checkbox_fields(df)
        df = _mark_checkbox_rows(df)

    # ---------- sort & write ----------
    def _sort_key_row(r):
        page_sort = (999999 if r["Page"] == "" else int(r["Page"]))
        return page_sort, str(r["Section"]).lower(), str(r["Field"]).lower(), int(r["Index"])

    df = df.sort_values(
        by=["Page", "Section", "Field", "Index"],
        key=lambda s: s.apply(
            (lambda x: (999999 if (isinstance(x, str) and x == "") else x)) if s.name == "Page" else (lambda x: x)
        )
    ).reset_index(drop=True)

    _ensure_parent_dir(out_path)
    if out_path.lower().endswith(".csv"):
        df.to_csv(out_path, index=False, encoding="utf-8-sig")
    else:
        try:
            df.to_excel(out_path, index=False)
        except Exception as e:
            fallback = os.path.splitext(out_path)[0] + ".csv"
            df.to_csv(fallback, index=False, encoding="utf-8-sig")
            print(f"⚠️  Could not write Excel ({e}). Wrote CSV instead: {fallback}")
            return fallback

    print(f"📤 Lookup template written to {out_path} with {len(df)} rows.")
    return out_path



def _call_builder_with_compat(builder, input_path: str, template_json: str):
    """
    Call a builder with broad signature compatibility:
    - build_XXX_template(path, template_json, lookup_rows=None, dry_run=False)
    - build_XXX_template(path, template_json=..., lookup_rows=None, dry_run=False)
    - build_XXX_template(doc_or_path=path, template_json=..., ...)
    """
    _ensure_parent_dir(template_json)
    try:
        return builder(input_path, template_json, lookup_rows=None, dry_run=False)
    except TypeError:
        pass
    try:
        return builder(input_path, template_json=template_json, lookup_rows=None, dry_run=False)
    except TypeError:
        pass
    try:
        return builder(doc_or_path=input_path, template_json=template_json, lookup_rows=None, dry_run=False)
    except TypeError:
        pass
    return builder(input_path, template_json)


def _load_template_field_count(template_json: str) -> int:
    try:
        with open(template_json, "r", encoding="utf-8") as f:
            tpl = json.load(f)
        return len(tpl.get("fields", []) or [])
    except Exception:
        return -1


def export_lookup_template(
        input_path: str,
        template_json: str = "template_fields.json",
        out_path: str = "lookup_template.xlsx",
        rebuild_template: bool = False,
        debug_import: bool = False
) -> str:
    """
    Ensures a template JSON exists (building it if needed), then exports the lookup sheet.
    Adds diagnostics so you can see which builder ran and how many fields were detected.
    """
    builder, kind = _pick_builder_by_input(input_path, verbose=True or debug_import)

    must_build = rebuild_template or not os.path.exists(template_json)
    if must_build:
        print(f"🧩 Building template → {template_json}")
        _call_builder_with_compat(builder, input_path, template_json)
        cnt = _load_template_field_count(template_json)
        print(f"🧩 Template saved to {template_json} with {cnt} fields.")
        if cnt == 0:
            print("⚠️  No fields were detected.")
            if kind == "pdf":
                print("   • If using PDF, confirm you imported the correct pdf_prefill.py.")
                print("     → Run:  python -c \"import pdf_prefill,inspect; print(pdf_prefill.__file__)\"")
                print("   • If it’s a dynamic/XFA PDF or lacks detectable lines/widgets, try flattening (Print to PDF).")
                print("   • Or run the builder in DRY mode to see detection logs.")
            else:
                print("   • For DOCX, verify the document has tables/boxes/underlines the builder can detect.")
    else:
        print(f"📄 Using existing template: {template_json} ({_load_template_field_count(template_json)} fields)")

    return export_lookup_template_from_json(
        template_json,
        out_path,
        docx_path=input_path if input_path.lower().endswith(".docx") else None
    )


# ---------------------------
# CLI
# ---------------------------
def main():
    ap = argparse.ArgumentParser(
        description="Export a blank lookup sheet (Section, Page, Field, Index, Value, Choices) from a template JSON"
    )
    ap.add_argument("--input", required=True, help="Input file (PDF or DOCX)")
    ap.add_argument("--template", default="template_fields.json",
                    help="Template JSON (rebuilt if missing or --rebuild-template)")
    ap.add_argument("--make-lookup", metavar="OUT.xlsx",
                    help="Path to write the blank lookup sheet (xlsx or csv).")
    ap.add_argument("--rebuild-template", action="store_true",
                    help="Force rebuild of the template JSON even if it already exists")
    ap.add_argument("--debug-import", action="store_true",
                    help="Print extra import info for the builder used.")
    args = ap.parse_args()

    if args.make_lookup:
        export_lookup_template(
            input_path=args.input,
            template_json=args.template,
            out_path=args.make_lookup,
            rebuild_template=args.rebuild_template,
            debug_import=args.debug_import,
        )
        return

    print("Nothing to do. Pass --make-lookup OUT.xlsx (or .csv) to export a blank lookup sheet.")


if __name__ == "__main__":
    main()
