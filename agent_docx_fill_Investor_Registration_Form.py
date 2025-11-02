#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Agentic DOCX filler:
- Understand input.docx structure (sections + fields)
- Read lookup_template.xlsx (Section, Field, Index, Value[, Page])
- Match & fill values into the correct place
- Save out.docx

Usage:
  python agent_docx_fill.py --input input.docx --lookup lookup_template.xlsx --output out.docx --dry-run
"""

import os, re, math, string, unicodedata, argparse, sys
from typing import List, Dict, Any, Optional, Tuple, Iterable, Set
from difflib import SequenceMatcher
from numbers import Number
import pandas as pd
import string  # (make sure these imports exist once at file top)
import re

def _norm_space(s: str) -> str:
    return re.sub(r"\s+", " ", (str(s or "")).strip())

def _strip_colon_like(s: str) -> str:
    return re.sub(r"[:：﹕꞉˸]\s*$", "", (s or "").strip())

def _cell_text(cell) -> str:
    return "\n".join(p.text for p in cell.paragraphs).strip()

UNDERLINE_RUN = re.compile(r"_+")
MOSTLY_SPACEY = re.compile(r"^[\s\u2000-\u200B\u00A0.\-–—]{3,}$")

def _blankish_cell(cell) -> bool:
    t2 = _cell_text(cell).strip()
    if not t2:
        return True
    return bool(UNDERLINE_RUN.search(t2) or MOSTLY_SPACEY.match(t2) or len(t2) <= 1)

def _short_header(txt: str) -> bool:
    t = _strip_colon_like(_norm_space(txt))
    return bool(t) and len(t) <= 60 and len(t.split()) <= 6

_GENERIC_STOPWORDS = {
    "the","a","an","and","or","of","for","to","in","on","by","with","at","as",
    "no","number","only","existing","shareholders","shareholder","details","form",
    "applicant","application","name","class","series","currency","cash","amount",
    "words","subscription","day","first","last","middle","initial","hid","id",
    "existing","shareholders","only"
}

def _sig_words(s: str) -> set[str]:
    """Lowercased significant words (>=3 letters) with generics removed."""
    import re
    words = re.findall(r"[A-Za-z]+", (s or "").lower())
    return {w for w in words if len(w) >= 3 and w not in _GENERIC_STOPWORDS}

def _has_meaningful_overlap(a: str, b: str) -> bool:
    aw = _sig_words(a)
    bw = _sig_words(b)
    # require at least one shared significant token
    return len(aw & bw) >= 1

def _split_two_labels_no_glyphs(text: str):
    """
    Fallback splitter when there's no checkbox glyph, e.g.:
      'Initial Subscription                      Additional Subscription'
    Returns ['Initial Subscription', 'Additional Subscription'] or None.
    Very conservative: only two short title-case chunks.
    """
    if not text:
        return None
    t = unicodedata.normalize("NFKC", text).strip()
    t = re.sub(r"\s+", " ", t)

    # 1) Try splitting by BIG gaps that usually separate options
    #    (we look in the *original* text for >=3 spaces to preserve where the visual gap was)
    parts_gap = [p.strip(" -:") for p in re.split(r"\s{3,}", text) if p.strip()]
    parts_gap = [p for p in parts_gap if 1 <= len(p.split()) <= 4]
    if len(parts_gap) == 2 and parts_gap[0].lower() != parts_gap[1].lower():
        # ensure both parts are reasonably title-case
        def _titleish(s):
            ws = s.split();
            return ws and sum(1 for w in ws if w[:1].isupper()) >= max(1, int(0.6*len(ws)))
        if _titleish(parts_gap[0]) and _titleish(parts_gap[1]):
            return parts_gap

    # 2) Title-case boundary fallback (no big gaps)
    words = t.split()
    if not (3 <= len(words) <= 8):
        return None

    def _is_titleish_word(w: str):
        return (w[:1].isupper() and (w[1:].islower() or w[1:] == "")) or w.isupper()

    if not all(_is_titleish_word(w) for w in words):
        return None

    # try equal halves first
    if len(words) % 2 == 0:
        mid = len(words) // 2
        left = " ".join(words[:mid]).strip(" -:")
        right = " ".join(words[mid:]).strip(" -:")
        if (1 <= len(left.split()) <= 4 and 1 <= len(right.split()) <= 4
                and left.lower() != right.lower()):
            return [left, right]

    # else find a sensible boundary at a capitalized word
    for i in range(2, len(words)-1):
        if words[i][0].isupper():
            left = " ".join(words[:i]).strip(" -:")
            right = " ".join(words[i:]).strip(" -:")
            if (1 <= len(left.split()) <= 4 and 1 <= len(right.split()) <= 4
                    and left.lower() != right.lower()):
                return [left, right]

    return None

# ---------------------------
# Utilities / shared
# ---------------------------
PUNCT = str.maketrans("", "", string.punctuation)

# --- Normalization helpers (strict vs loose) ---
PUNCT_NO_PARENS = str.maketrans("", "", (string.punctuation.replace("(", "").replace(")", "")))


def _normalize_strict(s: str) -> str:
    """
    Strict canonical form:
      - NFKC
      - collapse whitespace
      - lower-case
      - keep parentheses (so 'Applicant Name (Entity)' != 'Applicant Name')
      - strip trailing colon-like chars
    """
    t = unicodedata.normalize("NFKC", str(s or ""))
    t = re.sub(r"[:：﹕꞉˸፡︓]\s*$", "", t)
    t = re.sub(r"\s+", " ", t).strip().lower()
    return t

def _normalize_loose(s: str) -> str:
    """
    Loose canonical form for fuzzy fallback:
      - NFKC
      - drop parentheticals completely
      - remove punctuation (incl. parentheses after removal)
      - collapse whitespace & lower-case
    """
    t = unicodedata.normalize("NFKC", str(s or ""))
    t = re.sub(r"\(.*?\)", "", t)                   # <-- drop parentheses content
    t = re.sub(r"[:：﹕꞉˸፡︓]\s*$", "", t)
    t = re.sub(r"\s+", " ", t).strip().lower()
    t = t.translate(str.maketrans("", "", string.punctuation))
    return t

def _wordset_loose(s: str) -> set:
    """Word set used to guard fuzzy matches."""
    t = _normalize_loose(s)
    return {w for w in t.split() if len(w) >= 3}



def _sim(a: str, b: str) -> float:
    return SequenceMatcher(None, _normalize(a), _normalize(b)).ratio()

def to_text_value(v) -> str:
    try:
        import pandas as _pd
        if _pd.isna(v):
            return ""
    except Exception:
        pass
    if v is None:
        return ""
    if isinstance(v, Number):
        try:
            if isinstance(v, float) and not math.isfinite(v):
                return ""
        except Exception:
            pass
        try:
            if float(v).is_integer():
                return str(int(v))
        except Exception:
            pass
        s = f"{float(v):.12f}".rstrip("0").rstrip(".")
        return s or "0"
    s = str(v).strip()
    if not s:
        return ""
    m = re.fullmatch(r"([+-]?\d+)\.0+\b", s)
    if m:
        return m.group(1)
    m = re.fullmatch(r"([+-]?\d+\.\d*?[1-9])0+\b", s)
    if m:
        return m.group(1)
    return s

def _truthy(val: str) -> Optional[bool]:
    s = (str(val or "")).strip().lower()
    if s in {"y","yes","true","1","x","✓","check","checked"}:
        return True
    if s in {"n","no","false","0","uncheck","unchecked"}:
        return False
    return None

# ---------------------------
# Agents
# ---------------------------

class LookupAgent:
    """Loads and normalizes lookup rows from Excel/CSV."""
    def __init__(self, path: str):
        self.path = path
        self.rows: List[Dict[str, Any]] = []

    def load(self) -> List[Dict[str, Any]]:
        if not os.path.exists(self.path):
            raise FileNotFoundError(f"Lookup file not found: {self.path}")

        if self.path.lower().endswith((".xlsx", ".xls")):
            df = pd.read_excel(self.path)
        else:
            df = pd.read_csv(self.path)

        if {"Field", "Value"} - set(df.columns):
            raise ValueError("Lookup must have columns: Field, Value. Optional: Section, Page, Index")

        for col in ["Field", "Value", "Section"]:
            if col in df.columns:
                df[col] = df[col].astype(str).fillna("").map(lambda x: x.strip())

        if "Page" in df.columns:
            df["Page"] = pd.to_numeric(df["Page"], errors="coerce").astype("Int64")
        if "Index" in df.columns:
            idx_series = df["Index"].astype(str).str.replace(r"[^\d\-]+", "", regex=True).replace({"": None})
            df["Index"] = pd.to_numeric(idx_series, errors="coerce").astype("Int64")

        # drop empty Field/Value
        df = df[(df["Field"].astype(str).str.strip() != "") &
                (df["Value"].astype(str).str.strip() != "") &
                (df["Value"].astype(str).str.lower() != "nan")]

        rows: List[Dict[str, Any]] = []
        for _, r in df.iterrows():
            field = str(r.get("Field", "")).strip()
            value = to_text_value(r.get("Value", ""))
            section = str(r.get("Section", "")).strip() if "Section" in df.columns else ""
            page = int(r["Page"]) if ("Page" in df.columns and pd.notna(r["Page"])) else None
            index = int(r["Index"]) if ("Index" in df.columns and pd.notna(r["Index"])) else None
            rows.append({
                "Field": field,
                "Value": value,
                "Section": section,
                "Page": page,
                "Index": index,
                "field_norm_strict": _normalize_strict(field),
                "field_norm_loose": _normalize_loose(field),
                "section_norm_strict": _normalize_strict(section) if section else "",
                "section_norm_loose": _normalize_loose(section) if section else "",
                "field_words_loose": _wordset_loose(field),
            })

        self.rows = rows
        return rows


class DocumentUnderstandingAgent:
    """
    Scans the DOCX, finds:
      - Table label/value pairs (by columns 0,2,4,...)
      - Paragraph underline patterns ("Label: ______")
      - Section titles from table row-0 headers or "Section X:" paragraphs
    Produces a list of discovered fields with (section, label_short, index, placement).
    """
    def __init__(self, doc):
        self.doc = doc

    # --- low-level helpers ---
    @staticmethod
    def _iter_block_items(doc):
        from docx.text.paragraph import Paragraph
        from docx.table import Table
        from docx.oxml.text.paragraph import CT_P
        from docx.oxml.table import CT_Tbl

        parent_elm = doc._element
        for child in parent_elm.body.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, doc)
            elif isinstance(child, CT_Tbl):
                yield Table(child, doc)

    @staticmethod
    def _first_cell_text(cell) -> str:
        return " ".join(p.text for p in cell.paragraphs).strip()

    @staticmethod
    def _clean_spaces(s: str) -> str:
        return re.sub(r"\s+", " ", (s or "").strip())

    @staticmethod
    def _strip_trailing_paren(s: str) -> str:
        return re.sub(r"\s*\([^()]*\)\s*$", "", s or "").strip()

    @staticmethod
    def _dedupe_phrase(s: str) -> str:
        t = DocumentUnderstandingAgent._clean_spaces(s)
        if not t:
            return t
        parts = t.split()
        n = len(parts)
        if n % 2 == 0:
            mid = n // 2
            if parts[:mid] == parts[mid:]:
                return " ".join(parts[:mid])
        t2 = re.sub(r"(?:\s*\([^)]*\)\s*){2,}$",
                    lambda m: " " + m.group(0).strip().split(")")[0] + ")", t)
        return t2

    @staticmethod
    def _looks_like_section_title(text: str) -> bool:
        if not text:
            return False
        t = unicodedata.normalize("NFKC", text).strip()
        if len(t) > 180:
            return False
        if t.lower().startswith("section ") or t.endswith(":"):
            return True
        letters = [c for c in t if c.isalpha()]
        if not letters:
            return False
        caps_ratio = sum(1 for c in letters if c.isupper()) / max(1, len(letters))
        return caps_ratio >= 0.45 and len(t.split()) <= 15

    @staticmethod
    def _is_table_header_row(tbl, ri: int) -> bool:
        try:
            row = tbl.rows[ri]
        except Exception:
            return False
        merged_like = len(row.cells) == 1
        shaded_like = False
        try:
            for c in row.cells:
                if c._tc.xpath('.//w:tcPr/w:shd', namespaces=c._tc.nsmap):
                    shaded_like = True
                    break
        except Exception:
            pass
        row_text = " | ".join(DocumentUnderstandingAgent._first_cell_text(c) for c in row.cells).strip()
        simple = (":" not in row_text) and (len(row_text) <= 160)
        return merged_like or (shaded_like and simple)

    @staticmethod
    def _label_cols_for_table(tbl) -> list:
        try:
            ncols = len(tbl.rows[0].cells)
        except Exception:
            return [0]
        if ncols <= 1:
            return [0]
        return [c for c in range(0, ncols, 2)]

    def build_template(self, dry_run=False) -> List[Dict[str, Any]]:
        from docx.text.paragraph import Paragraph
        from docx.table import Table

        underline_pat = re.compile(r"^(.*?[:：﹕꞉˸፡︓])\s*[_\u2014\-.\s]{3,}$")

        page = 0                      # python-docx doesn't have pagination
        current_section = ""          # default section
        raw_fields: List[Dict[str, Any]] = []

        def _push(label: str, section: str, placement: str):
            lab = (label or "").strip()
            if not lab:
                return
            sec = (section or "").strip()
            lab_short = lab.split("\n", 1)[0].strip()
            raw_fields.append({
                "label": lab,
                "label_short": lab_short,
                "section": sec,
                "page": page,
                "placement": placement,
            })

        def _row0_mode_and_sections(tbl):
            """("single", caption,"") or ("percol","",{ci->hdr}) or ("none","",{})"""
            # Clean row-0 texts (strip trailing colon/paren)
            texts = []
            try:
                row0 = tbl.rows[0]
            except Exception:
                return ("none", "", {})
            for ci in range(len(row0.cells)):
                t = self._first_cell_text(row0.cells[ci]).strip()
                if t:
                    t = re.sub(r"[:：﹕꞉˸]\s*$", "", t).strip()
                    t = self._strip_trailing_paren(t)
                    t = self._clean_spaces(t)
                texts.append(t)

            nonempty = [t for t in texts if t]
            if not nonempty:
                return ("none", "", {})

            # Short multi-column headers => percol (e.g., Contact 1..4)
            short_cells = [t for t in nonempty if len(t.split()) <= 4]
            if len(nonempty) >= 2 and len(short_cells) >= max(2, int(0.6 * len(nonempty))) and len(set(short_cells)) >= 2:
                col_map, last = {}, ""
                for ci, t in enumerate(texts):
                    if t:
                        last = t
                    col_map[ci] = last
                for k, v in list(col_map.items()):
                    col_map[k] = self._dedupe_phrase(v) if v else v
                return ("percol", "", col_map)

            # Otherwise join all as a caption
            joined = self._clean_spaces(" ".join(nonempty))
            joined = self._dedupe_phrase(joined)
            return ("single", joined, {})

        # ---- scan document ----
        for block in self._iter_block_items(self.doc):
            if isinstance(block, Paragraph):
                txt = (block.text or "").strip()
                if not txt:
                    continue
                if self._looks_like_section_title(txt):
                    current_section = self._strip_trailing_paren(txt)
                    continue
                m = underline_pat.match(txt)
                if m:
                    label = re.sub(r"[:：﹕꞉˸፡︓]\s*$", "", m.group(1)).strip()
                    if label:
                        _push(label, current_section, "para_underline")
                    continue

            if isinstance(block, Table):
                header_present = self._is_table_header_row(block, 0)
                mode, joined_caption, col_sections = _row0_mode_and_sections(block) if header_present else ("none", "", {})

                if mode == "single" and joined_caption:
                    table_sec = joined_caption
                    current_section = table_sec
                elif mode == "percol":
                    table_sec = current_section
                else:
                    table_sec = current_section

                start_row = 1 if header_present else 0
                label_cols = self._label_cols_for_table(block)
                for ri in range(start_row, len(block.rows)):
                    row = block.rows[ri]
                    for ci in label_cols:
                        try:
                            label_cell = row.cells[ci]
                        except Exception:
                            continue
                        label = self._first_cell_text(label_cell).strip()
                        if not label:
                            continue
                        label_clean = self._clean_spaces(re.sub(r"[:：﹕꞉˸፡︓]\s*$", "", label).strip())
                        label_clean = self._strip_trailing_paren(label_clean)
                        if not label_clean:
                            continue
                        if mode == "percol":
                            sec_for_pair = self._clean_spaces(col_sections.get(ci) or table_sec or "")
                        else:
                            sec_for_pair = table_sec or ""
                        _push(label_clean, sec_for_pair, "table_pair")

        # per-section indexing
        fields = []
        counters: Dict[Tuple[str, str], int] = {}
        for f in raw_fields:
            key = (f["section"].lower(), f["label_short"].lower())
            counters[key] = counters.get(key, 0) + 1
            fields.append({**f, "index": counters[key]})

        if dry_run:
            print(f"[DRY] Discovered fields: {len(fields)}")
            for f in fields:
                sec = f["section"] or "—"
                print(f"  • {f['label_short']} | Sec: {sec} | Idx: {f['index']} | {f['placement']}")
        return fields


class MatchingAgent:
    """Resolves best value for a (section,label,index) against lookup rows."""
    STOPWORDS = {
        "name", "number", "no", "date", "day", "month", "year", "address",
        "phone", "email", "applicant", "the", "and", "or", "of"
    }

    def __init__(self, rows: List[Dict[str, Any]]):
        self.rows = rows or []
        # Ensure normalization keys exist (robust to older LookupAgent output)
        for r in self.rows:
            f = r.get("Field", "")
            s = r.get("Section", "")
            r.setdefault("field_norm", self._normalize_strict(f))
            r.setdefault("section_norm", self._normalize_strict(s) if s else "")
            r.setdefault("_field_norm_loose", self._normalize_loose(f))
            r.setdefault("_section_norm_loose", self._normalize_loose(s) if s else "")

    # ---------- normalization helpers ----------
    @staticmethod
    def _normalize_strict(s: str) -> str:
        import unicodedata, re
        s = unicodedata.normalize("NFKC", str(s or ""))
        s = re.sub(r"\(.*?\)", "", s)
        s = re.sub(r"[:：﹕꞉˸፡︓]\s*$", "", s)
        s = re.sub(r"\s+", " ", s).strip().lower()
        return s

    @staticmethod
    def _normalize_loose(s: str) -> str:
        import unicodedata, re, string
        s = unicodedata.normalize("NFKC", str(s or ""))
        s = re.sub(r"\(.*?\)", "", s)
        s = re.sub(r"[:：﹕꞉˸፡︓]\s*$", "", s)
        s = re.sub(r"\s+", " ", s).strip().lower()
        s = s.translate(str.maketrans("", "", string.punctuation))
        return s

    @classmethod
    def _wordset(cls, s: str) -> set:
        import re
        words = {w for w in re.split(r"\W+", str(s).lower()) if len(w) >= 3}
        return {w for w in words if w not in cls.STOPWORDS}

    @staticmethod
    def _sim_loose(a: str, b: str) -> float:
        from difflib import SequenceMatcher
        return SequenceMatcher(None, a, b).ratio()

    # ---------- resolver ----------
    def resolve(
            self,
            label: str,
            section: str,
            index: int,
            page: Optional[int] = None,
            *,
            min_fuzzy: float = 0.82,
            strict_index: bool = True,
            exact_only: bool = False,
            require_token_overlap: bool = True,
    ) -> Optional[str]:
        """
        - exact_only=True disables fuzzy fallback entirely.
        - require_token_overlap=True demands at least one **non-stopword** token match.
        """
        label_strict = self._normalize_strict(label)
        section_strict = self._normalize_strict(section) if section else ""
        label_loose = self._normalize_loose(label)
        section_loose = self._normalize_loose(section) if section else ""

        # ---- Step 1: strict exact field match ----
        candidates = [r for r in self.rows if r.get("field_norm") == label_strict]
        if not candidates:
            # sometimes authoring introduced slightly different strict norms; try loose
            candidates = [r for r in self.rows if r.get("field_norm") == label_loose]

        if candidates:
            # prefer same section if provided
            if section_strict:
                same_sec = [r for r in candidates if r.get("section_norm") == section_strict]
                if same_sec:
                    candidates = same_sec

            # prefer same page if provided
            if page is not None:
                same_pg = [r for r in candidates if r.get("Page") is not None and int(r["Page"]) == int(page)]
                if same_pg:
                    candidates = same_pg

            # index strictness
            if strict_index:
                with_idx = [r for r in candidates if r.get("Index") is not None]
                if with_idx:
                    exact_idx = [r for r in with_idx if int(r["Index"]) == int(index)]
                    if exact_idx:
                        candidates = exact_idx
                    else:
                        return None

            def score(r):
                s = 0
                if section_strict and r.get("section_norm") == section_strict:
                    s += 6
                if r.get("Index") is not None and int(r["Index"]) == int(index):
                    s += 3
                return s

            best = max(candidates, key=score, default=None)
            # --- Final cascade fallback:
            # If we still have nothing and caller asked for index N,
            # but there's exactly one row in lookup with the same field (loose) and NO Index,
            # use that value for any row (helps matrix/per-column tables).
            if best is None and strict_index:
                loose_pool = [r for r in self.rows if r.get("_field_norm_loose") == label_loose]
                no_idx = [r for r in loose_pool if r.get("Index") is None or str(r.get("Index")).strip() == ""]
                if len(no_idx) == 1:
                    return no_idx[0]["Value"]

            return best["Value"] if best else None

        # ---- Step 2: fuzzy disabled? ----
        if exact_only:
            return None

        # ---- Step 3: fuzzy fallback (loose norms) ----
        field_pool = list({r.get("_field_norm_loose") for r in self.rows})
        field_pool = [x for x in field_pool if x]
        if not field_pool:
            return None

        best_norm = None
        best_sim = -1.0
        for x in field_pool:
            sim = self._sim_loose(label_loose, x)
            if sim > best_sim:
                best_sim, best_norm = sim, x

        if best_norm is None or best_sim < min_fuzzy:
            return None

        # token overlap with stopwords removed — prevents 'First name' → 'Applicant Name (Entity)'
        if require_token_overlap and not (self._wordset(label_loose) & self._wordset(best_norm)):
            return None

        candidates = [r for r in self.rows if r.get("_field_norm_loose") == best_norm]

        # section preference on loose norm
        if section_loose:
            same_sec = [r for r in candidates if r.get("_section_norm_loose") == section_loose]
            if same_sec:
                candidates = same_sec

        # index strictness again
        if strict_index:
            with_idx = [r for r in candidates if r.get("Index") is not None]
            if with_idx:
                exact_idx = [r for r in with_idx if int(r["Index"]) == int(index)]
                if exact_idx:
                    candidates = exact_idx
                else:
                    return None

        best = max(
            candidates,
            key=lambda r: (
                1 if (r.get("Index") is not None and int(r["Index"]) == int(index)) else 0,
                1 if (section_loose and r.get("_section_norm_loose") == section_loose) else 0,
            ),
            default=None,
        )
        return best["Value"] if best else None







class FillingAgent:
    """Writes matches into the DOCX (placeholders, checkboxes, tables, underlines)."""
    def __init__(self, doc, matcher: MatchingAgent):
        self.doc = doc
        self.matcher = matcher

    # --- shared doc iteration ---
    def _iter_all_paragraphs_and_cells(self) -> Iterable:
        for p in self.doc.paragraphs:
            yield p
        for tbl in self.doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p

    # --- helpers for writing ---
    @staticmethod
    def _set_paragraph_text_keep_simple_format(p, new_text: str):
        while p.runs:
            r = p.runs[0]
            r.clear()
            r.text = ""
            r.element.getparent().remove(r.element)
        p.add_run(new_text)

    @staticmethod
    def _first_cell_text(cell) -> str:
        return " ".join(p.text for p in cell.paragraphs).strip()

    @staticmethod
    def _cell_has_box_target(cell) -> bool:
        try:
            if cell._tc.xpath('.//w:tcPr/w:shd', namespaces=cell._tc.nsmap):
                return True
        except Exception:
            pass
        for p in cell.paragraphs:
            for r in p.runs:
                if (r.text or "").strip().upper() == "FORMTEXT":
                    return True
            vis = "".join((r.text or "") for r in p.runs)
            if vis and not re.search(r"[A-Za-z0-9]", vis):
                return True
        return False

    @staticmethod
    def _write_value_inside_box(cell, value: str):
        """Prefer writing inside legacy FORMTEXT result or placeholder/empty spots."""
        import re
        from docx.oxml.shared import OxmlElement, qn

        val = value or ""
        PLACEHOLDER_RE = re.compile(r"[_\-\u2014\.\s\u2002\u2003\u2007\u2009\u00A0]+")
        ALNUM_RE = re.compile(r"[A-Za-z0-9]")

        # 1) legacy formtext
        for para in cell.paragraphs:
            p = para._element
            instr_nodes = p.xpath(".//w:instrText")
            has_formtext = any(("FORMTEXT" in (n.text or "").upper()) for n in instr_nodes)
            if not has_formtext:
                continue

            result_runs = []
            got_separate = False
            for r in p.xpath("./w:r"):
                if r.xpath("./w:fldChar[@w:fldCharType='begin']"):
                    got_separate = False
                    result_runs = []
                    continue
                if r.xpath("./w:fldChar[@w:fldCharType='separate']"):
                    got_separate = True
                    continue
                if r.xpath("./w:fldChar[@w:fldCharType='end']"):
                    break
                if got_separate:
                    result_runs.append(r)

            if result_runs:
                first = result_runs[0]
                for rr in result_runs:
                    for t in rr.xpath("./w:t"):
                        t.getparent().remove(t)
                t = OxmlElement('w:t')
                t.set(qn('xml:space'), 'preserve')
                t.text = val
                first.append(t)
                return

        # 2) underline/placeholder group
        def glen(group):
            return sum(len((r.text or "").replace(" ", "")) for r in group)

        best_group = []
        shaded_paras = []
        box_like_paras = []

        def _para_is_shaded(para) -> bool:
            try:
                return bool(para._element.xpath('.//w:pPr/w:shd'))
            except Exception:
                return False

        for para in cell.paragraphs:
            if _para_is_shaded(para):
                shaded_paras.append(para)
            vis = "".join((r.text or "") for r in para.runs)
            if vis and not ALNUM_RE.search(vis):
                box_like_paras.append(para)

            cur, best_here = [], []
            for run in para.runs:
                t = run.text or ""
                if t and PLACEHOLDER_RE.fullmatch(t):
                    cur.append(run)
                else:
                    if glen(cur) > glen(best_here):
                        best_here = cur[:]
                    cur = []
            if glen(cur) > glen(best_here):
                best_here = cur[:]
            if glen(best_here) > glen(best_group):
                best_group = best_here

        if best_group:
            L = max(1, glen(best_group))
            best_group[0].text = val[:L]
            for r in best_group[1:]:
                r.text = ""
            return

        if shaded_paras:
            # just append text in shaded para
            empties = [r for r in shaded_paras[0].runs if not (r.text or "").strip()]
            if empties:
                empties[-1].text = val
            else:
                shaded_paras[0].add_run(val)
            return

        if FillingAgent._cell_has_box_target(cell):
            p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            p.add_run(val)
            return

        p = cell.paragraphs[-1] if cell.paragraphs else cell.add_paragraph()
        p.add_run(val)

    # --- actions ---
    def replace_placeholders(self, dry_run=False):
        pats = [
            re.compile(r"\{\{\s*(.*?)\s*\}\}"),
            re.compile(r"\[\[\s*(.*?)\s*\]\]"),
            re.compile(r"\$\{\s*(.*?)\s*\}"),
        ]
        def _sub_key(key: str) -> Optional[str]:
            # placeholder has no section/index context -> use fuzzy across all rows
            # Try as-is first
            return self.matcher.resolve(label=key, section="", index=1, page=None,
                                        min_fuzzy=0.82, strict_index=False)

        for p in self._iter_all_paragraphs_and_cells():
            old = p.text
            new = old
            for pat in pats:
                def _sub(m):
                    k = (m.group(1) or "").strip()
                    v = _sub_key(k)
                    return str(v) if v is not None else m.group(0)
                new = pat.sub(_sub, new)
            if new != old:
                if dry_run:
                    print(f"[DRY][DOCX] placeholder: '{old}' → '{new}'")
                else:
                    self._set_paragraph_text_keep_simple_format(p, new)

    def fill_checkboxes(self, dry_run: bool = False):
        """
        Toggle one or more checkbox tokens in a paragraph based on lookup values.

        Handles both orders and multiples per line:
          - '☐ Label' / '[ ] Label'   (token → label)
          - 'Label ☐' / 'Label [ ]'   (label → token)
          - Repeated in one paragraph: '☐ A    ☐ B    ☐ C'  or  'A ☐    B ☐'

        NEW: If no token is present but the line looks like a checkbox sentence
        (like the “benefit plan investor … (tick if applicable)” rows), resolve
        the whole sentence and prefix with ☒/☐ accordingly.
        """
        import re, unicodedata, string

        BOX_RE = re.compile(r'(?:\[\s*[xX✓]?\s*\]|[□☐☑☒])')

        def _norm(s: str) -> str:
            s = unicodedata.normalize("NFKC", str(s or "")).strip()
            return re.sub(r"\s+", " ", s)

        def _norm_key(s: str) -> str:
            t = _norm(s).lower()
            return t.translate(str.maketrans("", "", string.punctuation))

        def _truthy(val):
            if val is None: return None
            s = (str(val).strip().lower())
            if s in {"y","yes","true","1","x","✓","check","checked"}:   return True
            if s in {"n","no","false","0","uncheck","unchecked"}:       return False
            return None

        # Rebuild the paragraph text once instead of touching runs piecemeal.
        def _replace_nth_token_in_paragraph(p, nth: int, make_checked: bool) -> bool:
            full = "".join(r.text or "" for r in p.runs)
            matches = list(BOX_RE.finditer(full))
            if nth >= len(matches):
                return False
            m = matches[nth]
            tok_text = full[m.start():m.end()]
            if tok_text.strip().startswith("["):
                repl = "[x]" if make_checked else "[ ]"
            else:
                repl = "☒" if make_checked else "☐"
            new_text = full[:m.start()] + repl + full[m.end():]
            self._set_paragraph_text_keep_simple_format(p, new_text)
            return True

        # Build (token_index, label) pairs for a paragraph.
        def _pairs_from_paragraph_text(text: str):
            tokens = list(BOX_RE.finditer(text))
            if not tokens:
                return []
            pairs = []
            token_first = tokens[0].start() <= len(re.match(r"^\s*", text).group(0))
            if token_first:
                for i, m in enumerate(tokens):
                    start = m.end()
                    end = tokens[i + 1].start() if i + 1 < len(tokens) else len(text)
                    label = text[start:end].strip(" \t-:•–")
                    if label:
                        pairs.append((i, label))
            else:
                prev_end = 0
                for i, m in enumerate(tokens):
                    label = text[prev_end:m.start()].strip(" \t-:•–")
                    if label:
                        pairs.append((i, label))
                    prev_end = m.end()
            return pairs

        # Tidy long option labels (drop list markers / trailing helpers).
        def _tidy_option_label(s: str) -> str:
            t = s.strip()
            t = re.sub(r'^\s*(?:[\(\[]?[ivxlcdmIVXLCDM0-9a-zA-Z]+[\)\].:-]|\-|•)\s+', '', t)
            t = re.sub(r'\s*\([^()]*\)\s*$', '', t).strip()
            t = re.sub(r'\s+', ' ', t)
            return t

        # Heuristic: a full paragraph that *looks* like a checkbox sentence.
        TRAIL_COLON_RE = re.compile(r"[:：﹕꞉˸፡︓]\s*$")
        def _looks_checkbox_sentence(s: str) -> bool:
            t = _norm(s)
            if TRAIL_COLON_RE.search(t):
                return False
            if len(t) < 24 and len(t.split()) < 6:
                return False
            if re.match(r"^\([a-z]\)\s", t):   # (a) / (b) …
                return True
            KW = ("tick if applicable", "benefit plan investor", "plan investor", "tick if")
            tl = t.lower()
            return any(k in tl for k in KW)

        for p in self._iter_all_paragraphs_and_cells():
            raw = p.text or ""
            if not raw.strip():
                continue

            # ---------- Primary path: explicit tokens present ----------
            if BOX_RE.search(raw):
                pairs = _pairs_from_paragraph_text(raw)
                if not pairs:
                    continue

                cleaned = []
                for nth, lbl in pairs:
                    key = _norm_key(lbl)
                    if key in {"or"}:
                        continue
                    if ":" in lbl:
                        continue
                    lbl2 = _tidy_option_label(lbl)
                    if len(re.sub(r'[^A-Za-z]', '', lbl2)) < 4:
                        continue
                    if len(lbl2) > 200:
                        continue
                    cleaned.append((nth, lbl2))

                if not cleaned:
                    continue

                for nth, label in cleaned:
                    val = self.matcher.resolve(label=label,
                                               section="",
                                               index=1,
                                               page=None,
                                               min_fuzzy=0.78,
                                               strict_index=False)
                    yn = _truthy(val)
                    if yn is None:
                        if dry_run:
                            print(f"[DRY][DOCX] checkbox: '{label}' → (no match; leave as-is)")
                        continue
                    if dry_run:
                        print(f"[DRY][DOCX] checkbox: '{label}' → {'CHECK' if yn else 'UNCHECK'}")
                    else:
                        _replace_nth_token_in_paragraph(p, nth, make_checked=yn)
                continue  # done with this paragraph

            # ---------- Fallback path: no token, but line looks like a checkbox sentence ----------
            line = raw.strip()
            if not _looks_checkbox_sentence(line):
                continue

            label_for_lookup = _tidy_option_label(line)
            val = self.matcher.resolve(label=label_for_lookup,
                                       section="",
                                       index=1,
                                       page=None,
                                       min_fuzzy=0.80,
                                       strict_index=False)
            yn = _truthy(val)
            if yn is None:
                # also try the un-tidied text once
                val2 = self.matcher.resolve(label=line,
                                            section="",
                                            index=1,
                                            page=None,
                                            min_fuzzy=0.80,
                                            strict_index=False)
                yn = _truthy(val2)

            if yn is None:
                if dry_run:
                    print(f"[DRY][DOCX] checkbox (no-token; no match): '{label_for_lookup[:72]}'")
                continue

            new_text = f"{'☒' if yn else '☐'} {label_for_lookup}"
            if dry_run:
                print(f"[DRY][DOCX] checkbox (no-token): '{label_for_lookup[:72]}' → {'CHECK' if yn else 'UNCHECK'}")
            else:
                self._set_paragraph_text_keep_simple_format(p, new_text)




    def fill_two_column_checkbox_rows(self, dry_run: bool = False):
        """
        Tick checkboxes where a table row has a short 'box-like' cell adjacent
        to a long label cell. Works for either order: [box][label] or [label][box].
        Will insert ☒/☐ when no token exists.

        Notes:
          • Truly empty cells are not 'box-like' unless the adjacent cell clearly
            looks like a checkbox sentence (fallback below).
          • This keeps the HID row excluded and picks up the p3/p4 sentences.
        """
        import re, unicodedata, string

        BOX_RE         = re.compile(r'(?:\[\s*[xX✓]?\s*\]|[□☐☑☒])')
        LETTER_RE      = re.compile(r"[A-Za-z]")
        FILLER_SHORTRE = re.compile(r"^[-_.—–·•\u00A0\s]{1,6}$")           # short, non-empty filler
        TRAIL_COLON_RE = re.compile(r"[:：﹕꞉˸፡︓]\s*$")

        def _cell_text(cell) -> str:
            return " ".join(p.text for p in cell.paragraphs).strip()

        def _norm_inline(s: str) -> str:
            s = unicodedata.normalize("NFKC", str(s or "")).strip()
            s = (s.replace("“", '"').replace("”", '"')
                 .replace("‘", "'").replace("’", "'")
                 .replace("–", "-").replace("—", "-"))
            return re.sub(r"\s+", " ", s)

        def _truthy_local(val: str):
            if val is None: return None
            ss = (str(val).strip().lower())
            if ss in {"y","yes","true","1","x","✓","check","checked"}:   return True
            if ss in {"n","no","false","0","uncheck","unchecked"}:       return False
            return None

        # Box-like only if explicit token OR short non-empty filler.
        def _looks_box_cell(txt: str) -> bool:
            if BOX_RE.search(txt or ""):
                return True
            if not txt:
                return False                          # empty ≠ box-like
            compact = re.sub(r"\s+", "", txt)
            if not compact:
                return False
            return bool(FILLER_SHORTRE.match(txt))

        # Checkbox-sentence detector for fallback when the box cell is an empty shape.
        def _looks_checkbox_sentence(s: str) -> bool:
            t = _norm_inline(s).lower()
            if TRAIL_COLON_RE.search(t):
                return False                           # "Label:" → not a checkbox sentence
            if len(t) < 24 and len(t.split()) < 6:
                return False
            if re.match(r"^\([a-z]\)\s", t):           # (a) (b) (c) ...
                return True
            KW = ("tick if applicable", "benefit plan investor", "plan investor", "tick if")
            return any(k in t for k in KW)

        def _set_token_in_cell(cell, make_checked: bool) -> bool:
            if not cell.paragraphs:
                p = cell.add_paragraph()
            else:
                p = cell.paragraphs[0]
            full = "".join(r.text or "" for r in p.runs)
            m = BOX_RE.search(full)
            if m:
                tok_text = full[m.start():m.end()]
                repl = "[x]" if tok_text.strip().startswith("[") else "☒"
                if not make_checked:
                    repl = "[ ]" if tok_text.strip().startswith("[") else "☐"
                new_text = full[:m.start()] + repl + full[m.end():]
                self._set_paragraph_text_keep_simple_format(p, new_text)
                return True
            self._set_paragraph_text_keep_simple_format(p, ("☒" if make_checked else "☐"))
            return True

        for tbl in self.doc.tables:
            for row in tbl.rows:
                cells = row.cells
                if len(cells) < 2:
                    continue

                for ci in range(len(cells) - 1):
                    left_txt  = _cell_text(cells[ci])
                    right_txt = _cell_text(cells[ci + 1])

                    pair = None
                    # explicit token + long-ish label
                    if BOX_RE.search(left_txt or "") and _looks_checkbox_sentence(right_txt):
                        pair = (cells[ci], right_txt)
                    elif BOX_RE.search(right_txt or "") and _looks_checkbox_sentence(left_txt):
                        pair = (cells[ci + 1], left_txt)

                    # short non-empty filler + long checkbox sentence
                    if pair is None:
                        left_is_box  = _looks_box_cell(left_txt)
                        right_is_box = _looks_box_cell(right_txt)
                        if left_is_box and _looks_checkbox_sentence(right_txt):
                            pair = (cells[ci], right_txt)
                        elif right_is_box and _looks_checkbox_sentence(left_txt):
                            pair = (cells[ci + 1], left_txt)

                    # NEW fallback: empty cell (likely a drawn box) + checkbox-like sentence on the other side
                    if pair is None:
                        if (not left_txt or not LETTER_RE.search(left_txt)) and _looks_checkbox_sentence(right_txt):
                            pair = (cells[ci], right_txt)
                        elif (not right_txt or not LETTER_RE.search(right_txt)) and _looks_checkbox_sentence(left_txt):
                            pair = (cells[ci + 1], left_txt)

                    if not pair:
                        continue

                    token_cell, label_text = pair
                    label_for_lookup = _norm_inline(TRAIL_COLON_RE.sub("", label_text).strip())
                    val = self.matcher.resolve(
                        label=label_for_lookup,
                        section="",
                        index=1,
                        page=None,
                        min_fuzzy=0.82,
                        strict_index=False
                    )
                    yn = _truthy_local(val)
                    if yn is None:
                        if dry_run:
                            print(f"[DRY][DOCX] checkbox-2col (no match): '{label_for_lookup[:72]}'")
                        continue

                    if dry_run:
                        print(f"[DRY][DOCX] checkbox-2col: '{label_for_lookup[:72]}' → {'CHECK' if yn else 'UNCHECK'}")
                    else:
                        _set_token_in_cell(token_cell, yn)




    def fill_checkbox_option_groups(self, dry_run=False):
        """
        Detect rows that look like a horizontal checkbox option group:
          [Option A] [Option B] [Option C]
        followed by a row of empty/placeholder cells underneath.
        Then tick the cell under the selected option.

        Matching rules:
          • Prefer boolean rows for each option label:
                Field = 'Initial Subscription'  Value = yes/no
                Field = 'Additional Subscription'  Value = yes/no
          • Otherwise, try a single-choice row for the group:
                Field = <group name>  Value = one of the option labels
            Group name is taken as the nearest non-empty text to the LEFT
            of the first option in the same row (or the whole row’s combined text).
        """
        import re

        TRAIL_COLON_RE = re.compile(r"[:：﹕꞉˸፡︓]\s*$")
        LETTER_RE = re.compile(r"[A-Za-z]")

        def _cell_text(cell) -> str:
            return " ".join(p.text for p in cell.paragraphs).strip()

        def _is_short_label(t: str) -> bool:
            if not t:
                return False
            # strip trailing colon and compress spaces
            t = TRAIL_COLON_RE.sub("", t).strip()
            # short-ish label (≤ 4 words) and contains letters
            return len(t.split()) <= 4 and bool(LETTER_RE.search(t))

        def _is_placeholder_like(text: str) -> bool:
            if not text:
                return True
            t = re.sub(r"\s+", "", text)
            # treat as placeholder if it has no letters (digits allowed)
            return not LETTER_RE.search(t)

        def _resolve_bool_for_option(option_label: str, idx: int) -> Optional[bool]:
            v = self.matcher.resolve(label=option_label, section="", index=idx,
                                     page=None, min_fuzzy=0.82, strict_index=False)
            if v is None:
                return None
            s = (str(v).strip().lower())
            if s in {"y","yes","true","1","x","✓","check","checked"}:
                return True
            if s in {"n","no","false","0","uncheck","unchecked"}:
                return False
            return None

        def _resolve_choice_for_group(group_label: str, idx: int) -> Optional[str]:
            v = self.matcher.resolve(label=group_label, section="", index=idx,
                                     page=None, min_fuzzy=0.82, strict_index=False)
            if v is None:
                return None
            return str(v).strip()

        for tbl in self.doc.tables:
            if len(tbl.rows) < 2:
                continue

            # scan each pair of consecutive rows
            for ri in range(0, len(tbl.rows) - 1):
                row_labels = tbl.rows[ri]
                row_boxes  = tbl.rows[ri + 1]

                # gather contiguous short labels on the label row
                labels: list[tuple[int, str]] = []
                for ci, c in enumerate(row_labels.cells):
                    t = _cell_text(c)
                    if t:
                        t = TRAIL_COLON_RE.sub("", t).strip()
                    if _is_short_label(t):
                        labels.append((ci, t))

                # need at least two options to qualify as a group
                if len(labels) < 2:
                    continue

                # the row below should have placeholder-like cells at those columns
                # (don't overwrite if user typed something)
                if any(ci >= len(row_boxes.cells) or not _is_placeholder_like(_cell_text(row_boxes.cells[ci]))
                       for ci, _ in labels):
                    continue

                # choose an index hint: try to take the first numeric content to the LEFT
                # of the group in either labels or boxes row; fallback to 1
                idx_hint = 1
                left_ci  = min(ci for ci, _ in labels)
                for probe_ci in range(left_ci - 1, -1, -1):
                    t1 = _cell_text(row_labels.cells[probe_ci]) if probe_ci < len(row_labels.cells) else ""
                    t2 = _cell_text(row_boxes.cells[probe_ci]) if probe_ci < len(row_boxes.cells) else ""
                    m = re.search(r"\b(\d+)\b", t1 or t2)
                    if m:
                        try:
                            idx_hint = int(m.group(1))
                            break
                        except Exception:
                            pass

                # a "group label" for single-choice fallback: take the closest non-empty
                # text to the LEFT of the first option in the same label row; if none,
                # fall back to the whole label row text
                group_label = ""
                for probe_ci in range(left_ci - 1, -1, -1):
                    t = _cell_text(row_labels.cells[probe_ci])
                    if t:
                        group_label = TRAIL_COLON_RE.sub("", t).strip()
                        if group_label:
                            break
                if not group_label:
                    group_label = TRAIL_COLON_RE.sub("", " ".join(_cell_text(c) for c in row_labels.cells)).strip()

                # Evaluate/tick per option
                for ci, option in labels:
                    # 1) per-option boolean
                    yn = _resolve_bool_for_option(option, idx_hint)

                    # 2) single-choice field fallback
                    if yn is None and group_label:
                        chosen = _resolve_choice_for_group(group_label, idx_hint)
                        if chosen:
                            yn = (option.lower() == chosen.strip().lower())

                    if yn is None:
                        continue

                    target = row_boxes.cells[ci]
                    mark = "☒" if yn else "☐"
                    if dry_run:
                        print(f"[DRY][DOCX] checkbox-group: row {ri+1} idx={idx_hint} option='{option}' → {mark}")
                    else:
                        # write the mark into the box row cell without touching the label row
                        existing = _cell_text(target)
                        if existing and "☐" in existing:
                            # replace the first ☐ with ☒/☐
                            new = existing.replace("☐", mark, 1)
                            self._set_paragraph_text_keep_simple_format(target.paragraphs[0], new)
                        else:
                            self._write_value_inside_box(target, mark)


    def fill_inline_checkbox_groups(self, dry_run=False):
        """
        Detect and tick inline checkbox groups like:
            "Initial Subscription    Additional Subscription"
        inside one paragraph/cell (with or without existing box glyphs).

        Rules:
          • Split a line into 2–6 short options (<= 4 words each) separated by
            2+ spaces, tabs, or "  OR  ".
          • If a Field matching an option exists in lookup and is truthy → ☒ that option.
          • If no per-option fields, but a single Field matching the entire line exists,
            and its Value equals one of the options → ☒ that option.
          • Non-selected options get ☐ (only within that same line).
        """
        SEP_RE = re.compile(r"(?:\t+|\s{2,}|(?:\s+OR\s+))", re.IGNORECASE)
        # treat these as existing box glyphs that we won't duplicate
        BOX_RE = re.compile(r"^\s*[□☐☒\[\]\(\)]\s*")

        def split_options(text: str) -> list:
            # remove leading box if present on the whole line
            t = text.strip()
            t = re.sub(r"^\s*(?:[□☐☒]\s*|\[\s?[xX✓]?\s?\]\s*)", "", t)
            parts = [p.strip() for p in SEP_RE.split(t) if p.strip()]
            # keep only "shortish" items to avoid splitting normal sentences
            parts = [p for p in parts if len(p.split()) <= 4]
            # need at least 2 distinct short parts to count as a group
            if 2 <= len(parts) <= 6 and len(set(parts)) >= 2:
                return parts
            # NEW: fallback when there are no glyphs and no clear separators,
            # e.g., "Initial Subscription                      Additional Subscription"
            alt = _split_two_labels_no_glyphs(text)
            if alt and len(alt) >= 2:
                return alt
            return []

        def truthy_val_for(label: str) -> Optional[bool]:
            v = self.matcher.resolve(label=label, section="", index=1,
                                     page=None, min_fuzzy=0.82, strict_index=False)
            return _truthy(v)

        def chosen_by_group_value(full_line: str, options: list) -> Optional[str]:
            """
            Fallback: if the group text itself is a Field and its Value equals
            one of the option labels, select that option.
            """
            v = self.matcher.resolve(label=full_line, section="", index=1,
                                     page=None, min_fuzzy=0.82, strict_index=False)
            if not v:
                return None
            v_norm = re.sub(r"\s+", " ", str(v).strip()).lower()
            for opt in options:
                if re.sub(r"\s+", " ", opt.lower()) == v_norm:
                    return opt
            return None

        def render_group(options: list, selected: Optional[set], original: str) -> str:
            # If the original already has explicit boxes per option, rebuild cleanly
            # to avoid run/spacing issues. Otherwise, prefix each option with a box.
            items = []
            for o in options:
                box = "☒" if (selected and o in selected) else "☐"
                items.append(f"{box} {o}")
            return "  ".join(items)

        # Iterate every paragraph (body + table cells)
        for p in self._iter_all_paragraphs_and_cells():
            raw = p.text or ""
            line = raw.strip()
            if not line:
                continue

            options = split_options(line)
            if not options:
                continue  # not an inline group

            # 1) Try per-option fields
            selected = set()
            any_info = False
            for opt in options:
                yn = truthy_val_for(opt)
                if yn is not None:
                    any_info = True
                if yn is True:
                    selected.add(opt)

            # 2) Fallback to a single group field whose value equals an option label
            if not selected and not any_info:
                chosen = chosen_by_group_value(line, options)
                if chosen:
                    selected = {chosen}
                    any_info = True

            # Only touch the paragraph if we actually learned something
            if not any_info:
                continue

            new_text = render_group(options, selected, raw)
            if new_text != raw:
                if dry_run:
                    sel_str = ", ".join(sorted(selected)) if selected else "∅"
                    print(f"[DRY][DOCX] inline-checkbox: '{line}' → [{sel_str}]")
                else:
                    self._set_paragraph_text_keep_simple_format(p, new_text)



    def fill_underline_lines(self, dry_run=False):
        us_pat = re.compile(r"^(.*?[:：﹕꞉˸፡︓])\s*_+\s*$")
        for p in self._iter_all_paragraphs_and_cells():
            txt = p.text.strip()
            if not txt:
                continue
            m = us_pat.match(txt)
            if not m:
                continue
            label_full = m.group(1)
            lab = re.sub(r"[:：﹕꞉˸፡︓]\s*$", "", label_full).strip()
            if not lab:
                continue
            val = self.matcher.resolve(
                label=lab,
                section="",
                index=1,
                page=None,
                min_fuzzy=1.1,          # ← exact-only
                strict_index=False
            )



            if val is None:
                continue
            new_text = f"{label_full} {val}"
            if dry_run:
                print(f"[DRY][DOCX] underline: '{lab}' → '{val}'")
            else:
                self._set_paragraph_text_keep_simple_format(p, new_text)

    def fill_grid_tables(self, dry_run=False):
        """
        Fill multi-column 'grid' tables where the first row is headers and subsequent rows are values.
        Column j's header text is used as the field name; row i (1-based below header) is the Index.
        Resolves with table 'section' context to improve matching (e.g., 'Authorised Signatories').
        """
        import re
        LETTER_RE = re.compile(r"[A-Za-z]")
        TRAIL_COLON_RE = re.compile(r"[:：﹕꞉˸፡︓]\s*$")

        def _cell_text(cell) -> str:
            return " ".join(p.text for p in cell.paragraphs).strip()

        def _is_placeholder_like(text: str) -> bool:
            if not text:
                return True
            t = re.sub(r"\s+", "", text)
            return not LETTER_RE.search(t)

        def _clean_label(s: str) -> str:
            s = (s or "").strip()
            return TRAIL_COLON_RE.sub("", s).strip()

        def _table_section(tbl) -> str:
            """
            Try to infer a section name for this table:
              1) If row 0 looks like a caption/heading (single merged/shaded cell), use it.
              2) Otherwise, look for the nearest non-empty paragraph before the table.
            """
            # 1) row-0 “caption” like header?
            try:
                row0 = tbl.rows[0]
                row0_text = " ".join(_clean_label(_cell_text(c)) for c in row0.cells).strip()
                if row0_text and len(row0.cells) == 1:
                    return row0_text
            except Exception:
                pass

            # 2) previous visible paragraph
            try:
                # walk up the document XML to find the paragraph immediately preceding this table
                t_elm = tbl._element
                prev = t_elm.getprevious()
                while prev is not None:
                    if prev.tag.endswith("p"):
                        from docx.text.paragraph import Paragraph
                        p = Paragraph(prev, self.doc)
                        txt = (p.text or "").strip()
                        if txt and len(txt) <= 180:
                            # strip trailing colon & parenthetical tails
                            txt = TRAIL_COLON_RE.sub("", txt)
                            txt = re.sub(r"\s*\([^()]*\)\s*$", "", txt).strip()
                            return txt
                    prev = prev.getprevious()
            except Exception:
                pass
            return ""

        for tbl in self.doc.tables:
            if not tbl.rows or len(tbl.rows) < 2:
                continue

            # header labels
            header = tbl.rows[0]
            headers = [_clean_label(_cell_text(c)) for c in header.cells]
            nonempty = [h for h in headers if h]
            if len(nonempty) < max(2, int(0.6 * len(headers))):
                continue  # not a grid

            # if first data row is mostly placeholder, treat as grid
            row1 = tbl.rows[1]
            boxlike = sum(1 for c in row1.cells if _is_placeholder_like(_cell_text(c)))
            if boxlike < max(2, int(0.6 * len(headers))):
                continue

            sec = _table_section(tbl)  # ← key change: resolve with section context

            # body rows (Index = ri)
            for ri in range(1, len(tbl.rows)):
                row = tbl.rows[ri]
                row_index = ri

                for ci in range(min(len(headers), len(row.cells))):
                    field_label = headers[ci]
                    if not field_label:
                        continue

                    cell = row.cells[ci]
                    if not _is_placeholder_like(_cell_text(cell)):
                        continue  # don't overwrite typed text

                    # --- resolve in a cascade, always with section context ---
                    val = self.matcher.resolve(
                        label=field_label, section=sec, index=row_index,
                        page=None, min_fuzzy=1.1, strict_index=True, exact_only=True
                    )
                    if val is None:
                        # exact label in section, no index enforcement
                        val = self.matcher.resolve(
                            label=field_label, section=sec, index=row_index,
                            page=None, min_fuzzy=1.1, strict_index=False, exact_only=True
                        )
                    if val is None:
                        # fuzzy label in section
                        val = self.matcher.resolve(
                            label=field_label, section=sec, index=row_index,
                            page=None, min_fuzzy=0.82, strict_index=False, exact_only=False
                        )
                    if val is None:
                        # last resort: if there is exactly one row in lookup with this label (no index),
                        # use it for all indices in this grid column.
                        v1 = self.matcher.resolve(
                            label=field_label, section=sec, index=1,
                            page=None, min_fuzzy=1.1, strict_index=False, exact_only=True
                        )
                        # accept v1 only if the lookup contains exactly one such row (heuristic inside resolve)
                        val = v1

                    if val is None:
                        continue

                    if dry_run:
                        print(f"[DRY][DOCX] grid: r{row_index} c{ci} '{field_label}' (sec='{sec}') → '{val}'")
                    else:
                        self._write_value_inside_box(cell, str(val))


    def fill_tables(self, dry_run=False):
        """
        Smart table filling:
          • Normal tables: write to right/below/diagonal (preferring box-like cells).
          • Per-column-header tables (e.g. 'Contact 1', 'Contact 2', ... in row 0):
            write into the SAME CELL for each column (because there is no dedicated
            value cell to the right for that column).
        """
        from docx.table import Table

        def _row0_texts(tbl):
            if not tbl.rows:
                return []
            t = []
            for c in tbl.rows[0].cells:
                s = self._first_cell_text(c).strip()
                s = re.sub(r"[:：﹕꞉˸]\s*$", "", s).strip()
                s = re.sub(r"\s+", " ", s)
                t.append(s)
            return t

        def _looks_per_column_headers(texts):
            """Heuristic: >=2 short, distinct headers in row 0 means per-column layout."""
            nonempty = [x for x in texts if x]
            if len(nonempty) < 2:
                return False
            short = [x for x in nonempty if len(x.split()) <= 3]
            if len(short) < max(2, int(0.6 * len(nonempty))):
                return False
            return len(set(short)) >= 2

        for tbl in self.doc.tables:
            nrows = len(tbl.rows)
            if nrows == 0:
                continue

            row0_texts = _row0_texts(tbl)
            is_percol = _looks_per_column_headers(row0_texts)

            for ri, row in enumerate(tbl.rows):
                cells = row.cells
                if len(cells) < 1:
                    continue

                # For normal label→value tables we scan 0,2,4,...
                # For per-column tables every column is its own "pair" and we should
                # NOT write into the cell on the right (it belongs to the next column/Contact).
                col_iter = range(0, len(cells), 1 if is_percol else 2)

                for ci in col_iter:
                    label_cell = cells[ci]
                    lab_raw = self._first_cell_text(label_cell).strip()
                    if not lab_raw:
                        continue

                    # In per-column layout the label often includes just the field name
                    # ('Attn', 'Firm', etc.). Strip trailing colon-like characters.
                    lab = re.sub(r"[:：﹕꞉˸፡︓]\s*$", "", lab_raw).strip()
                    if not lab or len(lab) > 160:
                        continue

                    # Column-index-based index hint: 1-based per logical pair/column
                    idx_guess = (ci // (1 if is_percol else 2)) + 1

                    # strict exact first
                    val = self.matcher.resolve(
                        label=lab,
                        section="",
                        index=idx_guess,
                        page=None,
                        min_fuzzy=1.1,
                        strict_index=False,   # keep False here; these are often not indexed in lookup
                        exact_only=True,
                    )
                    if val is None:
                        # fuzzy fallback, still no index enforcement
                        val = self.matcher.resolve(
                            label=lab,
                            section="",
                            index=idx_guess,
                            page=None,
                            min_fuzzy=0.82,
                            strict_index=False,
                            exact_only=False,
                        )




                    if val is None:
                        continue

                    # ---------- TARGET SELECTION ----------
                    # If we detected a per-column header table, write IN-PLACE (same cell).
                    if is_percol:
                        target = label_cell
                        if dry_run:
                            print(f"[DRY][DOCX] table-percol (same-cell col {ci}): '{lab}' → '{val}'")
                        else:
                            self._write_value_inside_box(target, str(val))
                        continue

                    # Otherwise, use classic right/below/diagonal preference.
                    candidates = []
                    # right
                    if ci + 1 < len(cells):
                        candidates.append(cells[ci + 1])
                    # below
                    if ri + 1 < nrows:
                        candidates.append(tbl.rows[ri + 1].cells[ci])
                        # diagonal
                        if ci + 1 < len(tbl.rows[ri + 1].cells):
                            candidates.append(tbl.rows[ri + 1].cells[ci + 1])

                    if not candidates:
                        continue

                    # Prefer “box-like/empty” targets; keep order otherwise.
                    candidates.sort(key=lambda c: 0 if self._cell_has_box_target(c) else 1)
                    target = candidates[0]

                    if dry_run:
                        print(f"[DRY][DOCX] table-smart: '{lab}' → '{val}'")
                    else:
                        self._write_value_inside_box(target, str(val))

    def fill_matrix_tables(self, dry_run=False):
        """
        Fill 'matrix' tables:
          • Detect the header row by sliding: the first row with ≥2 short labels is the header.
          • Use subsequent blankish rows as Index=1..N data rows.
          • Resolve (Section, Field, Index) using the nearest heading above the table.
        """
        from docx.table import Table
        import re

        def _clean_label(s: str) -> str:
            s = (s or "").strip()
            return re.sub(r"[:：﹕꞉˸]\s*$", "", s).strip()

        def _is_boxlike(cell) -> bool:
            txt = "".join(p.text or "" for p in cell.paragraphs).strip()
            if not txt:
                return True
            vis = re.sub(r"[\s_\-\u2014\.\u2002\u2003\u2007\u2009\u00A0]+", "", txt)
            return not re.search(r"[A-Za-z0-9]", vis)

        def _short_header(txt: str) -> bool:
            t = _clean_label(txt)
            return bool(t) and len(t) <= 60 and len(t.split()) <= 6

        # Find the nearest paragraph text above a given table (best-effort section name)
        def _infer_section_for_table(tbl) -> str:
            body = tbl._element.getparent()
            if body is None:
                return ""
            # Walk siblings backwards from this table to find the last non-empty paragraph
            cur = tbl._element
            while cur is not None:
                cur = cur.getprevious()
                if cur is None:
                    break
                # paragraph
                if cur.tag.endswith("p"):
                    from docx.text.paragraph import Paragraph
                    p = Paragraph(cur, self.doc)
                    t = (p.text or "").strip()
                    if t:
                        return _clean_label(t)
            return ""

        # Normalize spelling and trim
        def _norm_section(s: str) -> str:
            s = (s or "").strip().lower()
            # prefer "authorised" internally
            s = s.replace("authorized", "authorised")
            return s

        for tbl in self.doc.tables:
            if not isinstance(tbl, Table) or not tbl.rows:
                continue

            nrows = len(tbl.rows)
            ncols = len(tbl.rows[0].cells) if nrows else 0
            if ncols < 2:
                continue

            # ---- sliding header detection ----
            header_row_idx = None
            header_labels = []
            r = 0
            while r < nrows:
                labels = [_clean_label(self._first_cell_text(c)) for c in tbl.rows[r].cells]
                short = [h for h in labels if _short_header(h)]
                if len(short) >= 2:
                    header_row_idx = r
                    # use only columns that look like headers (but keep their positions)
                    header_labels = [h if _short_header(h) else "" for h in labels]
                    break
                r += 1

            if header_row_idx is None:
                continue  # not a matrix

            # Count consecutive blankish data rows after header
            data_rows = 0
            rr = header_row_idx + 1
            while rr < nrows:
                blanks = sum(1 for c in tbl.rows[rr].cells if _is_boxlike(c))
                if blanks >= max(2, int(0.6 * ncols)):
                    data_rows += 1
                    rr += 1
                else:
                    break

            if data_rows < 1:
                continue  # header found but no blank data rows → not a matrix

            # Resolve nearest section and normalize
            section_guess = _infer_section_for_table(tbl)
            sec_norm = _norm_section(section_guess)

            # Fill each header column across Index rows
            for col_idx, field_label in enumerate(header_labels):
                field_label = _clean_label(field_label)
                if not field_label:
                    continue  # skip empty header cells

                for idx_val in range(1, data_rows + 1):
                    data_row = tbl.rows[header_row_idx + idx_val]
                    target_cell = data_row.cells[col_idx]

                    # only write into box-like cells; don't overwrite any typed text
                    if not _is_boxlike(target_cell):
                        continue

                    # 1) exact, section-aware
                    val = self.matcher.resolve(
                        label=field_label,
                        section=section_guess,   # keep original case for any exact matches first
                        index=idx_val,
                        page=None,
                        min_fuzzy=1.1,
                        strict_index=True,
                        exact_only=True
                    )

                    # 2) retry with normalized section (authorised/authorized tolerance)
                    if val is None:
                        val = self.matcher.resolve(
                            label=field_label,
                            section=sec_norm,
                            index=idx_val,
                            page=None,
                            min_fuzzy=1.1,
                            strict_index=True,
                            exact_only=True
                        )

                    # 3) relaxed fallback: allow fuzzy and/or sectionless
                    if val is None:
                        val = self.matcher.resolve(
                            label=field_label,
                            section="",             # tolerate missing/unknown section in sheet
                            index=idx_val,
                            page=None,
                            min_fuzzy=0.85,         # fuzzy match allowed
                            strict_index=True,
                            exact_only=False
                        )

                    if val is None or str(val).strip() == "":
                        continue

                    if dry_run:
                        print(f"[DRY][DOCX] table-matrix: Row {idx_val}, Col {col_idx}, '{field_label}' → '{val}'")
                    else:
                        self._write_value_inside_box(target_cell, str(val))



# ---------------------------
# Orchestrator
# ---------------------------
def run_agentic_fill(input_docx: str, lookup_path: str, output_docx: str, dry_run: bool = False) -> int:
    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError("python-docx is required. Install with: pip install python-docx") from e

    # 1) Load lookup
    lookup = LookupAgent(lookup_path)
    rows = lookup.load()
    print(f"🔎 Lookup rows loaded: {len(rows)}")

    # 2) Understand document structure
    doc = Document(input_docx)
    dua = DocumentUnderstandingAgent(doc)
    fields = dua.build_template(dry_run=dry_run)
    print(f"🧭 Discovered fields: {len(fields)}")

    # 3) Match / resolve
    matcher = MatchingAgent(rows)

    # 4) Fill
    # 4) Fill
    filler = FillingAgent(doc, matcher)

    # 0) text placeholders first
    filler.replace_placeholders(dry_run=dry_run)

    # 1) lightweight text/checkbox line features
    filler.fill_inline_checkbox_groups(dry_run=dry_run)
    filler.fill_checkboxes(dry_run=dry_run)
    filler.fill_two_column_checkbox_rows(dry_run=dry_run)

    # 2) RUN MATRIX EARLY (before other table routines)
    #    Detects “Authorized/Authorised Signatories” and the
    #    Last/First/Middle/DOB/POB grid reliably.
    filler.fill_matrix_tables(dry_run=dry_run)

    # 3) other table strategies (these can fill non-matrix tables)
    filler.fill_grid_tables(dry_run=dry_run)
    filler.fill_tables(dry_run=dry_run)

    # 4) final pass on underline lines and horizontal option groups
    filler.fill_underline_lines(dry_run=dry_run)
    filler.fill_checkbox_option_groups(dry_run=dry_run)


    if dry_run:
        print("Dry-run complete. No file written.")
        return 0

    doc.save(output_docx)
    print(f"✅ Wrote: {output_docx}")
    return 1


# ---------------------------
# CLI
# ---------------------------
def main():
    ap = argparse.ArgumentParser(description="Agentic DOCX filler")
    ap.add_argument("--input", required=True, help="Input .docx file")
    ap.add_argument("--lookup", required=True, help="Excel/CSV with Field,Value[,Section,Page,Index]")
    ap.add_argument("--output", required=True, help="Output .docx file")
    ap.add_argument("--dry-run", action="store_true", help="Analyze + print actions without writing file")
    args = ap.parse_args()

    ok = run_agentic_fill(args.input, args.lookup, args.output, dry_run=args.dry_run)
    sys.exit(0 if ok else 1)

if __name__ == "__main__":
    main()