# docx_prefill.py
import argparse, re, sys
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import pandas as pd
from rapidfuzz import fuzz, process
from docx import Document
from docx.text.run import Run
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree  # comes with python-docx

# =========================
# Patterns & constants
# =========================
UNI_BLANK_CLASS = r"\u00A0\u2000-\u200B\t "  # NBSP + EN/EM/thin spaces + tabs + normal space
VISUAL_BLANK_RE = re.compile(rf"^[{UNI_BLANK_CLASS}]{{3,}}$")
INLINE_VISUAL_BLANK_RE = re.compile(rf"[{UNI_BLANK_CLASS}]{{3,}}")
UNDERS_RE = re.compile(r"_{3,}")
LABEL_COLON_RE = re.compile(r"^\s*(?P<label>[A-Za-z0-9().,/\-\&% ]{2,120})\s*[:;]\s*$")
LABEL_INLINE_UNDERS_RE = re.compile(r"^\s*(?P<label>[A-Za-z0-9().,/\-\&% ]{2,120})\s*[:;]\s*_{3,}\s*$")
CHECKBOX_EMPTY = ["□", "☐", "[ ]"]
CHECKBOX_MARKS = ["☑", "☒", "[x]"]
YES_WORDS = {"yes","true","y","checked","1"}
NO_WORDS  = {"no","false","n","unchecked","0"}
FUZZ_THRESH = 86

SYNONYMS = {
    "last name": ["surname", "family name"],
    "first name": ["given name"],
    "middle initial": ["middle name", "mi"],
    "city, st zip": ["city state zip", "city st zip code"],
    "date of birth": ["dob", "date of birth (dd/mm/yyyy)"],
    "citizen of": ["citizenship", "country of citizenship"],
    "country of residence": ["residence country"],
}

# =========================
# Normalization helpers
# =========================
def norm(s: str) -> str:
    return re.sub(r"[^a-z0-9]+"," ", str(s).strip().lower()).strip()

def expand_synonyms(keys: List[str]) -> List[str]:
    out = set(keys)
    for k in list(keys):
        k0 = norm(k)
        for canon, alts in SYNONYMS.items():
            if k0 == canon:
                out.update(alts)
    return list(out)

# =========================
# SDT/content-control helpers (namespace-agnostic)
# =========================
def _iter_sdt_in_cell(cell):
    # namespace-agnostic
    return cell._tc.xpath(".//*[local-name()='sdt']")

def _first_child_by_localname(parent, name):
    res = parent.xpath(f".//*[local-name()='{name}']")
    return res[0] if res else None

def _set_sdt_text_inplace(sdt_elem, value: str) -> bool:
    """
    Fill the first <w:t> inside <w:sdtContent> *in place* so the control remains.
    (Do NOT clear sdtContent here.)
    """
    content = _first_child_by_localname(sdt_elem, "sdtContent")
    if content is None:
        return False

    # 1) Write into an existing <w:t> if present
    t = _first_child_by_localname(content, "t")
    if t is not None:
        t.text = str(value)
        return True

    # 2) Otherwise append a <w:t> under the first <w:r> to keep rPr formatting
    r = _first_child_by_localname(content, "r")
    if r is not None:
        t = OxmlElement(qn("w:t"))
        t.text = str(value)
        r.append(t)
        return True

    # 3) Last resort: add a new p>r>t
    p = OxmlElement(qn("w:p"))
    r = OxmlElement(qn("w:r"))
    t = OxmlElement(qn("w:t"))
    t.text = str(value)
    r.append(t); p.append(r); content.append(p)
    return True


# =========================
# Table/CSV readers
# =========================
def _read_csv_with_fallbacks(path: Path) -> pd.DataFrame:
    encodings = ["utf-8", "utf-8-sig", "cp1252", "latin1"]
    last_err = None
    for enc in encodings:
        try:
            return pd.read_csv(path, dtype=str, encoding=enc)
        except Exception as e:
            last_err = e
    raise RuntimeError(f"Failed to read CSV with common encodings. Last error: {last_err}")

def _read_table_any(path: Path, sheet: Optional[str]) -> pd.DataFrame:
    ext = path.suffix.lower()
    if ext in [".xlsx", ".xls"]:
        return pd.read_excel(path, dtype=str, sheet_name=sheet or 0)
    elif ext == ".csv":
        return _read_csv_with_fallbacks(path)
    else:
        raise RuntimeError(f"Unsupported file extension: {ext}. Use .csv, .xlsx, or .xls")

def _load_structured_df(df: pd.DataFrame) -> Dict:
    # Expect columns: Section | Page | Field | Index | Value | Choices
    df = df.fillna("")
    if "Index" in df.columns:
        df["Index"] = df["Index"].apply(lambda x: int(str(x).strip() or "1"))
    else:
        df["Index"] = 1

    by_key, by_field, fuzz_keys = {}, {}, []
    for _, row in df.iterrows():
        section = str(row.get("Section","")).strip()
        field   = str(row.get("Field","")).strip()
        index   = int(row.get("Index", 1))
        value   = str(row.get("Value",""))
        choices = str(row.get("Choices",""))

        sN = norm(section)
        fN = norm(field)
        by_key[(sN, fN, index)] = (value, choices)
        by_field.setdefault(fN, []).append((value, choices))
        fuzz_keys.append(field)

    return {
        "schema": "structured",
        "by_key": by_key,
        "by_field": by_field,
        "fuzz_keys": expand_synonyms(fuzz_keys),
        "raw_fields": fuzz_keys
    }

def _load_simple_df_row(df: pd.DataFrame, row_index: int) -> Dict:
    if row_index < 0 or row_index >= len(df.index):
        raise IndexError(f"Row {row_index} is out of range 0..{len(df.index)-1}")
    row = df.iloc[row_index].to_dict()
    norm_map = {norm(k): ("" if pd.isna(v) else str(v)) for k, v in row.items()}
    return {
        "schema": "simple",
        "raw_row": {k: ("" if pd.isna(v) else str(v)) for k, v in row.items()},
        "norm_row": norm_map,
        "fuzz_keys": expand_synonyms(list(row.keys()))
    }

def load_table(csv_or_xlsx: Path, sheet: Optional[str], row_index: int) -> Dict:
    df0 = _read_table_any(csv_or_xlsx, sheet)
    cols = [str(c).strip().lower() for c in df0.columns.tolist()]
    structured_cols = {"section","page","field","index","value","choices"}
    if structured_cols.issubset(set(cols)):
        return _load_structured_df(df0)
    else:
        return _load_simple_df_row(df0, row_index)

# =========================
# Run/paragraph replacements (preserve formatting)
# =========================
def copy_run_format(src_run: Run, dst_run: Run):
    """Copy run properties (rPr) so highlight/shading stays."""
    src_rPr = src_run._r.rPr
    if src_rPr is None:
        return
    dst_rPr = dst_run._r.get_or_add_rPr()
    # clear existing dst props
    for child in list(dst_rPr):
        dst_rPr.remove(child)
    # shallow-copy is fine here
    for child in src_rPr:
        dst_rPr.append(child)

def replace_blank_segment_in_run(run: Run, value: str) -> bool:
    """
    If a run contains a blank/underscore segment, replace JUST that segment,
    copying formatting to keep inline grey background if present.
    """
    t = run.text or ""
    # whole-run blanks (NBSP clusters / underscores)
    if VISUAL_BLANK_RE.fullmatch(t) or UNDERS_RE.fullmatch(t):
        run.text = str(value)
        return True
    # inline blank cluster inside the run
    m = INLINE_VISUAL_BLANK_RE.search(t) or UNDERS_RE.search(t)
    if m:
        left, right = t[:m.start()], t[m.end():]
        run.text = left  # keep this run and its formatting
        new_val = run._paragraph.add_run(str(value))
        copy_run_format(run, new_val)
        if right:
            run._paragraph.add_run(right)
        return True
    return False

def replace_blank_in_paragraph(par, value: str) -> bool:
    """
    Replace ONLY blank/underscore segments in the paragraph,
    preserving any inline shading on the replaced run.
    """
    # pass 1: run-by-run targeted replace
    for r in par.runs:
        if replace_blank_segment_in_run(r, value):
            return True
    # pass 2: rebuild paragraph if there is a paragraph-wide cluster
    txt = "".join(r.text for r in par.runs)
    m = INLINE_VISUAL_BLANK_RE.search(txt) or UNDERS_RE.search(txt)
    if m:
        left, right = txt[:m.start()], txt[m.end():]
        for r in par.runs:
            r.text = ""
        par.add_run(left)
        par.add_run(str(value))
        if right:
            par.add_run(right)
        return True
    return False

def replace_blank_runs_in_cell(cell, value: str) -> bool:
    """
    Replace blank/underscore segments INSIDE the cell without clearing the cell,
    so cell shading (grey box) remains. Prefer replacing an existing blank run.
    """
    for p in cell.paragraphs:
        # try run-level replacement to keep inline shading
        for r in p.runs:
            if replace_blank_segment_in_run(r, value):
                return True
        # fallback: paragraph-level replacement
        if replace_blank_in_paragraph(p, value):
            return True
    # ultimate fallback: append value as its own run (keeps cell shading)
    if cell.paragraphs:
        cell.paragraphs[0].add_run(str(value))
    else:
        cell.add_paragraph(str(value))
    return True

# =========================
# Lookups
# =========================
def best_lookup_simple(label: str, values: Dict) -> Tuple[str,str]:
    labN = norm(label)
    if labN in values["norm_row"]:
        return (str(values["norm_row"][labN]), "")
    choice = process.extractOne(label, values["fuzz_keys"], scorer=fuzz.token_set_ratio)
    if choice and choice[1] >= FUZZ_THRESH:
        best_hdr = choice[0]
        for hdr, val in values["raw_row"].items():
            if fuzz.token_set_ratio(best_hdr, hdr) >= FUZZ_THRESH:
                return (str(val), "")
    return ("","")

def best_lookup_structured(
        label: str,
        section_hint: Optional[str],
        values: Dict,
        index_counter: Dict[str,int],
        section_first: bool
) -> Tuple[str,str]:
    labelN = norm(label)
    secN = norm(section_hint or "")
    idx_key = f"{secN}|{labelN}"
    index_counter[idx_key] = index_counter.get(idx_key, 0) + 1
    want_index = index_counter[idx_key]

    by_key = values["by_key"]
    by_field = values["by_field"]

    def try_section():
        if (secN, labelN, want_index) in by_key:
            return by_key[(secN, labelN, want_index)]
        if (secN, labelN, 1) in by_key:
            return by_key[(secN, labelN, 1)]
        return ("","")

    # >>> REPLACE THIS FUNCTION <<<
    def try_field():
        """Respect want_index when section is unknown."""
        arr = by_field.get(labelN, [])
        # try exact occurrence
        if 1 <= want_index <= len(arr):
            v, c = arr[want_index - 1]
            if str(v).strip():
                return (v, c)
        # otherwise do NOT fall back to “first non-empty” here,
        # because that causes duplicates like Street 1 everywhere.
        return ("","")
    # <<< END REPLACEMENT >>>

    if section_first:
        v, c = try_section()
        if v or c: return (str(v), str(c))
        v, c = try_field()
        if v or c: return (str(v), str(c))
    else:
        v, c = try_field()
        if v or c: return (str(v), str(c))
        v, c = try_section()
        if v or c: return (str(v), str(c))

    choice = process.extractOne(label, values["fuzz_keys"], scorer=fuzz.token_set_ratio)
    if choice and choice[1] >= FUZZ_THRESH:
        labN2 = norm(choice[0])
        if labN2 in by_field and by_field[labN2]:
            for v, c in by_field[labN2]:
                if str(v).strip():
                    return (str(v), str(c))
            v, c = by_field[labN2][0]
            return (str(v), str(c))
    return ("","")

# =========================
# Checkbox helpers
# =========================
def checkbox_line_apply(text: str, value: str, choices: str) -> str:
    val = str(value).strip()
    ch  = str(choices).strip()
    low = val.lower()

    if low in YES_WORDS:
        for sym in CHECKBOX_EMPTY:
            if sym in text:
                return text.replace(sym, "☑", 1)
        return text
    if low in NO_WORDS:
        t = text
        for mark in CHECKBOX_MARKS:
            t = t.replace(mark, "□")
        return t

    target = val or ch
    if not target:
        return text

    options = []
    for m in re.finditer(r"(□|☐)\s*([^\s].*?)(?=(\s{2,}|$))", text):
        options.append((m.start(), m.group(2)))
    if not options:
        return text
    best = process.extractOne(target, [o[1] for o in options], scorer=fuzz.token_set_ratio)
    if best and best[1] >= 80:
        idx = [o[1] for o in options].index(best[0])
        start = options[idx][0]
        before, after = text[:start], text[start:]
        after = after.replace("□","☑",1).replace("☐","☑",1)
        return before + after
    return text

# =========================
# DOCX helpers
# =========================
def para_text(par) -> str:
    return "".join(r.text for r in par.runs)

def set_par_text(par, text: str):
    for r in par.runs: r.text = ""
    par.add_run(text)

def is_visual_blank(s: str) -> bool:
    return bool(VISUAL_BLANK_RE.fullmatch(s)) or bool(UNDERS_RE.fullmatch(s))

def cell_text(cell) -> str:
    return "\n".join(p.text for p in cell.paragraphs).strip()

def set_cell_text(cell, text: str):
    for p in cell.paragraphs:
        for r in p.runs: r.text = ""
    cell.paragraphs[0].add_run(text)

# =========================
# Fill logic — paragraphs
# =========================
def fill_paragraphs(
        doc: Document,
        values: Dict,
        section_first: bool,
        section_hint: Optional[str]=None,
        idx_counter: Optional[Dict[str,int]]=None
) -> bool:
    changed = False
    pars = doc.paragraphs
    idx_counter = idx_counter or {}

    def lookup(label: str):
        if values["schema"] == "structured":
            return best_lookup_structured(label, section_hint, values, idx_counter, section_first)
        else:
            v, c = best_lookup_simple(label, values)
            return v, c

    for i, par in enumerate(pars):
        text = para_text(par)

        # Label : ____ on same line
        m = LABEL_INLINE_UNDERS_RE.match(text)
        if m:
            label = m.group("label").strip()
            val, _ = lookup(label)
            if val:
                set_par_text(par, f"{label}: {val}")
                changed = True
                continue

        # Label: + next line blank
        m2 = LABEL_COLON_RE.match(text)
        if m2 and i + 1 < len(pars):
            nxt_par = pars[i+1]
            nxt = para_text(nxt_par)
            if is_visual_blank(nxt) or INLINE_VISUAL_BLANK_RE.search(nxt):
                label = m2.group("label").strip()
                val, _ = lookup(label)
                if val:
                    if replace_blank_in_paragraph(nxt_par, str(val)):
                        changed = True
                        continue

        # Inline visual blanks (e.g., "City, ST Zip     ")
        if INLINE_VISUAL_BLANK_RE.search(text):
            left = text.split(INLINE_VISUAL_BLANK_RE.findall(text)[0])[0]
            label = left.strip().rstrip(":;")
            val, _ = lookup(label)
            if val:
                new = INLINE_VISUAL_BLANK_RE.sub(str(val), text, count=1)
                set_par_text(par, new)
                changed = True
                continue

        # Underscores without colon/semicolon: "... Name ____"
        if UNDERS_RE.search(text) and ":" not in text and ";" not in text:
            left = text.split("_")[0]
            label_guess = " ".join(left.strip().split()[-5:]).rstrip(":;")
            val, _ = lookup(label_guess)
            if val:
                set_par_text(par, UNDERS_RE.sub(str(val), text, count=1))
                changed = True
                continue

        # Checkbox lines
        if any(ch in text for ch in CHECKBOX_EMPTY + CHECKBOX_MARKS):
            lbl_m = re.search(r"([A-Za-z0-9 ().,/\-\&%]{2,120})\s*[:;]\s*$", text)
            label = lbl_m.group(1).strip() if lbl_m else ""
            val, choices = lookup(label) if label else ("","")
            if val or choices:
                new = checkbox_line_apply(text, val, choices)
                if new != text:
                    set_par_text(par, new)
                    changed = True

    return changed

# =========================
# Fill logic — tables
# =========================
def fill_tables(
        doc: Document,
        values: Dict,
        section_first: bool,
        section_hint: Optional[str] = None,
        idx_counter: Optional[Dict[str, int]] = None
) -> bool:
    changed = False
    idx_counter = idx_counter or {}

    def lookup(label: str):
        if values["schema"] == "structured":
            return best_lookup_structured(label, section_hint, values, idx_counter, section_first)
        else:
            v, c = best_lookup_simple(label, values)
            return v, c

    for tbl in doc.tables:
        for row in tbl.rows:
            cells = row.cells
            for ci, cell in enumerate(cells):
                # ---- SDT (content control) fill FIRST ----
                sdt_list = list(_iter_sdt_in_cell(cell))
                if sdt_list:
                    # Prefer label from left neighbor; otherwise try this cell's first line
                    inferred_label = ""
                    if ci > 0:
                        inferred_label = cell_text(cells[ci - 1]).strip().rstrip(":;")
                    if not inferred_label:
                        raw_txt_without_sdt = "\n".join(p.text for p in cell.paragraphs).strip()
                        inferred_label = (raw_txt_without_sdt.split("\n")[0] if raw_txt_without_sdt else "").strip().rstrip(":;")

                    if inferred_label:
                        val, _ = lookup(inferred_label)
                        if val:
                            wrote = False
                            for sdt in sdt_list:
                                if _set_sdt_text_inplace(sdt, str(val)):
                                    wrote = True
                                    break
                            if wrote:
                                changed = True
                                continue  # SDT handled; go next cell

                # From here on, treat as normal cell text
                txt = cell_text(cell)

                # (a) Checkbox in cell
                if any(ch in txt for ch in CHECKBOX_EMPTY + CHECKBOX_MARKS):
                    lbl = re.match(r"^\s*([A-Za-z0-9 ().,/\-\&%]{2,120})\s*[:;]", txt)
                    label = lbl.group(1).strip() if lbl else ""
                    val, choices = lookup(label) if label else ("", "")
                    if val or choices:
                        new = checkbox_line_apply(txt, val, choices)
                        if new != txt:
                            set_cell_text(cell, new)
                            changed = True
                            continue

                # (b) "Label ; ____" or "Label : ____" inside same cell
                m_inline = LABEL_INLINE_UNDERS_RE.match(txt)
                if m_inline:
                    label = m_inline.group("label").strip()
                    val, _ = lookup(label)
                    if val:
                        set_cell_text(cell, f"{label}: {val}")
                        changed = True
                        continue

                # (c) Left label (with/without colon) ➜ right cell is blank/underscores
                if ci + 1 < len(cells):
                    left_label_raw = txt.strip()
                    left_label = left_label_raw.rstrip(":;")
                    right_txt = cell_text(cells[ci + 1])
                    if left_label and (LABEL_COLON_RE.match(txt) or len(left_label) > 1):
                        if (not right_txt) or is_visual_blank(right_txt) or UNDERS_RE.search(right_txt) or INLINE_VISUAL_BLANK_RE.search(right_txt):
                            val, _ = lookup(left_label)
                            if val:
                                # try to keep the grey placeholder in the right cell
                                if not replace_blank_runs_in_cell(cells[ci + 1], str(val)):
                                    set_cell_text(cells[ci + 1], str(val))  # last resort
                                changed = True
                                continue

                # (d) Shaded/grey placeholder in this cell ➜ fill from left neighbor label
                shd = cell._tc.xpath(".//*[local-name()='shd']")
                if shd and (not txt or is_visual_blank(txt) or UNDERS_RE.search(txt) or INLINE_VISUAL_BLANK_RE.search(txt)):
                    if ci > 0:
                        left_label = cell_text(cells[ci - 1]).strip().rstrip(":;")
                        if left_label:
                            val, _ = lookup(left_label)
                            if val:
                                if replace_blank_runs_in_cell(cell, str(val)):
                                    changed = True
                                    continue

                # NEW: shaded cell with NO useful left-label → infer label from the nearest non-empty cell ABOVE (same column)
                if shd and (not txt or is_visual_blank(txt) or UNDERS_RE.search(txt) or INLINE_VISUAL_BLANK_RE.search(txt)):
                    # try left first (existing behavior)
                    left_label = ""
                    if ci > 0:
                        left_label = cell_text(cells[ci - 1]).strip().rstrip(":;")

                    use_above_header = False
                    header_label = ""
                    if not left_label:
                        # scan upward in the same column for a header/label (first non-empty text)
                        col_idx = ci
                        # find our row index (ri) inside this table
                        try:
                            ri = next(ridx for ridx, r in enumerate(tbl.rows) if cell in r.cells)
                        except StopIteration:
                            ri = None
                        if ri is not None:
                            for rj in range(ri - 1, -1, -1):
                                cand = cell_text(tbl.rows[rj].cells[col_idx]).strip()
                                if cand:
                                    header_label = cand.rstrip(":;")
                                    break
                        if header_label:
                            use_above_header = True

                    label_to_use = left_label or header_label
                    if label_to_use:
                        val, _ = lookup(label_to_use)
                        if val:
                            if replace_blank_runs_in_cell(cell, str(val)):
                                changed = True
                                continue


                # (e) Inline visual blanks in cell
                if INLINE_VISUAL_BLANK_RE.search(txt):
                    left = txt.split(INLINE_VISUAL_BLANK_RE.findall(txt)[0])[0]
                    label = left.strip().rstrip(":;")
                    val, _ = lookup(label)
                    if val:
                        new = INLINE_VISUAL_BLANK_RE.sub(str(val), txt, count=1)
                        set_cell_text(cell, new)
                        changed = True
                        continue

                # (f) Underscores without punctuation
                if UNDERS_RE.search(txt) and ":" not in txt and ";" not in txt:
                    left = txt.split("_")[0]
                    guess = " ".join(left.strip().split()[-5:]).rstrip(":;")
                    val, _ = lookup(guess)
                    if val:
                        set_cell_text(cell, UNDERS_RE.sub(str(val), txt, count=1))
                        changed = True
                        continue

    return changed

# =========================
# Pipeline
# =========================
def process_docx(doc_path: Path, values: Dict, section_first: bool) -> Document:
    doc = Document(doc_path)
    section_hint = None
    idx_counter: Dict[str,int] = {}

    changed = False
    changed |= fill_paragraphs(doc, values, section_first, section_hint, idx_counter)
    changed |= fill_tables(doc, values, section_first, section_hint, idx_counter)

    # Final pass: blank paragraph directly under a label line
    pars = doc.paragraphs
    for i, par in enumerate(pars):
        t = para_text(par)
        if VISUAL_BLANK_RE.fullmatch(t) or UNDERS_RE.fullmatch(t):
            if i > 0:
                prev = para_text(pars[i-1]).strip()
                m = LABEL_COLON_RE.match(prev)
                label = m.group("label").strip() if m else prev.rstrip(":;")
                if label:
                    if values["schema"] == "structured":
                        v, _ = best_lookup_structured(label, section_hint, values, idx_counter, section_first)
                    else:
                        v, _ = best_lookup_simple(label, values)
                    if v:
                        set_par_text(par, str(v))
                        changed = True

    return doc

# =========================
# CLI
# =========================
def main():
    ap = argparse.ArgumentParser(description="Prefill .docx from CSV/Excel (structured or simple).")
    ap.add_argument("--doc", "--input", dest="doc", required=True, help="Input DOCX template")
    ap.add_argument("--csv", required=True, help="CSV/XLSX/XLS with fields")
    ap.add_argument("--sheet", default=None, help="Worksheet name for Excel files (optional)")
    ap.add_argument("--row", type=int, default=0, help="Row index for simple CSV (ignored for structured)")
    ap.add_argument("--out", "--output", dest="out", required=True, help="Output DOCX path")
    ap.add_argument("--section-first", action="store_true", help="Prioritize Section+Field over Field-only")
    args = ap.parse_args()

    try:
        values = load_table(Path(args.csv), sheet=args.sheet, row_index=args.row)
    except Exception as e:
        print(f"[CSV/Excel] {e}", file=sys.stderr); sys.exit(2)

    try:
        out_doc = process_docx(Path(args.doc), values, args.section_first)
        out_doc.save(args.out)
        print(f"✅ Wrote {args.out}")
    except Exception as e:
        print(f"[DOCX] {e}", file=sys.stderr); sys.exit(3)

if __name__ == "__main__":
    main()
